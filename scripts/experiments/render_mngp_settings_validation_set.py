from __future__ import annotations

import subprocess
import sys
import winreg
from dataclasses import dataclass
from pathlib import Path

from PIL import Image, ImageChops, ImageStat


ROOT = Path(__file__).resolve().parents[2]
MNOVA_EXE = Path(r"C:\Program Files\Mestrelab Research S.L\MestReNova\MestReNova.exe")
QS_SCRIPT = ROOT / "scripts" / "experiments" / "render_with_mngp_settings.qs"
REGISTRY_SUBKEY = r"Software\Mestrelab Research S.L.\MestReNova\NMR"
REGISTRY_VALUE = "Spectrum Properties"


@dataclass(frozen=True)
class RenderCase:
    label: str
    input_path: Path
    profile_path: Path


def main() -> int:
    output_dir = ROOT / "output" / "mngp_validation" / "settings_api_four_styles"
    output_dir.mkdir(parents=True, exist_ok=True)

    cases = [
        RenderCase(
            "3a_1H_classic",
            ROOT / "examples" / "spectra_2" / "3a" / "da9081_1H" / "fid",
            ROOT / "examples" / "mngp_styles" / "classic_1H.mngp",
        ),
        RenderCase(
            "3a_1H_grid",
            ROOT / "examples" / "spectra_2" / "3a" / "da9081_1H" / "fid",
            ROOT / "examples" / "mngp_styles" / "grid_1H.mngp",
        ),
        RenderCase(
            "3a_13C_classic",
            ROOT / "examples" / "spectra_2" / "3a" / "da9081_13C" / "fid",
            ROOT / "examples" / "mngp_styles" / "classic_13C.mngp",
        ),
        RenderCase(
            "3a_13C_grid",
            ROOT / "examples" / "spectra_2" / "3a" / "da9081_13C" / "fid",
            ROOT / "examples" / "mngp_styles" / "grid_13C.mngp",
        ),
    ]

    original_settings = read_spectrum_properties()
    rendered: list[Path] = []
    try:
        for case in cases:
            print(f"[mngp] {case.label}", flush=True)
            image_path = output_dir / f"{case.label}.png"
            status_path = output_dir / f"{case.label}.status.txt"
            render_one(case, image_path, status_path)
            restore_spectrum_properties(original_settings)
            rendered.append(image_path)
    finally:
        restore_spectrum_properties(original_settings)

    report_path = output_dir / "grid_difference_report.tsv"
    report_path.write_text(grid_report(rendered), encoding="utf-8")

    docx_path = output_dir / "3a_mngp_styles_validation.docx"
    build_docx(docx_path, cases, rendered, report_path)
    print(f"[mngp] docx: {docx_path}", flush=True)
    print(f"[mngp] folder: {output_dir}", flush=True)
    return 0


def render_one(case: RenderCase, image_path: Path, status_path: Path) -> None:
    for path in (MNOVA_EXE, QS_SCRIPT, case.input_path, case.profile_path):
        if not path.exists():
            raise FileNotFoundError(path)
    for path in (image_path, status_path):
        if path.exists():
            path.unlink()

    sf_arg = ",".join(
        [
            "renderWithMNGPSettings",
            mnova_arg(case.input_path),
            mnova_arg(case.profile_path),
            mnova_arg(image_path),
            mnova_arg(status_path),
        ]
    )
    completed = subprocess.run(
        [str(MNOVA_EXE), "-w", str(QS_SCRIPT), "-sf", sf_arg],
        cwd=image_path.parent,
        check=False,
        timeout=180,
        capture_output=True,
        text=True,
        errors="replace",
    )
    (image_path.parent / f"{case.label}.stdout.txt").write_text(completed.stdout or "", encoding="utf-8")
    (image_path.parent / f"{case.label}.stderr.txt").write_text(completed.stderr or "", encoding="utf-8")

    status = status_path.read_text(encoding="utf-8", errors="replace") if status_path.exists() else ""
    if "DONE" not in status:
        raise RuntimeError(
            f"MestReNova did not complete {case.label}; return code {completed.returncode}; status:\n{status}"
        )
    if not image_path.exists() or image_path.stat().st_size < 1000:
        raise RuntimeError(f"MestReNova did not create image for {case.label}: {image_path}")


def read_spectrum_properties() -> bytes:
    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, REGISTRY_SUBKEY, 0, winreg.KEY_READ)
    try:
        value, value_type = winreg.QueryValueEx(key, REGISTRY_VALUE)
    finally:
        winreg.CloseKey(key)
    if value_type != winreg.REG_BINARY:
        raise RuntimeError(f"Unexpected registry type for {REGISTRY_VALUE}: {value_type}")
    return bytes(value)


def restore_spectrum_properties(value: bytes) -> None:
    key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, REGISTRY_SUBKEY, 0, winreg.KEY_SET_VALUE)
    try:
        winreg.SetValueEx(key, REGISTRY_VALUE, 0, winreg.REG_BINARY, value)
    finally:
        winreg.CloseKey(key)


def build_docx(docx_path: Path, cases: list[RenderCase], images: list[Path], report_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches

    document = Document()
    document.add_heading("MestReNova .mngp validation", level=1)
    document.add_paragraph(
        "Images were rendered by applying .mngp through MestReNova Settings('NMR')."
        "Spectrum Properties in a .qs script, without GUI automation."
    )
    for case, image_path in zip(cases, images):
        document.add_heading(case.label, level=2)
        document.add_paragraph(str(case.profile_path))
        document.add_picture(str(image_path), width=Inches(6.7))
    document.add_heading("Image Difference", level=2)
    document.add_paragraph(report_path.read_text(encoding="utf-8"))
    document.save(docx_path)


def grid_report(images: list[Path]) -> str:
    rows = ["nucleus\tdiff_score\tclassic\tgrid"]
    rows.append(f"1H\t{image_difference_score(images[0], images[1]):.6f}\t{images[0].name}\t{images[1].name}")
    rows.append(f"13C\t{image_difference_score(images[2], images[3]):.6f}\t{images[2].name}\t{images[3].name}")
    return "\n".join(rows) + "\n"


def image_difference_score(a: Path, b: Path) -> float:
    with Image.open(a).convert("RGB") as image_a, Image.open(b).convert("RGB") as image_b:
        if image_a.size != image_b.size:
            image_b = image_b.resize(image_a.size)
        diff = ImageChops.difference(image_a, image_b)
        stat = ImageStat.Stat(diff)
        return sum(stat.mean) / (255.0 * 3.0)


def mnova_arg(path: Path) -> str:
    return str(path.resolve()).replace("\\", "/")


if __name__ == "__main__":
    sys.exit(main())
