from __future__ import annotations

import re
import json
import subprocess
import sys
import time
import zipfile
import os
from copy import deepcopy
from pathlib import Path
from tempfile import mkstemp
from xml.etree import ElementTree as ET

import pythoncom
import win32com.client as win32
from docx import Document
from rdkit import Chem
from rdkit.Chem import rdMolDescriptors

from .domain.references import parse_reference_keys
from .domain.reactions import reaction_from_fields
from .models import Compound
from .structure_metadata import extract_structure_metadata_by_row


NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "v": "urn:schemas-microsoft-com:vml",
    "o": "urn:schemas-microsoft-com:office:office",
    "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
}

ET.register_namespace("w", NS["w"])
ET.register_namespace("r", NS["r"])
ET.register_namespace("v", NS["v"])
ET.register_namespace("o", NS["o"])
ET.register_namespace("w10", "urn:schemas-microsoft-com:office:word")


def read_word_compounds(path: str | Path, extract_structure_metadata: bool = False) -> list[Compound]:
    """Read the first Word table and extract text fields plus formulas from OLE structures."""
    path = str(Path(path).resolve())
    structure_metadata = extract_structure_metadata_by_row(path)
    if not extract_structure_metadata:
        return _read_word_compounds_without_com(path, structure_metadata)

    pythoncom.CoInitialize()
    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    doc = None

    try:
        doc = word.Documents.Open(path, False, False, False)
        if doc.Tables.Count == 0:
            raise ValueError("Word input must contain at least one table.")
        table = doc.Tables(1)
        headers = [_cell_text(table.Cell(1, col)) for col in range(1, table.Columns.Count + 1)]
        missing_name_rows: list[int] = []
        rows_data = []

        for row in range(2, table.Rows.Count + 1):
            values = [_cell_text(table.Cell(row, col)) for col in range(1, table.Columns.Count + 1)]
            fields = _map_row(headers, values)
            metadata = structure_metadata.get(row)
            number = fields.get("number") or str(row - 1)
            formula = fields.get("formula", "") or (metadata.formula if metadata else "")
            name = fields.get("name", "") or (metadata.name if metadata else "")
            if extract_structure_metadata:
                formula = formula or _formula_from_structure_in_row(doc, table, row)
            if not name:
                missing_name_rows.append(row)
            rows_data.append((row, fields, number, formula, name, metadata is not None))

        generated_names = _chemdraw_names_for_rows(path, missing_name_rows)
        compounds: list[Compound] = []

        for row, fields, number, formula, name, has_structure in rows_data:
            name = name or generated_names.get(row, "") or f"Compound {number}"
            compounds.append(
                Compound(
                    number=number,
                    name=_format_generated_name(name),
                    preparation=fields.get("preparation", ""),
                    yield_text=fields.get("yield_text", ""),
                    color=fields.get("color", ""),
                    state=fields.get("state", ""),
                    melting_point=_clean_empty_value(fields.get("melting_point", "")),
                    rf=fields.get("rf", ""),
                    formula=formula,
                    hrms_label=fields.get("hrms_label") or "HRMS (ESI-TOF) m/z",
                    hrms_adduct=fields.get("hrms_adduct") or _adduct_from_headers(headers) or "[M+H]+",
                    hrms_found=fields.get("hrms_found", ""),
                    h1_nmr=fields.get("h1_nmr", ""),
                    h1_conditions=fields.get("h1_conditions", ""),
                    h1_spectrum_path=fields.get("h1_spectrum_path", ""),
                    c13_nmr=fields.get("c13_nmr", ""),
                    c13_conditions=fields.get("c13_conditions", ""),
                    c13_spectrum_path=fields.get("c13_spectrum_path", ""),
                    extra_nmr=fields.get("extra_nmr", ""),
                    ir=fields.get("ir", ""),
                    elemental_analysis=_elemental_analysis_from_fields(fields),
                    reaction=reaction_from_fields(fields),
                    references=parse_reference_keys(fields.get("references", "")),
                    has_word_structure=has_structure,
                )
            )

        return compounds
    finally:
        if doc is not None:
            doc.Close(False)
        word.Quit(False)
        pythoncom.CoUninitialize()


def _read_word_compounds_without_com(path: str, structure_metadata) -> list[Compound]:
    document = Document(path)
    if not document.tables:
        raise ValueError("Word input must contain at least one table.")
    table = document.tables[0]
    if not table.rows:
        raise ValueError("Word input table is empty.")
    headers = [_clean_docx_cell_text(cell.text) for cell in table.rows[0].cells]
    rows_data = []
    missing_name_rows = []

    for row_index, row in enumerate(table.rows[1:], start=2):
        values = [_clean_docx_cell_text(cell.text) for cell in row.cells]
        fields = _map_row(headers, values)
        metadata = structure_metadata.get(row_index)
        number = fields.get("number") or str(row_index - 1)
        formula = fields.get("formula", "") or (metadata.formula if metadata else "")
        name = fields.get("name", "") or (metadata.name if metadata else "")
        if not name:
            missing_name_rows.append(row_index)
        rows_data.append((row_index, fields, metadata, number, formula, name))

    generated_names = _chemdraw_names_for_rows(path, missing_name_rows)
    compounds: list[Compound] = []

    for row_index, fields, metadata, number, formula, name in rows_data:
        name = name or generated_names.get(row_index, "") or f"Compound {number}"
        compounds.append(
            Compound(
                number=number,
                name=_format_generated_name(name),
                preparation=fields.get("preparation", ""),
                yield_text=fields.get("yield_text", ""),
                color=fields.get("color", ""),
                state=fields.get("state", ""),
                melting_point=_clean_empty_value(fields.get("melting_point", "")),
                rf=fields.get("rf", ""),
                formula=formula,
                hrms_label=fields.get("hrms_label") or "HRMS (ESI-TOF) m/z",
                hrms_adduct=fields.get("hrms_adduct") or _adduct_from_headers(headers) or "[M+H]+",
                hrms_found=fields.get("hrms_found", ""),
                h1_nmr=fields.get("h1_nmr", ""),
                h1_conditions=fields.get("h1_conditions", ""),
                h1_spectrum_path=fields.get("h1_spectrum_path", ""),
                c13_nmr=fields.get("c13_nmr", ""),
                c13_conditions=fields.get("c13_conditions", ""),
                c13_spectrum_path=fields.get("c13_spectrum_path", ""),
                extra_nmr=fields.get("extra_nmr", ""),
                ir=fields.get("ir", ""),
                elemental_analysis=_elemental_analysis_from_fields(fields),
                reaction=reaction_from_fields(fields),
                references=parse_reference_keys(fields.get("references", "")),
                has_word_structure=metadata is not None,
            )
        )

    return compounds


def _chemdraw_names_for_rows(path: str, rows: list[int], timeout: int = 240) -> dict[int, str]:
    if not rows:
        return {}
    if getattr(sys, "frozen", False):
        command = [sys.executable, "--si-generator-chemdraw-names", path]
    else:
        command = [sys.executable, "-m", "si_generator.chemdraw_names", path]
    command += ["--rows", ",".join(str(row) for row in rows)]
    env = os.environ.copy()
    env["PYTHONIOENCODING"] = "utf-8"
    env["PYTHONUTF8"] = "1"
    try:
        completed = subprocess.run(
            command,
            check=False,
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=timeout,
            env=env,
        )
    except (OSError, subprocess.TimeoutExpired) as exc:
        print(f"[ChemDraw warning] Could not generate structure names: {exc}", flush=True)
        return {}
    if completed.returncode != 0 or not completed.stdout.strip():
        message = (completed.stderr or completed.stdout or "").strip()
        if message:
            print(f"[ChemDraw warning] Could not generate structure names: {message}", flush=True)
        return {}
    try:
        raw = json.loads(completed.stdout)
    except json.JSONDecodeError:
        return {}
    return {int(row): str(name).strip() for row, name in raw.items() if str(name).strip()}


def paste_word_structures(
    input_docx: str | Path,
    output_docx: str | Path,
    compounds: list[Compound],
    main_top_offset_pt: float = 12,
    appendix_top_offset_pt: float = 0,
) -> None:
    """Copy OLE structures from the first column of the input table into output placeholders."""
    _paste_word_structures_in_package(
        Path(input_docx),
        Path(output_docx),
        compounds,
        main_top_offset_pt=main_top_offset_pt,
        appendix_top_offset_pt=appendix_top_offset_pt,
    )


def _paste_word_structures_with_word(input_docx: str | Path, output_docx: str | Path, compounds: list[Compound]) -> None:
    input_docx = str(Path(input_docx).resolve())
    output_docx = str(Path(output_docx).resolve())
    pythoncom.CoInitialize()
    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    source = target = None

    try:
        source = word.Documents.Open(input_docx, False, True, False)
        target = word.Documents.Open(output_docx, False, False, False)
        table = source.Tables(1)
        row_by_number = _row_by_number(table)

        for compound in compounds:
            row = row_by_number.get(compound.number)
            if row is None:
                continue
            structure = _structure_object_in_row(source, table, row)
            if structure is None:
                continue

            source_width = structure.Width
            source_height = structure.Height
            marker = f"[[STRUCTURE:{compound.number}]]"
            finder = target.Content.Find
            finder.ClearFormatting()
            finder.Text = marker
            if not finder.Execute():
                continue

            structure.Range.Copy()
            rng = finder.Parent
            rng.Text = ""
            rng.Select()
            old_shape_count = target.Shapes.Count
            old_inline_count = target.InlineShapes.Count
            word.Selection.Paste()

            if target.Shapes.Count > old_shape_count:
                pasted = target.Shapes(target.Shapes.Count)
            elif target.InlineShapes.Count > old_inline_count:
                pasted = target.InlineShapes(old_inline_count + 1).ConvertToShape()
            else:
                continue
            _apply_structure_layout(pasted, source_width, source_height)

        target.Save()
    finally:
        if target is not None:
            target.Close(False)
        if source is not None:
            source.Close(False)
        word.Quit(False)
        pythoncom.CoUninitialize()


def _cell_text(cell) -> str:
    text = cell.Range.Text
    return text.replace("\r\x07", "").replace("\x07", "").replace("\r", "\n").strip()


def _clean_docx_cell_text(text: str) -> str:
    return text.replace("\x01", "").strip()


def _normalize_header(header: str) -> str:
    header = header.lower().replace("δ", "")
    return re.sub(r"[^a-z0-9+#]+", "", header)


def _map_row(headers: list[str], values: list[str]) -> dict[str, str]:
    result: dict[str, str] = {}
    blank_seen = 0

    for header, value in zip(headers, values):
        key = _normalize_header(header)
        if not key:
            blank_seen += 1
            if blank_seen == 2 and value:
                _split_appearance(value, result)
            continue
        if key in {"no", "number", "compound", "id"}:
            result["number"] = value
        elif key in {"name", "title", "compoundname"}:
            result["name"] = value
        elif "prep" in key or "procedure" in key:
            result["preparation"] = value
        elif "yield" in key:
            result["yield_text"] = value
        elif key in {"1hspectrumpath", "1hspectrum", "1hpath", "hnmrspectrumpath", "hnmrspectrum", "protonpath"}:
            result["h1_spectrum_path"] = value
        elif key in {"13cspectrumpath", "13cspectrum", "13cpath", "cnmrspectrumpath", "cnmrspectrum", "carbonpath"}:
            result["c13_spectrum_path"] = value
        elif key.startswith("1h") or key.startswith("hnmr"):
            result["h1_nmr"] = value
            result["h1_conditions"] = _conditions_from_header(header)
        elif key.startswith("13c") or "cnmr" in key:
            result["c13_nmr"] = value
            result["c13_conditions"] = _conditions_from_header(header)
        elif "hrms" in key:
            result["hrms_label"] = _hrms_label_from_header(header)
            result["hrms_found"] = _first_number(value)
            adduct = re.search(r"\[M[+-][A-Za-z0-9]+\]\+", header)
            if adduct:
                result["hrms_adduct"] = adduct.group(0)
        elif key in {"mp", "meltingpoint"}:
            result["melting_point"] = value
        elif key in {"color", "state", "appearance"}:
            _split_appearance(value, result)
        elif key == "rf":
            result["rf"] = value
        elif key == "formula":
            result["formula"] = value
        elif key == "ir":
            result["ir"] = value
        elif key in {"elementalanalysis", "elementaryanalysis", "analysis", "anal", "ea"}:
            result["elemental_analysis"] = value
        elif key in {"references", "refs", "referencekeys"}:
            result["references"] = value
        elif key in {"targetmmol", "reactiontargetmmol"}:
            result["target_mmol"] = value
        elif re.fullmatch(r"reagent\d+(?:name|role|formula|mw|equiv|equivalents|mmol|massmg|volumeul|density|densitygml|concentration|concentrationm)", key):
            result[_reaction_field_key(key)] = value
        elif "nmr" in key:
            result["extra_nmr"] = value

    return result


def _split_appearance(value: str, result: dict[str, str]) -> None:
    parts = value.rsplit(" ", 1)
    if len(parts) == 2:
        result["color"], result["state"] = parts
    else:
        result["state"] = value


def _elemental_analysis_from_fields(fields: dict[str, str]) -> dict[str, str]:
    value = fields.get("elemental_analysis", "")
    return {"found": value} if value else {}


def _reaction_field_key(key: str) -> str:
    match = re.fullmatch(r"reagent(\d+)(.+)", key)
    if not match:
        return key
    index, field = match.groups()
    aliases = {
        "massmg": "mass_mg",
        "volumeul": "volume_ul",
        "densitygml": "density_g_ml",
        "concentrationm": "concentration_m",
    }
    return f"reagent_{index}_{aliases.get(field, field)}"


def _first_number(value: str) -> str:
    match = re.search(r"\d+(?:\.\d+)?", value)
    return match.group(0) if match else ""


def _adduct_from_headers(headers: list[str]) -> str | None:
    joined = " ".join(headers)
    match = re.search(r"\[M[+-][A-Za-z0-9]+\]\+", joined)
    return match.group(0) if match else None


def _hrms_label_from_header(header: str) -> str:
    cleaned = re.sub(r"\s+", " ", header).strip()
    cleaned = re.sub(r":?\s*\[M[+-][A-Za-z0-9]+\]\+\s*$", "", cleaned).strip()
    cleaned = cleaned.rstrip(":")
    return "HRMS (ESI/Q-TOF) m/z" if cleaned.lower() == "hrms" else cleaned or "HRMS (ESI/Q-TOF) m/z"


def _conditions_from_header(header: str) -> str:
    match = re.search(r"\(([^)]*(?:MHz|Hz)[^)]*)\)", header)
    return match.group(1).strip() if match else ""


def _format_generated_name(name: str) -> str:
    name = name.strip()
    if not name:
        return name
    return name[0].upper() + name[1:]


def _name_from_structure_in_row(doc, table, row: int) -> str:
    names = _chemdraw_names_for_rows(str(Path(doc.FullName).resolve()), [row])
    return names.get(row, "")


def _formula_from_structure_in_row(doc, table, row: int) -> str:
    structure = _structure_object_in_row(doc, table, row)
    if structure is None:
        return ""

    try:
        molfile_data = _get_structure_data(structure, "chemical/x-mdl-molfile")
        molfile_text = molfile_data.decode("latin1", errors="ignore") if isinstance(molfile_data, bytes) else str(molfile_data)
        mol = Chem.MolFromMolBlock(molfile_text, sanitize=True)
        if mol is None:
            return ""
        return rdMolDescriptors.CalcMolFormula(mol)
    except Exception:
        return ""


def _get_structure_data(shape, data_format: str):
    ole = shape.OLEFormat
    try:
        ole.Activate()
    except Exception:
        try:
            ole.DoVerb(0)
        except Exception:
            pass

    for _ in range(5):
        pythoncom.PumpWaitingMessages()
        time.sleep(0.2)

    obj = ole.Object
    return obj.Objects.GetData(data_format)


def _structure_object_in_row(doc, table, row: int):
    start = table.Cell(row, 1).Range.Start
    end = table.Cell(row, table.Columns.Count).Range.End

    for cell_index in range(1, table.Columns.Count + 1):
        cell = table.Cell(row, cell_index)
        for index in range(1, cell.Range.InlineShapes.Count + 1):
            inline_shape = cell.Range.InlineShapes(index)
            try:
                class_type = inline_shape.OLEFormat.ClassType
            except Exception:
                continue
            if class_type:
                return inline_shape

    for index in range(1, doc.Shapes.Count + 1):
        shape = doc.Shapes(index)
        try:
            anchor = shape.Anchor.Start
            class_type = shape.OLEFormat.ClassType
        except Exception:
            continue
        if start <= anchor <= end and class_type:
            return shape
    return None


def _row_by_number(table) -> dict[str, int]:
    rows: dict[str, int] = {}
    for row in range(2, table.Rows.Count + 1):
        number = _cell_text(table.Cell(row, 1))
        if number:
            rows[number] = row
    return rows


def _paste_word_structures_in_package(
    input_docx: Path,
    output_docx: Path,
    compounds: list[Compound],
    main_top_offset_pt: float = 12,
    appendix_top_offset_pt: float = 0,
) -> None:
    source_objects = _source_structure_objects(input_docx)
    if not source_objects:
        return

    compound_numbers = [compound.number for compound in compounds]
    fd, tmp_name = mkstemp(suffix=".docx")
    os.close(fd)
    Path(tmp_name).unlink(missing_ok=True)

    with zipfile.ZipFile(input_docx, "r") as source_zip, zipfile.ZipFile(output_docx, "r") as target_zip:
        document = ET.fromstring(target_zip.read("word/document.xml"))
        rels = ET.fromstring(target_zip.read("word/_rels/document.xml.rels"))
        content_types = ET.fromstring(target_zip.read("[Content_Types].xml"))
        existing_names = set(target_zip.namelist())
        rel_counter = _next_rel_counter(rels)
        object_counter = _next_object_counter(existing_names)
        extra_files: dict[str, bytes] = {}

        for number in compound_numbers:
            source_object = source_objects.get(number)
            if source_object is None:
                continue
            markers = [f"[[STRUCTURE:{number}]]", f"[[SPECTRUM_STRUCTURE:{number}:1H]]", f"[[SPECTRUM_STRUCTURE:{number}:13C]]"]
            for marker in markers:
                while True:
                    run = _run_with_text(document, marker)
                    if run is None:
                        break

                    object_xml = deepcopy(source_object["object"])
                    image_rel = object_xml.find(".//v:imagedata", NS)
                    ole_rel = object_xml.find(".//o:OLEObject", NS)
                    if image_rel is None or ole_rel is None:
                        break

                    image_ext = Path(source_object["image_target"]).suffix or ".emf"
                    ole_ext = Path(source_object["ole_target"]).suffix or ".bin"
                    new_image_name = f"word/media/si_structure_{object_counter}{image_ext}"
                    new_ole_name = f"word/embeddings/si_structure_{object_counter}{ole_ext}"
                    object_counter += 1

                    image_rid = f"rIdSI{rel_counter}"
                    rel_counter += 1
                    ole_rid = f"rIdSI{rel_counter}"
                    rel_counter += 1

                    image_rel.set(f"{{{NS['r']}}}id", image_rid)
                    ole_rel.set(f"{{{NS['r']}}}id", ole_rid)
                    _apply_object_layout(
                        object_xml,
                        source_object["width"],
                        source_object["height"],
                        in_front=marker.startswith("[[SPECTRUM_STRUCTURE:"),
                        top_offset_pt=appendix_top_offset_pt if marker.startswith("[[SPECTRUM_STRUCTURE:") else main_top_offset_pt,
                    )

                    _add_relationship(rels, image_rid, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image", new_image_name.removeprefix("word/"))
                    _add_relationship(rels, ole_rid, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/oleObject", new_ole_name.removeprefix("word/"))
                    _ensure_default_content_type(content_types, image_ext.lstrip("."), "image/x-emf")
                    _ensure_default_content_type(content_types, ole_ext.lstrip("."), "application/vnd.openxmlformats-officedocument.oleObject")

                    extra_files[new_image_name] = source_zip.read("word/" + source_object["image_target"])
                    extra_files[new_ole_name] = source_zip.read("word/" + source_object["ole_target"])
                    _replace_run_text_with_object(run, object_xml)

        with zipfile.ZipFile(tmp_name, "w", zipfile.ZIP_DEFLATED) as out_zip:
            for item in target_zip.infolist():
                if item.filename in {"word/document.xml", "word/_rels/document.xml.rels", "[Content_Types].xml"}:
                    continue
                out_zip.writestr(item, target_zip.read(item.filename))
            _drop_stale_ignorable_prefixes(document)
            out_zip.writestr("word/document.xml", ET.tostring(document, encoding="utf-8", xml_declaration=True))
            out_zip.writestr("word/_rels/document.xml.rels", ET.tostring(rels, encoding="utf-8", xml_declaration=True))
            out_zip.writestr("[Content_Types].xml", ET.tostring(content_types, encoding="utf-8", xml_declaration=True))
            for name, data in extra_files.items():
                out_zip.writestr(name, data)

    Path(tmp_name).replace(output_docx)


def _source_structure_objects(input_docx: Path) -> dict[str, dict]:
    with zipfile.ZipFile(input_docx, "r") as archive:
        document = ET.fromstring(archive.read("word/document.xml"))
        rels = _rels_map(archive.read("word/_rels/document.xml.rels"))
        table = document.find(".//w:tbl", NS)
        if table is None:
            return {}

        objects: dict[str, dict] = {}
        for row in table.findall("w:tr", NS)[1:]:
            cells = row.findall("w:tc", NS)
            if not cells:
                continue
            number = "".join(cells[0].itertext()).strip()
            object_xml = row.find(".//w:object", NS)
            image_rel = object_xml.find(".//v:imagedata", NS) if object_xml is not None else None
            ole_rel = object_xml.find(".//o:OLEObject", NS) if object_xml is not None else None
            if not number or object_xml is None or image_rel is None or ole_rel is None:
                continue

            image_target = rels.get(image_rel.attrib.get(f"{{{NS['r']}}}id", ""))
            ole_target = rels.get(ole_rel.attrib.get(f"{{{NS['r']}}}id", ""))
            if not image_target or not ole_target:
                continue
            width, height = _object_size(object_xml)
            objects[number] = {
                "object": object_xml,
                "image_target": image_target,
                "ole_target": ole_target,
                "width": width,
                "height": height,
            }
        return objects


def _rels_map(data: bytes) -> dict[str, str]:
    root = ET.fromstring(data)
    return {rel.attrib["Id"]: rel.attrib["Target"] for rel in root if "Id" in rel.attrib and "Target" in rel.attrib}


def _object_size(object_xml) -> tuple[str, str]:
    shape = object_xml.find(".//v:shape", NS)
    style = shape.attrib.get("style", "") if shape is not None else ""
    width = _style_value(style, "width") or "98.35pt"
    height = _style_value(style, "height") or "64.5pt"
    return width, height


def _style_value(style: str, key: str) -> str:
    match = re.search(rf"(?:^|;){re.escape(key)}:([^;]+)", style)
    return match.group(1).strip() if match else ""


def _apply_object_layout(object_xml, width: str, height: str, in_front: bool = False, top_offset_pt: float = 0) -> None:
    shape = object_xml.find(".//v:shape", NS)
    if shape is None:
        return
    top_margin = f"{top_offset_pt:g}pt"
    shape.set(
        "style",
        f"position:absolute;left:0;text-align:left;margin-left:0pt;margin-top:{top_margin};"
        f"width:{width};height:{height};z-index:251659264",
    )
    shape.set(f"{{{NS['o']}}}allowoverlap", "t" if in_front else "f")
    wrap = shape.find("w10:wrap", {"w10": "urn:schemas-microsoft-com:office:word"})
    if wrap is None:
        wrap = ET.SubElement(shape, "{urn:schemas-microsoft-com:office:word}wrap")
    if in_front:
        wrap.set("type", "none")
        wrap.attrib.pop("side", None)
    else:
        wrap.set("type", "square")
        wrap.set("side", "right")


def _run_with_text(document, text: str):
    for run in document.findall(".//w:r", NS):
        if text in "".join(run.itertext()):
            return run
    return None


def _replace_run_text_with_object(run, object_xml) -> None:
    for child in list(run):
        run.remove(child)
    run.append(object_xml)


def _next_rel_counter(rels) -> int:
    max_id = 0
    for rel in rels:
        rel_id = rel.attrib.get("Id", "")
        match = re.fullmatch(r"rId(?:SI)?(\d+)", rel_id)
        if match:
            max_id = max(max_id, int(match.group(1)))
    return max_id + 1


def _next_object_counter(names: set[str]) -> int:
    counter = 1
    while f"word/embeddings/si_structure_{counter}.bin" in names or f"word/media/si_structure_{counter}.emf" in names:
        counter += 1
    return counter


def _add_relationship(rels, rel_id: str, rel_type: str, target: str) -> None:
    rel = ET.SubElement(rels, f"{{{NS['rel']}}}Relationship")
    rel.set("Id", rel_id)
    rel.set("Type", rel_type)
    rel.set("Target", target)


def _ensure_default_content_type(content_types, extension: str, content_type: str) -> None:
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    for item in content_types.findall(f"{{{ct_ns}}}Default"):
        if item.attrib.get("Extension", "").lower() == extension.lower():
            return
    default = ET.SubElement(content_types, f"{{{ct_ns}}}Default")
    default.set("Extension", extension)
    default.set("ContentType", content_type)


def _drop_stale_ignorable_prefixes(document) -> None:
    ignorable_attr = "{http://schemas.openxmlformats.org/markup-compatibility/2006}Ignorable"
    if ignorable_attr in document.attrib:
        document.attrib.pop(ignorable_attr, None)


def _apply_structure_layout(shape, width: float, height: float) -> None:
    shape.LockAspectRatio = True
    shape.Width = width
    shape.Height = height
    shape.WrapFormat.Type = 0
    shape.WrapFormat.DistanceTop = 0
    shape.WrapFormat.DistanceBottom = 0
    shape.WrapFormat.DistanceLeft = 0
    shape.WrapFormat.DistanceRight = 8
    shape.RelativeHorizontalPosition = 2
    shape.RelativeVerticalPosition = 2
    shape.Left = 0
    shape.Top = 0


def _clean_empty_value(value: str) -> str:
    value = value.strip()
    return "" if value in {"-", "—", "–"} else value
