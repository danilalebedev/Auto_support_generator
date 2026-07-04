from __future__ import annotations

import json
import shutil
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document

from ...domain.issues import compound_issue_counts, count_issues
from ...domain.manifest import load_manifest, manifest_has_errors
from ...domain.patching import set_manifest_output_paths, support_docx_from_manifest
from ...domain.requests import GenerateSIRequest
from ...input_table import read_compounds
from ...word_input import read_word_compounds
from ..state import AddCompoundsState


def load_add_manifest_node(state: AddCompoundsState) -> dict:
    request = state["request"]
    artifacts = {**state.get("artifacts", {}), "source_manifest": str(request.manifest_path)}
    try:
        manifest = load_manifest(request.manifest_path)
    except (OSError, json.JSONDecodeError, ValueError) as exc:
        issues = [
            *state.get("issues", []),
            _issue(
                "MANIFEST_LOAD_FAILED",
                "error",
                f"could not load manifest: {exc}",
                path=str(request.manifest_path),
            ),
        ]
        return {"manifest": {}, "artifacts": artifacts, "issues": issues, "status": "fail"}
    return {"manifest": manifest, "artifacts": artifacts}


def read_new_compounds_node(state: AddCompoundsState) -> dict:
    request = state["request"]
    try:
        if request.input_kind == "word":
            compounds = read_word_compounds(request.input_path, extract_structure_metadata=False)
        else:
            compounds = read_compounds(request.input_path)
    except Exception as exc:
        issues = [
            *state.get("issues", []),
            _issue("ADD_COMPOUNDS_INPUT_READ_FAILED", "error", f"could not read new compound table: {exc}", path=str(request.input_path)),
        ]
        return {"new_compounds": [], "issues": issues, "status": "fail"}
    return {"new_compounds": compounds, "artifacts": {**state.get("artifacts", {}), "new_compound_table": str(request.input_path)}}


def check_duplicate_compound_numbers_node(state: AddCompoundsState) -> dict:
    issues = list(state.get("issues", []))
    if manifest_has_errors(issues):
        return {"issues": issues, "status": "fail"}

    existing_numbers = _manifest_numbers(state.get("manifest", {}))
    duplicate_numbers = sorted(
        {
            compound.number.strip()
            for compound in state.get("new_compounds", [])
            if compound.number.strip() and compound.number.strip() in existing_numbers
        }
    )
    if not duplicate_numbers:
        return {"issues": issues}

    issues.extend(
        _issue(
            "DUPLICATE_COMPOUND_NUMBER",
            "error",
            f"new compound number '{number}' already exists in the manifest.",
            compound_id=existing_numbers[number],
        )
        for number in duplicate_numbers
    )
    return {
        "issues": issues,
        "status": "fail",
        "add_result": {
            "duplicate_numbers": duplicate_numbers,
            "added_ids": [],
            "generated_support_docx": "",
        },
    }


def generate_new_support_node(state: AddCompoundsState) -> dict:
    issues = list(state.get("issues", []))
    if manifest_has_errors(issues):
        return {"issues": issues, "status": "fail"}

    request = state["request"]
    temp_dir = request.output_docx.parent / "_add_compounds_work" / (state.get("run_id") or "run")
    temp_output = temp_dir / "new_compounds.docx"
    temp_dir.mkdir(parents=True, exist_ok=True)

    generate_request = GenerateSIRequest(
        input_path=request.input_path,
        input_kind=request.input_kind,
        output_path=temp_output,
        template_docx=request.template_docx,
        references_path=request.references_path,
        spectra_source=request.resolved_spectra_source,
        mnova_exe=request.mnova_exe,
        no_extract_nmr=request.no_extract_nmr,
        insert_spectra_as=request.insert_spectra_as,
        peak_threshold_fraction=request.peak_threshold_fraction,
        peak_threshold_fraction_1h=request.peak_threshold_fraction_1h,
        peak_threshold_fraction_13c=request.peak_threshold_fraction_13c,
        baseline_mode=request.baseline_mode,
        baseline_apply_1h=request.baseline_apply_1h,
        baseline_apply_13c=request.baseline_apply_13c,
        baseline_poly_order=request.baseline_poly_order,
        whittaker_lambda=request.whittaker_lambda,
        whittaker_asymmetry=request.whittaker_asymmetry,
        generate_loadings=request.generate_loadings,
        no_check_support=request.no_check_support,
    )

    from ...workflows.generate_si import output_path_from_state, run_generate_si

    generated_state = run_generate_si(generate_request)
    generated_output = output_path_from_state(generated_state)
    artifacts = {
        **state.get("artifacts", {}),
        "generated_support_docx": str(generated_output),
    }
    if generated_state.get("artifacts", {}).get("manifest"):
        artifacts["generated_manifest"] = generated_state["artifacts"]["manifest"]
    issues.extend(generated_state.get("issues", []))
    if manifest_has_errors(generated_state.get("issues", [])):
        issues.append(
            _issue(
                "ADD_COMPOUNDS_GENERATION_FAILED",
                "error",
                "new compound generation reported errors; existing support was not patched.",
                path=str(generated_output),
            )
        )
        return {"new_generate_state": generated_state, "artifacts": artifacts, "issues": issues, "status": "fail"}
    return {"new_generate_state": generated_state, "artifacts": artifacts, "issues": issues}


def append_new_blocks_node(state: AddCompoundsState) -> dict:
    issues = list(state.get("issues", []))
    if manifest_has_errors(issues):
        return {"issues": issues, "status": "fail"}

    request = state["request"]
    output_docx = Path(request.output_docx)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    try:
        source_docx = support_docx_from_manifest(state.get("manifest", {}), request.manifest_path, request.support_docx)
        generated_docx = Path(state.get("artifacts", {}).get("generated_support_docx", ""))
        if not source_docx.exists():
            raise FileNotFoundError(source_docx)
        if not generated_docx.exists():
            raise FileNotFoundError(generated_docx)
        if source_docx.resolve() == output_docx.resolve():
            issues.append(
                _issue(
                    "ADD_COMPOUNDS_IN_PLACE_OUTPUT",
                    "error",
                    "output DOCX must be different from the existing support DOCX.",
                    path=str(output_docx),
                )
            )
            return {"issues": issues, "status": "fail"}
        _append_generated_docx_text(source_docx, generated_docx, output_docx)
    except Exception as exc:
        issues.append(
            _issue(
                "ADD_COMPOUNDS_DOCX_APPEND_FAILED",
                "error",
                f"could not append new compound blocks: {exc}",
                path=str(output_docx),
            )
        )
        return {"issues": issues, "status": "fail"}

    issues.append(
        _issue(
            "ADD_COMPOUNDS_TEXT_ONLY_MERGE",
            "warning",
            "new compound blocks were appended as text-only; generated structures/spectra remain in the generated SI artifacts.",
            path=str(output_docx),
        )
    )
    return {"artifacts": {**state.get("artifacts", {}), "support_docx": str(output_docx)}, "issues": issues}


def write_add_manifest_node(state: AddCompoundsState) -> dict:
    issues = list(state.get("issues", []))
    if manifest_has_errors(issues):
        return {"issues": issues, "status": "fail"}

    request = state["request"]
    output_docx = Path(request.output_docx)
    output_manifest = output_docx.with_suffix(".manifest.json")
    generated_manifest = _generated_manifest(state)
    merged_manifest, add_result = _merge_manifest(
        state.get("manifest", {}),
        generated_manifest,
        run_id=state.get("run_id", ""),
        source_manifest=request.manifest_path,
        output_docx=output_docx,
        output_manifest=output_manifest,
    )
    output_manifest.parent.mkdir(parents=True, exist_ok=True)
    output_manifest.write_text(json.dumps(merged_manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    artifacts = {
        **state.get("artifacts", {}),
        "manifest": str(output_manifest),
    }
    return {"manifest": merged_manifest, "add_result": add_result, "artifacts": artifacts}


def write_add_compounds_report_node(state: AddCompoundsState) -> dict:
    request = state["request"]
    output_docx = Path(request.output_docx)
    report_path = output_docx.with_suffix(".add_report.json")
    issues = list(state.get("issues", []))
    status = "fail" if manifest_has_errors(issues) or state.get("status") == "fail" else "pass"
    report = {
        "run_id": state.get("run_id", ""),
        "status": status,
        "source_manifest": str(request.manifest_path),
        "source_support_docx": str(request.support_docx) if request.support_docx else "",
        "new_compound_table": str(request.input_path),
        "output_docx": str(output_docx),
        "strict_artifacts": request.strict_artifacts,
        "add_result": state.get("add_result", _empty_add_result()),
        "issue_counts": count_issues(issues),
        "compound_issue_counts": compound_issue_counts(issues),
        "issues": issues,
        "artifacts": {**state.get("artifacts", {}), "add_report": str(report_path)},
    }
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return {"status": status, "issues": issues, "artifacts": report["artifacts"]}


def route_add_compounds_after_load(state: AddCompoundsState) -> str:
    return "fail" if manifest_has_errors(state.get("issues", [])) or state.get("status") == "fail" else "continue"


def route_add_compounds_after_duplicate_check(state: AddCompoundsState) -> str:
    return "fail" if manifest_has_errors(state.get("issues", [])) or state.get("status") == "fail" else "continue"


def route_add_compounds_after_generation(state: AddCompoundsState) -> str:
    return "fail" if manifest_has_errors(state.get("issues", [])) or state.get("status") == "fail" else "continue"


def _append_generated_docx_text(source_docx: Path, generated_docx: Path, output_docx: Path) -> None:
    if source_docx.resolve() != output_docx.resolve():
        shutil.copy2(source_docx, output_docx)
    target = Document(output_docx)
    source = Document(generated_docx)
    if target.paragraphs and target.paragraphs[-1].text.strip():
        target.add_paragraph("")
    for paragraph in source.paragraphs:
        text = paragraph.text
        if not text.strip():
            target.add_paragraph("")
            continue
        target.add_paragraph(text)
    target.save(output_docx)


def _generated_manifest(state: AddCompoundsState) -> dict[str, Any]:
    generated_state = state.get("new_generate_state", {})
    if isinstance(generated_state, dict) and isinstance(generated_state.get("manifest"), dict):
        return generated_state["manifest"]
    manifest_path = state.get("artifacts", {}).get("generated_manifest")
    if manifest_path:
        return load_manifest(manifest_path)
    return {}


def _merge_manifest(
    old_manifest: dict[str, Any],
    new_manifest: dict[str, Any],
    *,
    run_id: str,
    source_manifest: Path,
    output_docx: Path,
    output_manifest: Path,
) -> tuple[dict[str, Any], dict[str, Any]]:
    merged = deepcopy(old_manifest)
    merged.setdefault("order", [])
    merged.setdefault("compounds", {})
    existing_ids = {str(item) for item in merged.get("compounds", {})}
    added_ids: list[str] = []

    for new_id in new_manifest.get("order", []):
        raw_entry = (new_manifest.get("compounds", {}) or {}).get(new_id)
        if not isinstance(raw_entry, dict):
            continue
        merged_id = _unique_compound_id(str(new_id), existing_ids)
        existing_ids.add(merged_id)
        entry = deepcopy(raw_entry)
        entry["id"] = merged_id
        entry["docx_block_id"] = f"compound:{merged_id}"
        entry["docx_bookmark"] = ""
        snapshot = entry.get("domain_snapshot")
        if isinstance(snapshot, dict):
            snapshot["id"] = merged_id
        merged["compounds"][merged_id] = entry
        merged["order"].append(merged_id)
        added_ids.append(merged_id)

    set_manifest_output_paths(merged, support_docx=output_docx, manifest_path=output_manifest)
    history = merged.setdefault("add_compounds_history", [])
    if not isinstance(history, list):
        merged["add_compounds_history"] = history = []
    add_result = {
        "added_ids": added_ids,
        "duplicate_numbers": [],
        "generated_support_docx": new_manifest.get("output_paths", {}).get("support_docx")
        or new_manifest.get("artifacts", {}).get("support_docx", ""),
    }
    history.append(
        {
            "run_id": run_id,
            "source_manifest": str(source_manifest),
            "output_manifest": str(output_manifest),
            "output_docx": str(output_docx),
            "result": add_result,
        }
    )
    return merged, add_result


def _manifest_numbers(manifest: dict[str, Any]) -> dict[str, str]:
    numbers: dict[str, str] = {}
    for compound_id, compound in (manifest.get("compounds", {}) or {}).items():
        if not isinstance(compound, dict):
            continue
        number = str(compound.get("number") or "").strip()
        if number:
            numbers[number] = str(compound_id)
    return numbers


def _unique_compound_id(base_id: str, existing_ids: set[str]) -> str:
    candidate = base_id or "compound"
    if candidate not in existing_ids:
        return candidate
    counter = 1
    while f"added_{candidate}_{counter}" in existing_ids:
        counter += 1
    return f"added_{candidate}_{counter}"


def _empty_add_result() -> dict[str, Any]:
    return {
        "added_ids": [],
        "duplicate_numbers": [],
        "generated_support_docx": "",
    }


def _issue(
    code: str,
    severity: str,
    message: str,
    *,
    compound_id: str = "",
    path: str = "",
) -> dict[str, str]:
    issue = {"code": code, "severity": severity, "message": message}
    if compound_id:
        issue["compound_id"] = compound_id
    if path:
        issue["path"] = path
    return issue
