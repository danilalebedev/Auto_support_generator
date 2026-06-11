from __future__ import annotations

import hashlib
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Pt

from .chemistry import parse_formula
from .domain.elemental_analysis import calculate_elemental_analysis_block, found_from_block
from .domain.massspec import build_hrms_block, hrms_adduct_text, hrms_found_text, hrms_label_text
from .domain.references import format_reference
from .domain.reactions import calculate_reaction_loadings
from .domain.types import JournalProfile
from .domain.types import ReferenceStore
from .models import Compound
from .render.document_model import build_si_document_model
from .render.si_document import DocumentBlock, SIDocument
from .style_config import DEFAULT_STYLE_CONFIG, apply_paragraph_style, apply_run_style, config_get


def build_document(
    compounds: list[Compound],
    output_path: str | Path,
    style_config: dict[str, Any] | None = None,
    template_path: str | Path | None = None,
    journal_profile: JournalProfile | None = None,
    reference_store: ReferenceStore | None = None,
) -> Path:
    return build_document_from_model(
        build_si_document_model(compounds, journal_profile, reference_store),
        output_path,
        style_config=style_config,
        template_path=template_path,
    )


def build_document_from_model(
    document_model: SIDocument,
    output_path: str | Path,
    style_config: dict[str, Any] | None = None,
    template_path: str | Path | None = None,
) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    style_config = style_config or DEFAULT_STYLE_CONFIG

    document = Document(str(template_path)) if template_path else Document()
    if template_path:
        document._body.clear_content()
    else:
        _configure_styles(document)

    _render_document_model(document, document_model, style_config)
    document.save(output_path)
    return output_path


def _render_document_model(document: Document, document_model: SIDocument, style_config: dict[str, Any]) -> None:
    sections = document_model.get("sections", [])
    for section in sections:
        blocks = section.get("blocks", [])
        if section.get("id") == "compound_descriptions":
            _render_compound_description_blocks(document, blocks, style_config)
        elif section.get("id") == "spectra_appendix" and blocks:
            _render_spectra_appendix_blocks(document, blocks, style_config)
        elif section.get("id") == "references" and blocks:
            _render_reference_blocks(document, blocks, style_config)


def _render_compound_description_blocks(document: Document, blocks: list[DocumentBlock], style_config: dict[str, Any]) -> None:
    for index, block in enumerate(blocks):
        if index:
            document.add_paragraph()
        first_index = len(document.paragraphs)
        _add_compound_block(document, block["content"], style_config)
        _add_bookmark_range(document.paragraphs[first_index], document.paragraphs[-1], block.get("bookmark", ""))


def _render_spectra_appendix_blocks(document: Document, blocks: list[DocumentBlock], style_config: dict[str, Any]) -> None:
    document.add_page_break()
    for index, block in enumerate(blocks):
        if index:
            document.add_page_break()
        first_index = len(document.paragraphs)
        _add_spectrum_page(document, block, style_config)
        _add_bookmark_range(document.paragraphs[first_index], document.paragraphs[-1], block.get("bookmark", ""))


def _render_reference_blocks(document: Document, blocks: list[DocumentBlock], style_config: dict[str, Any]) -> None:
    document.add_page_break()
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(0)
    apply_paragraph_style(title, style_config, "references.title")
    run = title.add_run("References")
    apply_run_style(run, style_config, "references.title")
    for block in blocks:
        content = block.get("content", {})
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        apply_paragraph_style(paragraph, style_config, "references.body")
        _add_bookmark_range(paragraph, paragraph, block.get("bookmark", ""))
        text = format_reference(content["reference"], int(content["index"]))
        paragraph.add_run(text)


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for section in document.sections:
        section.top_margin = Pt(56.7)
        section.bottom_margin = Pt(56.7)
        section.left_margin = Pt(85.05)
        section.right_margin = Pt(42.5)


def _add_compound_block(document: Document, compound: Compound, style_config: dict[str, Any]) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(0)
    apply_paragraph_style(title, style_config, "compound.title")
    title_run = title.add_run(f"{compound.name} ")
    apply_run_style(title_run, style_config, "compound.title")
    number_run = title.add_run(compound.label)
    apply_run_style(number_run, style_config, "compound.number")

    summary_parts = []
    if compound.preparation:
        summary_parts.append(compound.preparation.rstrip("."))
    if compound.yield_text:
        summary_parts.append(f"Yield {compound.yield_text}")
    appearance = " ".join(part for part in [compound.color, compound.state] if part)
    if appearance:
        summary_parts.append(appearance)
    if compound.melting_point:
        summary_parts.append(f"mp {compound.melting_point}")
    if compound.rf:
        summary_parts.append(compound.rf)

    if compound.structure_path or compound.has_word_structure:
        _add_structure_paragraph(document, compound)

    if summary_parts:
        text = "; ".join(summary_parts) + "." if summary_parts else ""
        _add_summary_paragraph(document, text, style_config)

    if compound.reaction:
        _add_reaction_loadings_line(document, compound, style_config)

    if compound.h1_nmr:
        _add_nmr_line(document, "1H NMR", compound.h1_nmr, compound.h1_conditions, style_config)
    if compound.c13_nmr:
        _add_nmr_line(document, "13C{1H} NMR", compound.c13_nmr, compound.c13_conditions, style_config)
    if compound.extra_nmr:
        _add_sentence_paragraph(document, compound.extra_nmr, style_config)
    if compound.formula and hrms_found_text(compound.hrms, compound.hrms_found):
        _add_hrms_line(document, compound, style_config)
    if compound.elemental_analysis:
        _add_elemental_analysis_line(document, compound, style_config)
    if compound.ir:
        _add_ir_line(document, compound.ir, style_config)
    if compound.nmr_check_warning:
        _add_nmr_warning(document, compound.nmr_check_warning)


def _add_sentence_paragraph(document: Document, text: str, style_config: dict[str, Any]):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    _add_chem_text_runs(paragraph, text, style_config)
    return paragraph


def _add_structure_paragraph(document: Document, compound: Compound):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.add_run(f"[[STRUCTURE:{compound.number}]]")
    return paragraph


def _add_summary_paragraph(document: Document, text: str, style_config: dict[str, Any]):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    apply_paragraph_style(paragraph, style_config, "compound.summary")
    paragraph.add_run(text)
    return paragraph


def _add_nmr_line(document: Document, label: str, text: str, conditions: str, style_config: dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    apply_paragraph_style(paragraph, style_config, "nmr")
    if label == "1H NMR":
        run = _add_isotope_run(paragraph, "1", style_config)
        apply_run_style(run, style_config, "nmr.label")
        run = paragraph.add_run("H NMR")
        apply_run_style(run, style_config, "nmr.label")
        if conditions:
            _add_conditions_runs(paragraph, conditions, style_config)
        paragraph.add_run(" ")
    elif label == "13C{1H} NMR":
        run = _add_isotope_run(paragraph, "13", style_config)
        apply_run_style(run, style_config, "nmr.label")
        run = paragraph.add_run("C{")
        apply_run_style(run, style_config, "nmr.label")
        run = _add_isotope_run(paragraph, "1", style_config)
        apply_run_style(run, style_config, "nmr.label")
        run = paragraph.add_run("H} NMR")
        apply_run_style(run, style_config, "nmr.label")
        if conditions:
            _add_conditions_runs(paragraph, conditions, style_config)
        paragraph.add_run(" ")
    else:
        run = paragraph.add_run(f"{label}: ")
        apply_run_style(run, style_config, "nmr.label")
    _add_chem_text_runs(paragraph, text.strip(), style_config, "nmr.body")


def _add_ir_line(document: Document, text: str, style_config: dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    apply_paragraph_style(paragraph, style_config, "ir")
    run = paragraph.add_run("IR (KBr, cm")
    apply_run_style(run, style_config, "ir.label")
    run = paragraph.add_run("-1")
    run.font.superscript = bool(config_get(style_config, "ir.unit.superscript_minus_one", True))
    apply_run_style(run, style_config, "ir.label")
    run = paragraph.add_run("): ")
    apply_run_style(run, style_config, "ir.label")
    paragraph.add_run(text.strip())


def _add_reaction_loadings_line(document: Document, compound: Compound, style_config: dict[str, Any]) -> None:
    block = compound.reaction
    if not block.get("formatted_text"):
        block = calculate_reaction_loadings(block)
        compound.reaction = block
    text = str(block.get("formatted_text", "")).strip()
    if not text:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    apply_paragraph_style(paragraph, style_config, "compound.summary")
    label = paragraph.add_run("Reaction loadings: ")
    label.bold = True
    paragraph.add_run(text.rstrip(".") + ".")


def _add_hrms_line(document: Document, compound: Compound, style_config: dict[str, Any]) -> None:
    hrms = compound.hrms or {}
    found_text = hrms_found_text(hrms, compound.hrms_found)
    if not hrms.get("calculated_mz") or not hrms.get("ion_formula"):
        hrms = build_hrms_block(
            formula=compound.formula,
            label=hrms_label_text(hrms, compound.hrms_label),
            adduct=hrms_adduct_text(hrms, compound.hrms_adduct),
            found_text=found_text,
            isotope_policy=str(hrms.get("isotope_policy", "auto_halogen")),
            isotope_labels=hrms.get("isotope_labels"),
        )
        compound.hrms = hrms
        if not compound.hrms_found:
            compound.hrms_found = found_text
        compound.hrms_calculated = float(hrms["calculated_mz"])
        compound.hrms_ion_formula = str(hrms["ion_formula"])

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    apply_paragraph_style(paragraph, style_config, "hrms")
    label_run = paragraph.add_run(f"{hrms.get('label') or compound.hrms_label}: ")
    apply_run_style(label_run, style_config, "hrms.label")
    _add_adduct_runs(paragraph, str(hrms.get("adduct") or compound.hrms_adduct), style_config)
    paragraph.add_run(" calcd for ")
    _add_formula_runs(paragraph, str(hrms.get("ion_formula") or compound.hrms_ion_formula), style_config, hrms.get("isotope_labels", {}))
    paragraph.add_run(f" {float(hrms['calculated_mz']):.4f}. Found {float(found_text):.4f}.")


def _add_elemental_analysis_line(document: Document, compound: Compound, style_config: dict[str, Any]) -> None:
    block = compound.elemental_analysis
    if not block.get("calculated"):
        block = calculate_elemental_analysis_block(compound.formula, found=found_from_block(block))
        compound.elemental_analysis = block
    calculated = block.get("calculated", {})
    found = block.get("found", {})
    if not calculated:
        return

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    apply_paragraph_style(paragraph, style_config, "elemental_analysis")
    label_run = paragraph.add_run("Anal.")
    apply_run_style(label_run, style_config, "elemental_analysis.label")
    body_run = paragraph.add_run(" Calcd for ")
    apply_run_style(body_run, style_config, "elemental_analysis.body")
    _add_formula_runs(paragraph, str(block.get("formula") or compound.formula), style_config)
    paragraph.add_run(": ")
    paragraph.add_run(_format_element_percentages(calculated))
    if found:
        paragraph.add_run(". Found: ")
        paragraph.add_run(_format_element_percentages(found))
    paragraph.add_run(".")


def _format_element_percentages(values: dict[str, float]) -> str:
    return "; ".join(f"{element}, {float(value):.2f}" for element, value in values.items())


def _add_nmr_warning(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(f"(Support check: {text})")
    run.font.color.rgb = RGBColor(192, 0, 0)
    run.bold = True


def _add_bookmark_range(start_paragraph, end_paragraph, name: str) -> None:
    if not name:
        return
    bookmark_id = str(int(hashlib.sha1(name.encode("utf-8")).hexdigest()[:8], 16))
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    insert_at = 1 if len(start_paragraph._p) and start_paragraph._p[0].tag == qn("w:pPr") else 0
    start_paragraph._p.insert(insert_at, start)
    end_paragraph._p.append(end)


def _add_spectrum_page(document: Document, block: DocumentBlock, style_config: dict[str, Any]) -> None:
    compound = block["content"]
    nucleus = block["nucleus"]
    image_path = block.get("image_path", "")
    mnova_path = block.get("mnova_path", "")
    embed_mode = block.get("embed_mode", "png")

    _add_spectrum_compound_title(document, compound, style_config)
    conditions = compound.h1_conditions if nucleus == "1H" else compound.c13_conditions
    _add_spectrum_nmr_title(document, nucleus, conditions, style_config)

    structure = document.add_paragraph()
    structure.paragraph_format.space_after = Pt(0)
    structure.add_run(f"[[SPECTRUM_STRUCTURE:{compound.number}:{nucleus}]]")

    if embed_mode in {"mnova", "both"} and mnova_path:
        mnova = document.add_paragraph()
        mnova.paragraph_format.space_after = Pt(0)
        mnova.add_run(f"[[MNOVA:{compound.number}:{nucleus}]]")

    if embed_mode in {"png", "both"} and image_path:
        picture = document.add_paragraph()
        picture.paragraph_format.space_after = Pt(0)
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        section = document.sections[-1]
        picture_width = section.page_width - section.left_margin - section.right_margin
        picture.add_run().add_picture(image_path, width=picture_width)


def _add_spectrum_compound_title(document: Document, compound: Compound, style_config: dict[str, Any]):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    apply_paragraph_style(paragraph, style_config, "appendix.title")
    run = paragraph.add_run(f"{compound.name} {compound.label}")
    apply_run_style(run, style_config, "appendix.title")
    return paragraph


def _add_spectrum_nmr_title(document: Document, nucleus: str, conditions: str, style_config: dict[str, Any]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    apply_paragraph_style(paragraph, style_config, "appendix.spectrum_title")
    if nucleus == "1H":
        run = _add_isotope_run(paragraph, "1", style_config)
        apply_run_style(run, style_config, "appendix.spectrum_title")
        run = paragraph.add_run("H NMR")
        apply_run_style(run, style_config, "appendix.spectrum_title")
    else:
        run = _add_isotope_run(paragraph, "13", style_config)
        apply_run_style(run, style_config, "appendix.spectrum_title")
        run = paragraph.add_run("C{")
        apply_run_style(run, style_config, "appendix.spectrum_title")
        run = _add_isotope_run(paragraph, "1", style_config)
        apply_run_style(run, style_config, "appendix.spectrum_title")
        run = paragraph.add_run("H} NMR")
        apply_run_style(run, style_config, "appendix.spectrum_title")
    if conditions:
        _add_conditions_runs(paragraph, conditions, style_config, "appendix.spectrum_title")


def _add_formula_runs(paragraph, formula: str, style_config: dict[str, Any], isotope_labels: dict[str, int] | None = None) -> None:
    if isotope_labels and config_get(style_config, "hrms.formula.isotope_labels", True):
        if _add_labeled_formula_runs(paragraph, formula, style_config, isotope_labels):
            return

    chunks = []
    current = ""
    current_is_digit = None

    for char in formula:
        is_digit = char.isdigit()
        if current and is_digit != current_is_digit:
            chunks.append((current, current_is_digit))
            current = ""
        current += char
        current_is_digit = is_digit
    if current:
        chunks.append((current, current_is_digit))

    for text, is_digit in chunks:
        run = paragraph.add_run(text)
        run.font.subscript = bool(is_digit and config_get(style_config, "hrms.formula.subscripts", True))
        if text in {"+", "-"}:
            run.font.superscript = bool(config_get(style_config, "hrms.formula.charge_superscript", True))


def _add_labeled_formula_runs(paragraph, formula: str, style_config: dict[str, Any], isotope_labels: dict[str, int]) -> bool:
    charge = ""
    formula_body = formula
    if formula_body.endswith(("+", "-")):
        charge = formula_body[-1]
        formula_body = formula_body[:-1]
    try:
        elements = parse_formula(formula_body)
    except ValueError:
        return False

    for element, count in elements.items():
        if element in isotope_labels:
            run = paragraph.add_run(str(isotope_labels[element]))
            run.font.superscript = bool(config_get(style_config, "hrms.formula.isotope_label_superscript", True))
        paragraph.add_run(element)
        if count != 1:
            run = paragraph.add_run(str(count))
            run.font.subscript = bool(config_get(style_config, "hrms.formula.subscripts", True))
    if charge:
        run = paragraph.add_run(charge)
        run.font.superscript = bool(config_get(style_config, "hrms.formula.charge_superscript", True))
    return True


def _add_isotope_run(paragraph, text: str, style_config: dict[str, Any]):
    run = paragraph.add_run(text)
    run.font.superscript = bool(config_get(style_config, "chem_formatting.isotope_numbers.superscript", True))
    return run


def _add_conditions_runs(paragraph, conditions: str, style_config: dict[str, Any], style_path: str = "nmr.conditions") -> None:
    run = paragraph.add_run(" (")
    apply_run_style(run, style_config, style_path)
    _add_formula_text_runs(paragraph, conditions, style_config, style_path)
    run = paragraph.add_run(")")
    apply_run_style(run, style_config, style_path)


def _add_adduct_runs(paragraph, adduct: str, style_config: dict[str, Any]) -> None:
    for index, char in enumerate(adduct):
        run = paragraph.add_run(char)
        if index == len(adduct) - 1 and char in "+-":
            run.font.superscript = bool(config_get(style_config, "hrms.adduct.superscript_charge", False))


def _add_formula_text_runs(paragraph, text: str, style_config: dict[str, Any], style_path: str = "") -> None:
    pos = 0
    for match in re.finditer(r"\b(?:CDCl3|DMSO-d6|C\d*H\d*(?:[A-Z][a-z]?\d*)*)\b", text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            if style_path:
                apply_run_style(run, style_config, style_path)
        formula = match.group(0)
        for char in formula:
            run = paragraph.add_run(char)
            if style_path:
                apply_run_style(run, style_config, style_path)
            run.font.subscript = bool(char.isdigit() and config_get(style_config, "chem_formatting.formulas.subscripts", True))
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if style_path:
            apply_run_style(run, style_config, style_path)


def _add_chem_text_runs(paragraph, text: str, style_config: dict[str, Any], style_path: str = "") -> None:
    if config_get(style_config, "chem_formatting.ranges.en_dash", False):
        text = re.sub(r"(?<=\d)\s+-\s+(?=\d)", "\u2013", text)

    pattern = re.compile(r"(?<![A-Za-z])(?:(\d+)(J)([A-Z]{1,3})?|\b(J)\b)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            if style_path:
                apply_run_style(run, style_config, style_path)
        if match.group(2):
            order, j_text, partner = match.group(1), match.group(2), match.group(3) or ""
            run = paragraph.add_run(order)
            if style_path:
                apply_run_style(run, style_config, style_path)
            run.font.superscript = bool(config_get(style_config, "chem_formatting.coupling_constants.order_superscript", False))
            run = paragraph.add_run(j_text)
            if style_path:
                apply_run_style(run, style_config, style_path)
            run.italic = bool(config_get(style_config, "chem_formatting.coupling_constants.j_italic", False))
            if partner:
                run = paragraph.add_run(partner)
                if style_path:
                    apply_run_style(run, style_config, style_path)
                run.font.subscript = bool(config_get(style_config, "chem_formatting.coupling_constants.coupling_partner_subscript", False))
        else:
            run = paragraph.add_run(match.group(4))
            if style_path:
                apply_run_style(run, style_config, style_path)
            run.italic = bool(config_get(style_config, "chem_formatting.coupling_constants.j_italic", False))
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if style_path:
            apply_run_style(run, style_config, style_path)
