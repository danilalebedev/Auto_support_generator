from __future__ import annotations

import json
import re
import subprocess
import sys
from dataclasses import dataclass, field, replace
from pathlib import Path
from typing import Any

from docx import Document

from ..structure_metadata import StructureMetadata, extract_structure_metadata_by_cell
from .compound import Compound
from .types import Issue, ReagentAmount


@dataclass(frozen=True)
class LoadingsWorkflowPaths:
    schema_docx: Path
    scope_docx: Path
    template_docx: Path | None = None


@dataclass(frozen=True)
class SchemaEntry:
    key: str
    label: str
    equivalents: float | None = None
    mw: float | None = None
    density_g_mL: float | None = None
    concentration_M: float | None = None


@dataclass(frozen=True)
class ScopeRow:
    product_number: str
    product_mass_mg: float | None
    product: StructureMetadata
    reagents: dict[str, StructureMetadata] = field(default_factory=dict)
    reagent_masses_mg: dict[str, float | None] = field(default_factory=dict)
    reagent_cells: dict[str, tuple[int, int, int]] = field(default_factory=dict)
    product_cell: tuple[int, int, int] = (0, 0, 0)

    @property
    def reagent_1_mass_mg(self) -> float | None:
        return self.reagent_masses_mg.get("Reagent_1")

    @property
    def reagent_1(self) -> StructureMetadata:
        return self.reagents.get("Reagent_1", StructureMetadata())

    @property
    def reagent_2(self) -> StructureMetadata:
        return self.reagents.get("Reagent_2", StructureMetadata())

    @property
    def reagent_1_cell(self) -> tuple[int, int, int]:
        return self.reagent_cells.get("Reagent_1", (0, 0, 0))

    @property
    def reagent_2_cell(self) -> tuple[int, int, int]:
        return self.reagent_cells.get("Reagent_2", (0, 0, 0))


def discover_loadings_workflow(base_dir: str | Path) -> LoadingsWorkflowPaths | None:
    base_dir = Path(base_dir)
    directories = [base_dir / "loadings", base_dir]
    for directory in directories:
        if not directory.exists():
            continue
        schema = _find_docx(directory, "Reaction_schema.docx")
        scope = _find_docx(directory, "Scope.docx")
        if schema and scope:
            return LoadingsWorkflowPaths(schema, scope)
    return None


def apply_loadings_workflow(
    compounds: list[Compound],
    base_dir: str | Path,
    paths: LoadingsWorkflowPaths | None = None,
    template_docx: str | Path | None = None,
    structure_names_by_cell: dict[tuple[int, int, int], str] | None = None,
) -> list[Issue]:
    paths = paths or discover_loadings_workflow(base_dir)
    if paths is None:
        return []

    issues: list[Issue] = []
    schema = read_reaction_schema(paths.schema_docx)
    template_path = Path(template_docx) if template_docx else paths.template_docx or _default_si_template_path()
    template = read_characterization_template(template_path)
    if structure_names_by_cell is None:
        structure_names_by_cell, name_issues = _structure_names_for_template(paths.scope_docx, template)
        issues.extend(name_issues)
    scope_rows = read_scope(paths.scope_docx, structure_names_by_cell=structure_names_by_cell)
    compounds_by_number = {compound.number.strip(): compound for compound in compounds if compound.number.strip()}
    mismatch = _scope_input_number_mismatch(compounds_by_number, scope_rows, paths.scope_docx)
    if mismatch:
        return [mismatch]

    for row in scope_rows:
        compound = compounds_by_number.get(row.product_number)
        row_issues = _apply_scope_row(compound, row, schema, template)
        issues.extend(row_issues)

    return issues


def _scope_input_number_mismatch(
    compounds_by_number: dict[str, Compound],
    scope_rows: list[ScopeRow],
    scope_path: Path,
) -> Issue | None:
    input_numbers = set(compounds_by_number)
    scope_numbers = {row.product_number.strip() for row in scope_rows if row.product_number.strip()}
    if input_numbers == scope_numbers:
        return None
    missing_in_scope = sorted(input_numbers - scope_numbers)
    extra_in_scope = sorted(scope_numbers - input_numbers)
    message_parts = [
        "Reaction loadings cannot be calculated because compound numbers in Scope.docx and compound table do not match.",
        f"Input compounds: {', '.join(sorted(input_numbers)) or '<none>'}.",
        f"Scope products: {', '.join(sorted(scope_numbers)) or '<none>'}.",
    ]
    if missing_in_scope:
        message_parts.append(f"Missing in Scope.docx: {', '.join(missing_in_scope)}.")
    if extra_in_scope:
        message_parts.append(f"Extra in Scope.docx: {', '.join(extra_in_scope)}.")
    return {
        "code": "LOADINGS_SCOPE_INPUT_MISMATCH",
        "severity": "error",
        "message": " ".join(message_parts),
        "path": str(scope_path),
    }


def read_reaction_schema(path: str | Path) -> dict[str, SchemaEntry]:
    document = Document(str(path))
    if not document.tables:
        return {}

    table = document.tables[0]
    headers = [_normalize_header(cell.text) for cell in table.rows[0].cells]
    name_col = _header_index(headers, "reagents", "reagent")
    equiv_col = _header_index(headers, "equiv", "equivalents")
    mw_col = _header_index(headers, "mwgmol", "mw")
    density_col = _header_index(headers, "densitygml", "density")
    concentration_col = _header_index(headers, "concentrationm", "concentration")

    entries: dict[str, SchemaEntry] = {}
    for row in table.rows[1:]:
        cells = row.cells
        label = _cell_text(cells, name_col)
        if not label:
            continue
        key = _schema_key(label)
        entries[key] = SchemaEntry(
            key=key,
            label=label,
            equivalents=_float_or_none(_cell_text(cells, equiv_col)),
            mw=_float_or_none(_cell_text(cells, mw_col)),
            density_g_mL=_float_or_none(_cell_text(cells, density_col)),
            concentration_M=_float_or_none(_cell_text(cells, concentration_col)),
        )
    return entries


def read_scope(
    path: str | Path,
    structure_names_by_cell: dict[tuple[int, int, int], str] | None = None,
) -> list[ScopeRow]:
    document = Document(str(path))
    if not document.tables:
        return []

    table = document.tables[0]
    headers = [_normalize_header(cell.text) for cell in table.rows[0].cells]
    reagent_cols = _reagent_columns(headers)
    reagent_mass_cols = _reagent_mass_columns(headers)
    product_col = _header_index(headers, "product")
    product_number_col = _header_index(headers, "productnumber", "number")
    product_mass_col = _header_index(headers, "massofproductmg", "productmassmg")
    metadata_by_cell = extract_structure_metadata_by_cell(path)

    rows: list[ScopeRow] = []
    for row_index, row in enumerate(table.rows[1:], start=2):
        cells = row.cells
        product_number = _cell_text(cells, product_number_col)
        if not product_number:
            continue
        reagents: dict[str, StructureMetadata] = {}
        reagent_masses_mg: dict[str, float | None] = {}
        reagent_cells: dict[str, tuple[int, int, int]] = {}
        for reagent_key in sorted(set(reagent_cols) | set(reagent_mass_cols), key=_reagent_sort_key):
            reagent_col = reagent_cols.get(reagent_key, -1)
            reagent_cell = (1, row_index, reagent_col + 1) if reagent_col >= 0 else (0, 0, 0)
            reagent_cells[reagent_key] = reagent_cell
            reagents[reagent_key] = _metadata_with_name(
                metadata_by_cell.get(reagent_cell, StructureMetadata()),
                structure_names_by_cell,
                reagent_cell,
            )
            reagent_masses_mg[reagent_key] = _float_or_none(_cell_text(cells, reagent_mass_cols.get(reagent_key, -1)))
        rows.append(
            ScopeRow(
                product_number=product_number,
                product_mass_mg=_float_or_none(_cell_text(cells, product_mass_col)),
                product=_metadata_with_name(
                    metadata_by_cell.get((1, row_index, product_col + 1), StructureMetadata()),
                    structure_names_by_cell,
                    (1, row_index, product_col + 1),
                ),
                reagents=reagents,
                reagent_masses_mg=reagent_masses_mg,
                reagent_cells=reagent_cells,
                product_cell=(1, row_index, product_col + 1),
            )
        )
    return rows


def read_characterization_template(path: str | Path) -> str:
    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    loadings_paragraph = next((text for text in paragraphs if _paragraph_has_loadings_placeholders(text)), "")
    text = loadings_paragraph or "\n".join(paragraphs)
    return _repair_known_template_placeholders(text)


def _apply_scope_row(
    compound: Compound,
    row: ScopeRow,
    schema: dict[str, SchemaEntry],
    template: str,
) -> list[Issue]:
    issues: list[Issue] = []
    reagent_1_schema = schema.get("Reagent_1")
    if reagent_1_schema is None:
        return [
            {
                "code": "LOADINGS_SCHEMA_INCOMPLETE",
                "severity": "warning",
                "message": "Reaction_schema.docx must contain a Reagent_1 row.",
                "compound_id": compound.id or compound.number,
            }
        ]

    reagent_1_mw = _metadata_mw(row.reagent_1) or reagent_1_schema.mw
    product_mw = _metadata_mw(row.product)
    target_mmol = _limiting_mmol(row.reagent_1_mass_mg, reagent_1_mw, reagent_1_schema.equivalents)
    if target_mmol is None:
        issues.append(_compound_issue(compound, "LOADINGS_LIMITING_MISSING", "Cannot calculate limiting reagent mmol."))
        return issues

    reagent_amounts: dict[str, dict[str, Any]] = {}
    for key, entry in _numbered_reagent_schema_entries(schema):
        metadata = row.reagents.get(key, StructureMetadata())
        mw = _metadata_mw(metadata) or entry.mw
        if key == "Reagent_1":
            reagent_amounts[key] = _amount_from_mass(
                schema_entry=entry,
                metadata=metadata,
                mass_mg=row.reagent_masses_mg.get(key),
                mw=mw,
                mmol=target_mmol * (entry.equivalents or 1.0),
            )
            reagent_amounts[key]["number"] = _infer_precursor_number(row.product_number)
        else:
            reagent_amounts[key] = _amount_from_equivalents(entry, target_mmol, metadata, mw)

    for key, entry in schema.items():
        if _is_numbered_reagent_key(key):
            continue
        reagent_amounts[key] = _amount_from_equivalents(entry, target_mmol, StructureMetadata(), entry.mw)

    product_mmol = _safe_div(row.product_mass_mg, product_mw)
    percent_yield = product_mmol / target_mmol * 100 if product_mmol is not None and target_mmol else None
    rf_value, rf_system = _split_rf(compound.rf)

    values = _base_template_values(
        compound=compound,
        row=row,
        target_mmol=target_mmol,
        product_mmol=product_mmol,
        percent_yield=percent_yield,
        rf_value=rf_value,
        rf_system=rf_system,
    )
    for key, amount in reagent_amounts.items():
        values[_token_key(key)] = _display_schema_name(schema[key].label)
        values.update(_amount_template_values(key, amount))

    preparation = _render_template(template, values)
    compound.preparation = preparation.rstrip(".")
    compound.yield_text = _yield_text(row.product_mass_mg, percent_yield)
    if row.product.formula and not compound.formula:
        compound.formula = row.product.formula
    compound.reaction = {
        "scale_basis": "limiting_reagent",
        "target_mmol": round(target_mmol, 4),
        "limiting_reagent": "Reagent_1",
        "reagents": [_reaction_reagent(key, amount, schema.get(key)) for key, amount in reagent_amounts.items()],
        "template_values": values,
        "preparation_includes_summary": True,
        "hide_loadings_line": True,
        "source": "loadings_workflow",
    }
    return issues


def _base_template_values(
    compound: Compound,
    row: ScopeRow,
    target_mmol: float,
    product_mmol: float | None,
    percent_yield: float | None,
    rf_value: str,
    rf_system: str,
) -> dict[str, str]:
    return {
        _token_key("Product.name"): _metadata_display(row.product, compound.name),
        _token_key("Product.number"): row.product_number,
        _token_key("Product.mg"): _format_mass(row.product_mass_mg),
        _token_key("Product.g"): _format_scaled_amount(_scale_value(row.product_mass_mg, 1 / 1000)),
        _token_key("Product.kg"): _format_scaled_amount(_scale_value(row.product_mass_mg, 1 / 1_000_000)),
        _token_key("Product.mmol"): _format_mmol(product_mmol),
        _token_key("Product.mol"): _format_scaled_amount(_scale_value(product_mmol, 1 / 1000)),
        _token_key("Product.yield.percent"): _format_percent(percent_yield),
        _token_key("Product.appearance"): _compound_appearance(compound),
        _token_key("Product.mp"): _strip_temperature_unit(compound.melting_point),
        _token_key("Product.rf.value"): rf_value,
        _token_key("Product.rf.system"): rf_system,
        # Legacy aliases are kept for older local templates, but new templates should use Product.*.
        _token_key("number_Product"): row.product_number,
        _token_key("number_Reagent_1"): _infer_precursor_number(row.product_number),
        _token_key("mg_yield_Product"): _format_mass(row.product_mass_mg),
        _token_key("percent_yield_Product"): _format_percent(percent_yield),
        _token_key("Product.precursor_number"): _infer_precursor_number(row.product_number),
        _token_key("Product.mass.mg"): _format_mass(row.product_mass_mg),
        _token_key("Product.yield.mg"): _format_mass(row.product_mass_mg),
        _token_key("yield.Product.mg"): _format_mass(row.product_mass_mg),
        _token_key("yield.Product.percent"): _format_percent(percent_yield),
        _token_key("color"): _compound_appearance(compound),
        _token_key("appearance"): _compound_appearance(compound),
        _token_key("mp"): _strip_temperature_unit(compound.melting_point),
        _token_key("Rf"): rf_value,
        _token_key("system_Rf"): rf_system,
        _token_key("rf.value"): rf_value,
        _token_key("rf.system"): rf_system,
        _token_key("target_mmol"): _format_mmol(target_mmol),
        _token_key("Reaction.target.mmol"): _format_mmol(target_mmol),
    }


def _amount_template_values(name: str, amount: dict[str, Any]) -> dict[str, str]:
    volume_ul = _volume_ul(amount)
    volume_ml = _volume_ml(amount)
    keys = {
        _token_key(f"{name}.name"): str(amount.get("name") or ""),
        _token_key(f"{name}.mg"): _format_mass(amount.get("mass_mg")),
        _token_key(f"{name}.g"): _format_scaled_amount(_scale_value(amount.get("mass_mg"), 1 / 1000)),
        _token_key(f"{name}.kg"): _format_scaled_amount(_scale_value(amount.get("mass_mg"), 1 / 1_000_000)),
        _token_key(f"{name}.mmol"): _format_mmol(amount.get("mmol")),
        _token_key(f"{name}.mol"): _format_scaled_amount(_scale_value(amount.get("mmol"), 1 / 1000)),
        _token_key(f"{name}.mcl"): _format_volume_ul(volume_ul),
        _token_key(f"{name}.ml"): _format_volume_ml(volume_ml),
        _token_key(f"{name}.l"): _format_scaled_amount(_scale_value(volume_ml, 1 / 1000)),
        _token_key(f"{name}.eq"): _format_equivalents(amount.get("equivalents")),
        # Legacy aliases are kept for older local templates, but new templates should use object.attribute aliases above.
        _token_key(f"name_{name}"): str(amount.get("name") or ""),
        _token_key(f"mg_{name}"): _format_mass(amount.get("mass_mg")),
        _token_key(f"mol_{name}"): _format_mmol(amount.get("mmol")),
        _token_key(f"mmol_{name}"): _format_mmol(amount.get("mmol")),
        _token_key(f"uL_{name}"): _format_volume_ul(amount.get("volume_uL")),
        _token_key(f"\u03bcL_{name}"): _format_volume_ul(amount.get("volume_uL")),
        _token_key(f"ml_{name}"): _format_volume_ml(amount.get("volume_mL")),
        _token_key(f"{name}.mass.mg"): _format_mass(amount.get("mass_mg")),
        _token_key(f"{name}.volume.uL"): _format_volume_ul(amount.get("volume_uL")),
        _token_key(f"{name}.uL"): _format_volume_ul(amount.get("volume_uL")),
        _token_key(f"{name}.\u03bcL"): _format_volume_ul(amount.get("volume_uL")),
        _token_key(f"{name}.volume.mL"): _format_volume_ml(amount.get("volume_mL")),
        _token_key(f"{name}.mL"): _format_volume_ml(amount.get("volume_mL")),
        _token_key(f"{name}.equiv"): _format_equivalents(amount.get("equivalents")),
        _token_key(f"{name}.formula"): str(amount.get("formula") or ""),
        _token_key(f"{name}.number"): str(amount.get("number") or ""),
    }
    if name == "Reagent_1":
        keys[_token_key("mg_Reagent 1")] = _format_mass(amount.get("mass_mg"))
        keys[_token_key("Reagent_1.number")] = str(amount.get("number") or "")
    keys.update(
        {
            _token_key(f"{name}.name"): str(amount.get("name") or ""),
            _token_key(f"{name}.mg"): _format_mass(amount.get("mass_mg")),
            _token_key(f"{name}.g"): _format_scaled_amount(_scale_value(amount.get("mass_mg"), 1 / 1000)),
            _token_key(f"{name}.kg"): _format_scaled_amount(_scale_value(amount.get("mass_mg"), 1 / 1_000_000)),
            _token_key(f"{name}.mmol"): _format_mmol(amount.get("mmol")),
            _token_key(f"{name}.mol"): _format_scaled_amount(_scale_value(amount.get("mmol"), 1 / 1000)),
            _token_key(f"{name}.mcl"): _format_volume_ul(volume_ul),
            _token_key(f"{name}.ml"): _format_volume_ml(volume_ml),
            _token_key(f"{name}.l"): _format_scaled_amount(_scale_value(volume_ml, 1 / 1000)),
            _token_key(f"{name}.eq"): _format_equivalents(amount.get("equivalents")),
        }
    )
    return keys


def _amount_from_mass(
    schema_entry: SchemaEntry,
    metadata: StructureMetadata,
    mass_mg: float | None,
    mw: float | None,
    mmol: float | None,
) -> dict[str, Any]:
    return _compact_amount(
        {
            "name": _metadata_display(metadata, schema_entry.label),
            "formula": metadata.formula,
            "mw": mw,
            "equivalents": schema_entry.equivalents,
            "mmol": mmol,
            "mass_mg": mass_mg,
            "volume_uL": _volume_from_mass(mass_mg, schema_entry.density_g_mL),
        }
    )


def _amount_from_equivalents(
    schema_entry: SchemaEntry,
    target_mmol: float,
    metadata: StructureMetadata,
    mw: float | None,
) -> dict[str, Any]:
    mmol = target_mmol * schema_entry.equivalents if schema_entry.equivalents is not None else None
    mass_mg = mmol * mw if mmol is not None and mw is not None else None
    volume_mL = _safe_div(target_mmol, schema_entry.concentration_M)
    return _compact_amount(
        {
            "name": _metadata_display(metadata, schema_entry.label),
            "formula": metadata.formula,
            "mw": mw,
            "equivalents": schema_entry.equivalents,
            "mmol": mmol,
            "mass_mg": mass_mg,
            "volume_uL": _volume_from_mass(mass_mg, schema_entry.density_g_mL),
            "volume_mL": volume_mL,
            "concentration_M": schema_entry.concentration_M,
            "density_g_mL": schema_entry.density_g_mL,
        }
    )


def _reaction_reagent(key: str, amount: dict[str, Any], schema_entry: SchemaEntry | None) -> ReagentAmount:
    role = "solvent" if schema_entry and schema_entry.concentration_M else "reagent"
    reagent: ReagentAmount = {
        "name": str(amount.get("name") or key),
        "role": role,  # type: ignore[typeddict-item]
    }
    for source_key in ("formula", "mw", "equivalents", "mmol", "mass_mg", "volume_uL", "volume_mL", "density_g_mL", "concentration_M"):
        if amount.get(source_key) not in {None, ""}:
            reagent[source_key] = amount[source_key]  # type: ignore[literal-required]
    return reagent


def _numbered_reagent_schema_entries(schema: dict[str, SchemaEntry]) -> list[tuple[str, SchemaEntry]]:
    return sorted(
        ((key, entry) for key, entry in schema.items() if _is_numbered_reagent_key(key)),
        key=lambda item: _reagent_sort_key(item[0]),
    )


def _is_numbered_reagent_key(key: str) -> bool:
    return bool(re.fullmatch(r"Reagent_\d+", key))


def _render_template(template: str, values: dict[str, str]) -> str:
    template = _cleanup_optional_template_fragments(template, values)

    def replace(match: re.Match[str]) -> str:
        raw_key = match.group(1)
        return values.get(_token_key(raw_key), match.group(0))

    text = re.sub(r"\{([^{}]+)\}", replace, template)
    text = re.sub(r"\s+([,.;])", r"\1", text)
    text = re.sub(r"\(\s+", "(", text)
    text = re.sub(r"\s+\)", ")", text)
    text = re.sub(r"\s{2,}", " ", text)
    return text.strip()


def _cleanup_optional_template_fragments(template: str, values: dict[str, str]) -> str:
    text = template
    if not _any_value(values, "Product.yield.percent", "yield.Product.percent", "percent_yield_Product"):
        text = _remove_fragment_with_any_placeholder(
            text,
            ("Product.yield.percent", "yield.Product.percent", "percent_yield_Product"),
            r"\s*\(\s*{placeholder}\s*\)",
        )
    if not _any_value(values, "Product.mg", "Product.yield.mg", "yield.Product.mg", "mg_yield_Product"):
        text = _remove_fragment_with_any_placeholder(
            text,
            ("Product.mg", "Product.yield.mg", "yield.Product.mg", "mg_yield_Product"),
            r"\s*Yield\s+{placeholder}\s*mg\s*(?:\([^)]*\))?\s*[;.]?",
        )
    if not values.get(_token_key("Product.appearance")):
        text = _remove_fragment_with_placeholder(text, "Product.appearance", r"\s*;\s*{placeholder}")
        text = _remove_fragment_with_placeholder(text, "Product.appearance", r"{placeholder}\s*;\s*")
    if not values.get(_token_key("Product.mp")):
        text = _remove_fragment_with_placeholder(text, "Product.mp", r"\s*(?:[;.]\s*)?mp\s+{placeholder}\s*(?:°C|deg\.?\s*C|C)?")
    if not values.get(_token_key("Product.rf.value")):
        text = _remove_fragment_with_placeholder(
            text,
            "Product.rf.value",
            r"\s*[;.]\s*Rf\s*=\s*{placeholder}\s*(?:\(\s*" + _placeholder_any_pattern(("Product.rf.system",)) + r"\s*\))?",
        )
    if values.get(_token_key("Product.rf.value")) and not values.get(_token_key("Product.rf.system")):
        text = _remove_fragment_with_placeholder(text, "Product.rf.system", r"\s*\(\s*{placeholder}\s*\)")
    return _normalize_optional_template_cleanup(text)


def _remove_fragment_with_placeholder(text: str, key: str, pattern_template: str) -> str:
    return _remove_fragment_with_any_placeholder(text, (key,), pattern_template)


def _remove_fragment_with_any_placeholder(text: str, keys: tuple[str, ...], pattern_template: str) -> str:
    for key in keys:
        placeholder = _placeholder_any_pattern((key,))
        text = re.sub(pattern_template.replace("{placeholder}", placeholder), "", text, flags=re.IGNORECASE)
    return text


def _placeholder_any_pattern(keys: tuple[str, ...]) -> str:
    alternatives = []
    for key in keys:
        escaped = re.escape(key)
        alternatives.append(r"\{\s*" + escaped + r"\s*\}")
    return r"(?:" + "|".join(alternatives) + r")"


def _any_value(values: dict[str, str], *keys: str) -> bool:
    return any(values.get(_token_key(key)) for key in keys)


def _normalize_optional_template_cleanup(text: str) -> str:
    text = re.sub(r"\s+([,.;])", r"\1", text)
    text = re.sub(r";\s*;", ";", text)
    text = re.sub(r"\(\s*\)", "", text)
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r";\s*\.", ".", text)
    text = re.sub(r"\s+\.", ".", text)
    return text.strip()


def _repair_known_template_placeholders(template: str) -> str:
    template = template.replace("{mg_ K2CO3},", "{mg_ K2CO3} mg,")
    template = template.replace("{\u03bcL_AcOH},", "{\u03bcL_AcOH} \u03bcL,")
    return template.replace("{K2CO3} ({\u03bcL_AcOH}", "{AcOH} ({\u03bcL_AcOH}").replace(
        "{K2CO3} ({uL_AcOH}", "{AcOH} ({uL_AcOH}"
    )


def _find_docx(directory: Path, expected_name: str) -> Path | None:
    expected = expected_name.lower()
    for path in directory.glob("*.docx"):
        if path.name.lower() == expected:
            return path
    return None


def _default_si_template_path() -> Path:
    return Path(__file__).resolve().parents[1] / "templates" / "SI_template.docx"


def _structure_names_for_template(scope_path: Path, template: str) -> tuple[dict[tuple[int, int, int], str], list[Issue]]:
    if not _template_requests_structure_names(template):
        return {}, []
    rows = read_scope(scope_path)
    cells = sorted({cell for row in rows for cell in row.reagent_cells.values()} | {row.product_cell for row in rows})
    cells = [cell for cell in cells if cell != (0, 0, 0)]
    if not cells:
        return {}, []
    return _chemdraw_names_for_cells(scope_path, cells)


def _template_requests_structure_names(template: str) -> bool:
    for match in re.finditer(r"\{([^{}]+)\}", template):
        key = _token_key(match.group(1))
        if key.startswith("name") or key.endswith(".name"):
            return True
    return False


def _paragraph_has_loadings_placeholders(text: str) -> bool:
    keys = {_token_key(match.group(1)) for match in re.finditer(r"\{([^{}]+)\}", text)}
    loading_product_keys = {
        "product.mg",
        "product.g",
        "product.kg",
        "product.mmol",
        "product.mol",
        "product.yield.percent",
        "product.appearance",
        "product.mp",
        "product.rf.value",
        "product.rf.system",
        "product.precursor.number",
        "product.mass.mg",
        "product.yield.mg",
    }
    if keys & loading_product_keys:
        return True
    if any(
        key.startswith(prefix)
        for key in keys
        for prefix in (
            "name.reagent",
            "mg.reagent",
            "mmol.reagent",
            "number.product",
            "mg.yield.product",
            "reagent.",
            "solvent.",
        )
    ):
        return True
    return any(_is_named_loading_alias(key) for key in keys)


def _is_named_loading_alias(key: str) -> bool:
    if key.split(".", 1)[0] in {"product", "reagent", "solvent"}:
        return False
    return bool(re.match(r"^[a-z0-9]+(?:\.[a-z0-9]+)?\.(?:name|mg|g|kg|mmol|mol|mcl|ml|l|eq)$", key))


def _chemdraw_names_for_cells(scope_path: Path, cells: list[tuple[int, int, int]], timeout: int = 240) -> tuple[dict[tuple[int, int, int], str], list[Issue]]:
    cell_arg = ",".join(_format_cell_coordinate(cell) for cell in cells)
    if getattr(sys, "frozen", False):
        command = [sys.executable, "--si-generator-chemdraw-names", str(scope_path), "--cells", cell_arg]
    else:
        command = [sys.executable, "-m", "si_generator.chemdraw_names", str(scope_path), "--cells", cell_arg]
    try:
        completed = subprocess.run(command, capture_output=True, text=True, timeout=timeout, check=False)
    except Exception as exc:
        return {}, [_loadings_name_issue(scope_path, str(exc))]
    if completed.returncode != 0:
        detail = (completed.stderr or completed.stdout or "ChemDraw name generation failed.").strip()
        return {}, [_loadings_name_issue(scope_path, detail)]
    try:
        raw = json.loads(completed.stdout or "{}")
    except json.JSONDecodeError as exc:
        return {}, [_loadings_name_issue(scope_path, f"Cannot parse ChemDraw name output: {exc}")]
    return {_parse_cell_coordinate(key): str(value).strip() for key, value in raw.items() if str(value).strip()}, []


def _loadings_name_issue(scope_path: Path, detail: str) -> Issue:
    return {
        "code": "LOADINGS_NAME_GENERATION_FAILED",
        "severity": "warning",
        "message": "Could not generate reagent/product names with ChemDraw. Formula fallback will be used.",
        "path": str(scope_path),
        "detail": detail,
    }


def _format_cell_coordinate(cell: tuple[int, int, int]) -> str:
    return f"{cell[0]}:{cell[1]}:{cell[2]}"


def _parse_cell_coordinate(value: str) -> tuple[int, int, int]:
    parts = [int(part.strip()) for part in value.split(":")]
    if len(parts) != 3:
        raise ValueError(f"Invalid cell coordinate: {value}")
    return parts[0], parts[1], parts[2]


def _metadata_with_name(
    metadata: StructureMetadata,
    structure_names_by_cell: dict[tuple[int, int, int], str] | None,
    cell: tuple[int, int, int],
) -> StructureMetadata:
    name = (structure_names_by_cell or {}).get(cell, "").strip()
    if not name:
        return metadata
    return replace(metadata, name=name)


def _cell_text(cells: Any, index: int) -> str:
    if index < 0 or index >= len(cells):
        return ""
    return cells[index].text.strip()


def _header_index(headers: list[str], *aliases: str) -> int:
    wanted = {_normalize_header(alias) for alias in aliases}
    for index, header in enumerate(headers):
        if header in wanted:
            return index
    return -1


def _reagent_columns(headers: list[str]) -> dict[str, int]:
    columns: dict[str, int] = {}
    for index, header in enumerate(headers):
        match = re.fullmatch(r"reagent(\d+)", header)
        if match:
            columns[f"Reagent_{int(match.group(1))}"] = index
    return columns


def _reagent_mass_columns(headers: list[str]) -> dict[str, int]:
    columns: dict[str, int] = {}
    patterns = (
        re.compile(r"massofreagent(\d+)mg"),
        re.compile(r"reagent(\d+)massmg"),
        re.compile(r"massreagent(\d+)mg"),
    )
    for index, header in enumerate(headers):
        for pattern in patterns:
            match = pattern.fullmatch(header)
            if match:
                columns[f"Reagent_{int(match.group(1))}"] = index
                break
    return columns


def _reagent_sort_key(key: str) -> tuple[int, str]:
    match = re.fullmatch(r"Reagent_(\d+)", key)
    return (int(match.group(1)), key) if match else (10_000, key)


def _normalize_header(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", text.lower())


def _schema_key(label: str) -> str:
    key = re.sub(r"\s+", "_", label.strip())
    return re.sub(r"[^A-Za-z0-9_]+", "", key)


def _token_key(text: str) -> str:
    normalized = text.replace("\u03bc", "u")
    return re.sub(r"[^a-z0-9]+", ".", normalized.lower()).strip(".")


def _display_schema_name(label: str) -> str:
    if label.startswith("Solvent_"):
        return label.split("_", 1)[1]
    return label


def _metadata_display(metadata: StructureMetadata, fallback: str) -> str:
    return metadata.name or metadata.formula or _display_schema_name(fallback)


def _metadata_mw(metadata: StructureMetadata) -> float | None:
    return metadata.molecular_weight or None


def _limiting_mmol(mass_mg: float | None, mw: float | None, equivalents: float | None) -> float | None:
    amount = _safe_div(mass_mg, mw)
    if amount is None:
        return None
    if equivalents:
        return amount / equivalents
    return amount


def _safe_div(numerator: float | None, denominator: float | None) -> float | None:
    if numerator is None or denominator in {None, 0}:
        return None
    return numerator / float(denominator)


def _volume_from_mass(mass_mg: float | None, density_g_mL: float | None) -> float | None:
    if mass_mg is None or not density_g_mL:
        return None
    return mass_mg / density_g_mL


def _float_or_none(value: str | float | int | None) -> float | None:
    if value in {None, ""}:
        return None
    try:
        return float(str(value).replace(",", "."))
    except ValueError:
        return None


def _split_rf(value: str) -> tuple[str, str]:
    text = value.strip()
    match = re.match(r"(.+?)\s*\((.+)\)\s*$", text)
    if not match:
        return text, ""
    return match.group(1).strip(), match.group(2).strip()


def _infer_precursor_number(product_number: str) -> str:
    match = re.match(r"(\d+)(.*)", product_number.strip())
    if not match:
        return product_number
    number = max(int(match.group(1)) - 1, 0)
    return f"{number}{match.group(2)}"


def _strip_temperature_unit(value: str) -> str:
    return re.sub(r"\s*(?:deg\.?\s*C|degrees?\s*C|C|°C)\s*$", "", value.strip(), flags=re.IGNORECASE)


def _normalize_sentence_piece(value: str) -> str:
    return value.strip().rstrip(".;")


def _compound_appearance(compound: Compound) -> str:
    return _normalize_sentence_piece(" ".join(part for part in [compound.color, compound.state] if part))


def _yield_text(mass_mg: float | None, percent_yield: float | None) -> str:
    mass = _format_mass(mass_mg)
    percent = _format_percent(percent_yield)
    if mass and percent:
        return f"{mass} mg ({percent})"
    return mass


def _format_mass(value: Any) -> str:
    parsed = _to_float(value)
    if parsed is None:
        return ""
    if abs(parsed) >= 10:
        return f"{parsed:.0f}"
    return _format_decimal(parsed, 1)


def _scale_value(value: Any, factor: float) -> float | None:
    parsed = _to_float(value)
    return None if parsed is None else parsed * factor


def _format_scaled_amount(value: Any) -> str:
    parsed = _to_float(value)
    if parsed is None:
        return ""
    return f"{parsed:.6f}".rstrip("0").rstrip(".")


def _format_mmol(value: Any) -> str:
    parsed = _to_float(value)
    return _format_decimal(parsed, 2) if parsed is not None else ""


def _format_equivalents(value: Any) -> str:
    parsed = _to_float(value)
    return _format_decimal(parsed, 2) if parsed is not None else ""


def _format_percent(value: Any) -> str:
    parsed = _to_float(value)
    return f"{parsed:.0f}%" if parsed is not None else ""


def _format_volume_ul(value: Any) -> str:
    parsed = _to_float(value)
    if parsed is None:
        return ""
    if abs(parsed) >= 10:
        return f"{parsed:.0f}"
    return _format_decimal(parsed, 1)


def _format_volume_ml(value: Any) -> str:
    parsed = _to_float(value)
    if parsed is None:
        return ""
    return _format_decimal(parsed, 2)


def _volume_ul(amount: dict[str, Any]) -> float | None:
    volume_ul = _to_float(amount.get("volume_uL"))
    if volume_ul is not None:
        return volume_ul
    volume_ml = _to_float(amount.get("volume_mL"))
    return None if volume_ml is None else volume_ml * 1000


def _volume_ml(amount: dict[str, Any]) -> float | None:
    volume_ml = _to_float(amount.get("volume_mL"))
    if volume_ml is not None:
        return volume_ml
    volume_ul = _to_float(amount.get("volume_uL"))
    return None if volume_ul is None else volume_ul / 1000


def _format_decimal(value: float | None, places: int) -> str:
    if value is None:
        return ""
    return f"{value:.{places}f}".rstrip("0").rstrip(".")


def _to_float(value: Any) -> float | None:
    if value in {None, ""}:
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def _compact_amount(amount: dict[str, Any]) -> dict[str, Any]:
    return {key: value for key, value in amount.items() if value not in {None, ""}}


def _compound_issue(compound: Compound, code: str, message: str) -> Issue:
    return {
        "code": code,
        "severity": "warning",
        "message": message,
        "compound_id": compound.id or compound.number,
    }
