from __future__ import annotations

import hashlib
import re
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentObject
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .domain.compound import Compound
from .domain.elemental_analysis import calculate_elemental_analysis_block, found_from_block
from .domain.ir import parse_ir_block
from .domain.massspec import build_hrms_block, hrms_adduct_text, hrms_found_text, hrms_label_text
from .domain.references import format_reference
from .domain.reactions import calculate_reaction_loadings
from .mnova_ole import MnovaOleTarget, embed_mnova_ole_objects, preview_size_pt
from .render.si_document import DocumentBlock, SIDocument
from .runtime_paths import bundled_resource_path


DEFAULT_TEMPLATE_RESOURCE = Path("si_generator/templates/SI_template.docx")

PLACEHOLDER_RE = re.compile(r"\[\{([^{}]+)\}\]|\{([^{}]+)\}")
NMR_LABEL_RE = re.compile(r"13C(?:\{1H\})?(?=\s*NMR\b)|1H(?=\s*NMR\b)")
STEREOCHEMISTRY_RE = re.compile(r"\((?:\d*[EZ](?:,\d*[EZ])*)\)")
RF_RE = re.compile(r"\bRf\b")
GP_RE = re.compile(r"\bGP\d+\b")
COMPOUND_NUMBER_RE = re.compile(r"(?<![A-Za-z0-9.])\d+[a-z](?![A-Za-z0-9])")
CHEM_TOKEN_RE = re.compile(r"(?<![A-Za-z0-9])(?:\[[A-Za-z0-9+\-]+\][+\-]|[A-Z][A-Za-z0-9]*[+\-]?)(?![A-Za-z0-9])")
KNOWN_ISOTOPE_LABELS = ("13", "15", "18", "29", "31", "35", "37", "79", "81", "2")


def default_template_path() -> Path:
    candidates = [
        bundled_resource_path(DEFAULT_TEMPLATE_RESOURCE, package_file=__file__),
        Path(__file__).resolve().parent / "templates" / "SI_template.docx",
        Path(__file__).resolve().parents[2] / "examples" / "templates" / "SI_template.docx",
    ]
    for path in candidates:
        if path.exists():
            return path
    return candidates[0]


def render_document_from_template(
    document_model: SIDocument,
    output_path: str | Path,
    template_path: str | Path | None = None,
) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    template = Path(template_path) if template_path else default_template_path()

    template_document = Document(str(template))
    segments = _split_template_segments(template_document)
    output_document = Document(str(template))
    output_document._body.clear_content()
    mnova_ole_targets: list[MnovaOleTarget] = []

    for section in document_model.get("sections", []):
        blocks = section.get("blocks", [])
        if section.get("id") == "compound_descriptions":
            _render_compound_blocks(output_document, segments["compound"], blocks)
        elif section.get("id") == "spectra_appendix" and blocks:
            _render_spectrum_blocks(output_document, segments, blocks, mnova_ole_targets)
        elif section.get("id") == "references" and blocks:
            _render_reference_blocks(output_document, blocks)

    output_document.save(output_path)
    if mnova_ole_targets:
        embed_mnova_ole_objects(output_path, mnova_ole_targets)
    return output_path


def _split_template_segments(document: DocumentObject) -> dict[str, list[Paragraph]]:
    pages: list[list[Paragraph]] = [[]]
    for paragraph in document.paragraphs:
        if _paragraph_has_page_break(paragraph):
            pages.append([])
            continue
        pages[-1].append(paragraph)

    return {
        "compound": pages[0] if pages else [],
        "appendix_1h": pages[1] if len(pages) > 1 else [],
        "appendix_13c": pages[2] if len(pages) > 2 else (pages[1] if len(pages) > 1 else []),
    }


def _paragraph_has_page_break(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath('.//w:br[@w:type="page"]'))


def _render_compound_blocks(document: DocumentObject, template_paragraphs: list[Paragraph], blocks: list[DocumentBlock]) -> None:
    for index, block in enumerate(blocks):
        if index:
            document.add_paragraph()
        first_index = len(document.paragraphs)
        compound: Compound = block["content"]
        values = _compound_values(compound)
        _render_template_paragraphs(document, template_paragraphs, values, compound=compound)
        _add_bookmark_range(document.paragraphs[first_index], document.paragraphs[-1], block.get("bookmark", ""))


def _render_spectrum_blocks(
    document: DocumentObject,
    segments: dict[str, list[Paragraph]],
    blocks: list[DocumentBlock],
    mnova_ole_targets: list[MnovaOleTarget],
) -> None:
    first = True
    for block in blocks:
        if first:
            document.add_page_break()
            first = False
        else:
            document.add_page_break()
        first_index = len(document.paragraphs)
        compound: Compound = block["content"]
        nucleus = str(block.get("nucleus") or "")
        template_paragraphs = segments["appendix_1h"] if nucleus == "1H" else segments["appendix_13c"]
        values = _compound_values(compound)
        values.update(_spectrum_values(compound, nucleus))
        _render_template_paragraphs(
            document,
            template_paragraphs,
            values,
            compound=compound,
            spectrum_block=block,
            mnova_ole_targets=mnova_ole_targets,
        )
        _add_bookmark_range(document.paragraphs[first_index], document.paragraphs[-1], block.get("bookmark", ""))


def _render_reference_blocks(document: DocumentObject, blocks: list[DocumentBlock]) -> None:
    document.add_page_break()
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(0)
    title.add_run("References").bold = True
    for block in blocks:
        content = block.get("content", {})
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        _add_bookmark_range(paragraph, paragraph, block.get("bookmark", ""))
        paragraph.add_run(format_reference(content["reference"], int(content["index"])))


def _render_template_paragraphs(
    document: DocumentObject,
    template_paragraphs: list[Paragraph],
    values: dict[str, str],
    *,
    compound: Compound,
    spectrum_block: DocumentBlock | None = None,
    mnova_ole_targets: list[MnovaOleTarget] | None = None,
) -> None:
    for template_paragraph in template_paragraphs:
        text = template_paragraph.text
        if spectrum_block is None and _is_structure_placeholder_paragraph(text) and not (compound.has_word_structure or compound.structure_path):
            continue
        if _should_skip_paragraph(text, values):
            continue
        if _is_spectrum_artifact_paragraph(text):
            _render_spectrum_artifact(document, spectrum_block, mnova_ole_targets)
            continue

        paragraph = _clone_paragraph(document, template_paragraph)
        _remove_empty_optional_fragments(paragraph, values)
        _replace_placeholders(paragraph, values)
        _apply_inline_formatting(paragraph)
        if compound.nmr_check_warning and ("{compound.support_warning}" in text or "{Product.support.warning}" in text):
            _style_warning_paragraph(paragraph)


def _clone_paragraph(document: DocumentObject, paragraph: Paragraph) -> Paragraph:
    new_p = deepcopy(paragraph._p)
    body = document._body._element
    sect_pr = body.sectPr
    if sect_pr is None:
        body.append(new_p)
    else:
        body.insert(body.index(sect_pr), new_p)
    return Paragraph(new_p, document._body)


def _replace_placeholders(paragraph: Paragraph, values: dict[str, str]) -> None:
    runs = list(paragraph.runs)
    if not runs:
        return

    run_texts = [run.text or "" for run in runs]
    full_text = "".join(run_texts)
    matches = list(PLACEHOLDER_RE.finditer(full_text))
    if not matches:
        return

    spans: list[tuple[int, int, int]] = []
    position = 0
    for index, text in enumerate(run_texts):
        end = position + len(text)
        spans.append((position, end, index))
        position = end

    segments: list[tuple[str, Any]] = []
    cursor = 0
    for match in matches:
        _append_original_run_segments(segments, runs, run_texts, spans, cursor, match.start())
        key = match.group(1) or match.group(2) or ""
        replacement = values.get(_key(key), "")
        if replacement:
            source_index = _run_index_at(spans, match.start())
            segments.append((replacement, deepcopy(runs[source_index]._r.rPr)))
        cursor = match.end()
    _append_original_run_segments(segments, runs, run_texts, spans, cursor, len(full_text))
    _replace_paragraph_runs(paragraph, segments)


def _remove_empty_optional_fragments(paragraph: Paragraph, values: dict[str, str]) -> None:
    runs = list(paragraph.runs)
    if not runs:
        return
    full_text = "".join(run.text or "" for run in runs)
    cleaned = _cleanup_optional_text_fragments(full_text, values)
    if cleaned == full_text:
        return
    first_run = next((run for run in runs if run.text), runs[0])
    _replace_paragraph_runs(paragraph, [(cleaned, deepcopy(first_run._r.rPr))])


def _cleanup_optional_text_fragments(text: str, values: dict[str, str]) -> str:
    cleaned = text
    if not values.get("anal.found"):
        cleaned = _remove_fragment_with_placeholder(cleaned, "anal.found", r"\s*Found\s*:\s*{placeholder}\s*\.?")
    if not _any_value(values, "product.yield.percent", "yield.product.percent", "percent.yield.product"):
        cleaned = _remove_fragment_with_any_placeholder(
            cleaned,
            ("product.yield.percent", "yield.product.percent", "percent.yield.product"),
            r"\s*\(\s*{placeholder}\s*\)",
        )
    if not _any_value(values, "product.mg", "product.yield.mg", "yield.product.mg", "mg.yield.product"):
        cleaned = _remove_fragment_with_any_placeholder(
            cleaned,
            ("product.mg", "product.yield.mg", "yield.product.mg", "mg.yield.product"),
            r"\s*Yield\s+{placeholder}\s*mg\s*(?:\([^)]*\))?\s*[;.]?",
        )
    if not values.get("product.appearance"):
        cleaned = _remove_fragment_with_placeholder(cleaned, "product.appearance", r"\s*;\s*{placeholder}")
        cleaned = _remove_fragment_with_placeholder(cleaned, "product.appearance", r"{placeholder}\s*;\s*")
    if not values.get("product.mp"):
        cleaned = _remove_fragment_with_placeholder(cleaned, "product.mp", r"\s*(?:[;.]\s*)?mp\s+{placeholder}\s*(?:°C|deg\.?\s*C|C)?")
    if not values.get("product.rf.value"):
        cleaned = _remove_fragment_with_placeholder(
            cleaned,
            "product.rf.value",
            r"\s*[;.]\s*Rf\s*=\s*{placeholder}\s*(?:\(\s*" + _placeholder_any_pattern(("product.rf.system",)) + r"\s*\))?",
        )
    if values.get("product.rf.value") and not values.get("product.rf.system"):
        cleaned = _remove_fragment_with_placeholder(cleaned, "product.rf.system", r"\s*\(\s*{placeholder}\s*\)")
    return _normalize_optional_cleanup(cleaned)


def _remove_fragment_with_placeholder(text: str, key: str, pattern_template: str) -> str:
    return _remove_fragment_with_any_placeholder(text, (key,), pattern_template)


def _remove_fragment_with_any_placeholder(text: str, keys: tuple[str, ...], pattern_template: str) -> str:
    for key in keys:
        placeholder = _placeholder_any_pattern((key,))
        text = re.sub(pattern_template.replace("{placeholder}", placeholder), "", text, flags=re.IGNORECASE)
    return text


def _placeholder_any_pattern(keys: tuple[str, ...]) -> str:
    alternatives = []
    for key in keys:
        escaped = re.escape(key)
        alternatives.append(r"\[\{\s*" + escaped + r"\s*\}\]")
        alternatives.append(r"\{\s*" + escaped + r"\s*\}")
    return r"(?:" + "|".join(alternatives) + r")"


def _any_value(values: dict[str, str], *keys: str) -> bool:
    return any(values.get(_key(key)) for key in keys)


def _normalize_optional_cleanup(text: str) -> str:
    text = re.sub(r"\s+([,.;])", r"\1", text)
    text = re.sub(r";\s*;", ";", text)
    text = re.sub(r"\(\s*\)", "", text)
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r";\s*\.", ".", text)
    text = re.sub(r"\s+\.", ".", text)
    return text.strip()


def _append_original_run_segments(
    segments: list[tuple[str, Any]],
    runs: list[Run],
    run_texts: list[str],
    spans: list[tuple[int, int, int]],
    start: int,
    end: int,
) -> None:
    if end <= start:
        return
    for run_start, run_end, run_index in spans:
        overlap_start = max(start, run_start)
        overlap_end = min(end, run_end)
        if overlap_end <= overlap_start:
            continue
        text = run_texts[run_index][overlap_start - run_start : overlap_end - run_start]
        if text:
            segments.append((text, deepcopy(runs[run_index]._r.rPr)))


def _run_index_at(spans: list[tuple[int, int, int]], position: int) -> int:
    for start, end, index in spans:
        if start <= position < end:
            return index
    return spans[-1][2] if spans else 0


def _replace_paragraph_runs(paragraph: Paragraph, segments: list[tuple[str, Any]]) -> None:
    parent = paragraph._p
    insert_index = 0
    for index, child in enumerate(list(parent)):
        if child.tag == qn("w:pPr"):
            insert_index = index + 1
        if child.tag == qn("w:r"):
            parent.remove(child)

    offset = 0
    for text, run_properties in segments:
        if not text:
            continue
        new_r = OxmlElement("w:r")
        if run_properties is not None:
            new_r.append(deepcopy(run_properties))
        text_element = OxmlElement("w:t")
        if text[:1].isspace() or text[-1:].isspace():
            text_element.set(qn("xml:space"), "preserve")
        text_element.text = text
        new_r.append(text_element)
        parent.insert(insert_index + offset, new_r)
        offset += 1


def _apply_inline_formatting(paragraph: Paragraph) -> None:
    if "[[" in paragraph.text:
        return
    for run in list(paragraph.runs):
        text = run.text
        if not text:
            continue
        segments = _inline_format_segments(text)
        if not any(segment[1] for segment in segments):
            continue
        _replace_run_with_segments(paragraph, run, segments)


def _inline_format_segments(text: str) -> list[tuple[str, dict[str, bool]]]:
    patterns = [
        ("stereochemistry", STEREOCHEMISTRY_RE),
        ("nmr", NMR_LABEL_RE),
        ("rf", RF_RE),
        ("gp", GP_RE),
        ("compound_number", COMPOUND_NUMBER_RE),
        ("chem", CHEM_TOKEN_RE),
    ]
    segments: list[tuple[str, dict[str, bool]]] = []
    position = 0
    while position < len(text):
        best: tuple[int, int, int, str, re.Match[str]] | None = None
        for priority, (kind, pattern) in enumerate(patterns):
            match = pattern.search(text, position)
            if not match:
                continue
            candidate = (match.start(), priority, match.end(), kind, match)
            if best is None or candidate[:2] < best[:2]:
                best = candidate
        if best is None:
            segments.append((text[position:], {}))
            break

        start, _priority, end, kind, match = best
        if start > position:
            segments.append((text[position:start], {}))

        token = match.group(0)
        if kind == "stereochemistry":
            segments.extend(_stereochemistry_segments(token))
        elif kind == "nmr":
            segments.extend(_nmr_label_segments(token))
        elif kind == "rf":
            segments.extend([("R", {}), ("f", {"subscript": True})])
        elif kind in {"gp", "compound_number"}:
            segments.append((token, {"bold": True}))
        else:
            chemical_segments = _chemical_token_segments(token)
            if chemical_segments:
                segments.extend(chemical_segments)
            else:
                segments.append((token, {}))
        position = end
    return _merge_adjacent_segments(segments)


def _nmr_label_segments(token: str) -> list[tuple[str, dict[str, bool]]]:
    if token == "1H":
        return [("1", {"superscript": True}), ("H", {})]
    if token == "13C":
        return [("13", {"superscript": True}), ("C", {})]
    return [("13", {"superscript": True}), ("C{", {}), ("1", {"superscript": True}), ("H}", {})]


def _stereochemistry_segments(token: str) -> list[tuple[str, dict[str, bool]]]:
    return [(char, {"italic": True} if char in {"E", "Z"} else {}) for char in token]


def _chemical_token_segments(token: str) -> list[tuple[str, dict[str, bool]]] | None:
    if token.startswith("[") and token[-1:] in {"+", "-"}:
        return [(token[:-1], {}), (token[-1], {"superscript": True})]
    if not any(char.isdigit() for char in token) and token[-1:] not in {"+", "-"}:
        return None
    if not re.fullmatch(r"[A-Z][A-Za-z0-9]*[+\-]?", token):
        return None

    body = token
    charge = ""
    if body[-1:] in {"+", "-"}:
        charge = body[-1]
        body = body[:-1]

    segments: list[tuple[str, dict[str, bool]]] = []
    index = 0
    while index < len(body):
        char = body[index]
        if char.isdigit():
            end = index + 1
            while end < len(body) and body[end].isdigit():
                end += 1
            segments.extend(_formula_digit_segments(body[index:end], next_char=body[end : end + 1]))
            index = end
            continue
        end = index + 1
        while end < len(body) and not body[end].isdigit():
            end += 1
        segments.append((body[index:end], {}))
        index = end
    if charge:
        segments.append((charge, {"superscript": True}))
    return segments


def _formula_digit_segments(digits: str, *, next_char: str) -> list[tuple[str, dict[str, bool]]]:
    if next_char.isupper():
        for label in sorted(KNOWN_ISOTOPE_LABELS, key=len, reverse=True):
            if digits.endswith(label) and len(digits) > len(label):
                return [
                    (digits[: -len(label)], {"subscript": True}),
                    (label, {"superscript": True}),
                ]
    return [(digits, {"subscript": True})]


def _merge_adjacent_segments(segments: list[tuple[str, dict[str, bool]]]) -> list[tuple[str, dict[str, bool]]]:
    merged: list[tuple[str, dict[str, bool]]] = []
    for text, formatting in segments:
        if not text:
            continue
        if merged and merged[-1][1] == formatting:
            merged[-1] = (merged[-1][0] + text, formatting)
        else:
            merged.append((text, formatting))
    return merged or [("", {})]


def _replace_run_with_segments(paragraph: Paragraph, run: Run, segments: list[tuple[str, dict[str, bool]]]) -> None:
    parent = run._r.getparent()
    if parent is None:
        return
    index = parent.index(run._r)
    run_properties = deepcopy(run._r.rPr)
    parent.remove(run._r)
    for offset, (text, formatting) in enumerate(segments):
        new_r = OxmlElement("w:r")
        if run_properties is not None:
            new_r.append(deepcopy(run_properties))
        text_element = OxmlElement("w:t")
        if text[:1].isspace() or text[-1:].isspace():
            text_element.set(qn("xml:space"), "preserve")
        text_element.text = text
        new_r.append(text_element)
        parent.insert(index + offset, new_r)
        new_run = Run(new_r, paragraph)
        if formatting.get("bold"):
            new_run.bold = True
        if formatting.get("italic"):
            new_run.italic = True
        if formatting.get("subscript"):
            new_run.font.superscript = False
            new_run.font.subscript = True
        if formatting.get("superscript"):
            new_run.font.subscript = False
            new_run.font.superscript = True


def _render_spectrum_artifact(
    document: DocumentObject,
    block: DocumentBlock | None,
    mnova_ole_targets: list[MnovaOleTarget] | None = None,
) -> None:
    if not block:
        return
    compound: Compound = block["content"]
    nucleus = str(block.get("nucleus") or "")
    embed_mode = str(block.get("embed_mode") or "png")
    image_path = str(block.get("image_path") or "")
    mnova_path = str(block.get("mnova_path") or "")

    if embed_mode == "mnova" and image_path and mnova_path and mnova_ole_targets is not None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        section = document.sections[-1]
        picture_width = section.page_width - section.left_margin - section.right_margin
        width_pt, height_pt = preview_size_pt(image_path, _emu_to_pt(picture_width))
        marker = f"[[MNOVA_OLE:{len(mnova_ole_targets) + 1}:{compound.number}:{nucleus}]]"
        paragraph.add_run(marker)
        mnova_ole_targets.append(
            MnovaOleTarget(
                marker=marker,
                mnova_path=Path(mnova_path),
                preview_path=Path(image_path),
                width_pt=width_pt,
                height_pt=height_pt,
            )
        )
        return

    if embed_mode in {"png", "mnova"} and image_path:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        section = document.sections[-1]
        picture_width = section.page_width - section.left_margin - section.right_margin
        paragraph.add_run().add_picture(image_path, width=picture_width)


def _placeholder_keys(text: str) -> set[str]:
    return {_key(match.group(1) or match.group(2) or "") for match in PLACEHOLDER_RE.finditer(text)}


def _is_structure_placeholder_paragraph(text: str) -> bool:
    stripped = text.strip()
    if stripped.startswith("[[STRUCTURE:") or stripped.startswith("[[SPECTRUM_STRUCTURE:"):
        return True
    keys = _placeholder_keys(stripped)
    return bool(keys) and keys.issubset({"product.structure", "compound.number.structure", "spectrum.structure.marker"})


def _is_spectrum_artifact_paragraph(text: str) -> bool:
    if "[[SPECTRUM:" in text:
        return True
    keys = _placeholder_keys(text)
    return bool(
        keys
        & {
            "product.nmr.1h.picture",
            "product.nmr.13c.picture",
            "compound.number.nmr.1h.picture",
            "compound.number.nmr.13c.picture",
            "spectrum.picture",
        }
    )


def _should_skip_paragraph(text: str, values: dict[str, str]) -> bool:
    keys = _placeholder_keys(text)
    if not keys:
        return False
    if any(key.startswith("nmr.1h.") for key in keys) and not values.get("nmr.1h.peaks"):
        return True
    if any(key.startswith("nmr.13c.") for key in keys) and not values.get("nmr.13c.peaks"):
        return True
    if "nmr.extra" in keys and not values.get("nmr.extra"):
        return True
    if any(key.startswith("hrms.") for key in keys) and not values.get("hrms.found"):
        return True
    if any(key.startswith("anal.") for key in keys) and not values.get("anal.calculated"):
        return True
    if any(key.startswith("ir.") for key in keys) and not values.get("ir.peaks"):
        return True
    if ("compound.preparation" in keys or "product.preparation" in keys) and not values.get("product.preparation"):
        return True
    if "reaction.loadings" in keys and not values.get("reaction.loadings"):
        return True
    if _has_loadings_placeholders(keys) and not values.get("product.number"):
        return True
    if ("compound.support_warning" in keys or "product.support.warning" in keys) and not values.get("product.support.warning"):
        return True
    return False


def _emu_to_pt(value: int) -> float:
    return float(value) / 12700


def _compound_values(compound: Compound) -> dict[str, str]:
    loadings_values = {str(key): str(value) for key, value in compound.reaction.get("template_values", {}).items()}
    product_values = _product_values(compound)
    values = {
        "compound.name": compound.name,
        "compound.number": compound.number,
        "compound.label": compound.label,
        "compound.number.structure": f"[[STRUCTURE:{compound.number}]]",
        "product.name": compound.name,
        "product.number": compound.number,
        "product.structure": f"[[STRUCTURE:{compound.number}]]",
        "compound.preparation": "" if loadings_values else _summary_text(compound),
        "product.preparation": "" if loadings_values else _summary_text(compound),
        "compound.support_warning": f"(Support check: {compound.nmr_check_warning})" if compound.nmr_check_warning else "",
        "product.support.warning": f"(Support check: {compound.nmr_check_warning})" if compound.nmr_check_warning else "",
        "reaction.loadings": _reaction_loadings_text(compound),
        "nmr.1h.label": _nmr_label_from_text(compound.h1_nmr, "1H NMR"),
        "nmr.1h.conditions": compound.h1_conditions,
        "nmr.1h.peaks": _nmr_peaks_text(compound.h1_nmr),
        "nmr.13c.label": _nmr_label_from_text(compound.c13_nmr, "13C NMR"),
        "nmr.13c.conditions": compound.c13_conditions,
        "nmr.13c.peaks": _nmr_peaks_text(compound.c13_nmr),
        "nmr.extra": compound.extra_nmr.strip(),
    }
    values.update(product_values)
    values.update(_hrms_values(compound))
    values.update(_elemental_values(compound))
    values.update(_ir_values(compound))
    values.update(loadings_values)
    return {_key(key): value for key, value in values.items()}


def _spectrum_values(compound: Compound, nucleus: str) -> dict[str, str]:
    if nucleus == "1H":
        return {
            "spectrum.nucleus": "1H",
            "spectrum.label": _nmr_label_from_text(compound.h1_nmr, "1H NMR"),
            "spectrum.conditions": compound.h1_conditions,
            "spectrum.structure.marker": f"[[SPECTRUM_STRUCTURE:{compound.number}:1H]]",
            "spectrum.picture": f"[[SPECTRUM:{compound.number}:1H]]",
            "compound.number.structure": f"[[SPECTRUM_STRUCTURE:{compound.number}:1H]]",
            "compound.number.nmr.1h.picture": f"[[SPECTRUM:{compound.number}:1H]]",
            "product.structure": f"[[SPECTRUM_STRUCTURE:{compound.number}:1H]]",
            "product.nmr.1h.picture": f"[[SPECTRUM:{compound.number}:1H]]",
        }
    return {
        "spectrum.nucleus": "13C",
        "spectrum.label": _nmr_label_from_text(compound.c13_nmr, "13C NMR"),
        "spectrum.conditions": compound.c13_conditions,
        "spectrum.structure.marker": f"[[SPECTRUM_STRUCTURE:{compound.number}:13C]]",
        "spectrum.picture": f"[[SPECTRUM:{compound.number}:13C]]",
        "compound.number.structure": f"[[SPECTRUM_STRUCTURE:{compound.number}:13C]]",
        "compound.number.nmr.13c.picture": f"[[SPECTRUM:{compound.number}:13C]]",
        "product.structure": f"[[SPECTRUM_STRUCTURE:{compound.number}:13C]]",
        "product.nmr.13c.picture": f"[[SPECTRUM:{compound.number}:13C]]",
    }


def _product_values(compound: Compound) -> dict[str, str]:
    mass_mg = _product_mass_mg(compound)
    rf_value, rf_system = _split_rf(compound.rf)
    return {
        "product.mg": _format_mass_mg(mass_mg),
        "product.g": _format_scaled_amount(_scale_value(mass_mg, 1 / 1000)),
        "product.kg": _format_scaled_amount(_scale_value(mass_mg, 1 / 1_000_000)),
        "product.mmol": "",
        "product.mol": "",
        "product.yield.percent": _product_yield_percent(compound),
        "product.appearance": _compound_appearance(compound),
        "product.mp": _strip_temperature_unit(compound.melting_point),
        "product.rf.value": rf_value,
        "product.rf.system": rf_system,
    }


def _product_mass_mg(compound: Compound) -> float | None:
    match = re.search(r"(\d+(?:[.,]\d+)?)\s*mg\b", compound.yield_text or "", flags=re.IGNORECASE)
    return _to_float(match.group(1)) if match else None


def _product_yield_percent(compound: Compound) -> str:
    match = re.search(r"(\d+(?:[.,]\d+)?)\s*%", compound.yield_text or "")
    if not match:
        return ""
    value = _to_float(match.group(1))
    return _format_scaled_amount(value) + "%" if value is not None else ""


def _split_rf(value: str) -> tuple[str, str]:
    text = str(value or "").strip()
    match = re.match(r"(.+?)\s*\((.+)\)\s*$", text)
    if not match:
        return text, ""
    return match.group(1).strip(), match.group(2).strip()


def _compound_appearance(compound: Compound) -> str:
    return " ".join(part.strip() for part in [compound.color, compound.state] if part and part.strip()).rstrip(".;")


def _strip_temperature_unit(value: str) -> str:
    return re.sub(r"\s*(?:deg\.?\s*C|degrees?\s*C|°C|C)\s*$", "", str(value or "").strip(), flags=re.IGNORECASE)


def _scale_value(value: float | None, factor: float) -> float | None:
    return None if value is None else value * factor


def _format_mass_mg(value: float | None) -> str:
    if value is None:
        return ""
    if abs(value) >= 10:
        return f"{value:.0f}"
    return _format_scaled_amount(value)


def _format_scaled_amount(value: float | None) -> str:
    if value is None:
        return ""
    return f"{value:.6f}".rstrip("0").rstrip(".")


def _to_float(value: Any) -> float | None:
    if value in {None, ""}:
        return None
    try:
        return float(str(value).replace(",", "."))
    except (TypeError, ValueError):
        return None


def _summary_text(compound: Compound) -> str:
    parts = []
    preparation_includes_summary = bool(compound.reaction.get("preparation_includes_summary"))
    if compound.preparation:
        parts.append(compound.preparation.rstrip("."))
    if compound.yield_text and not preparation_includes_summary:
        parts.append(f"Yield {compound.yield_text}")
    appearance = " ".join(part for part in [compound.color, compound.state] if part)
    if appearance and not preparation_includes_summary:
        parts.append(appearance)
    if compound.melting_point and not preparation_includes_summary:
        parts.append(f"mp {compound.melting_point}")
    if compound.rf and not preparation_includes_summary:
        parts.append(compound.rf)
    return "; ".join(parts) + "." if parts else ""


def _reaction_loadings_text(compound: Compound) -> str:
    if not compound.reaction or compound.reaction.get("hide_loadings_line"):
        return ""
    block = compound.reaction
    if not block.get("formatted_text"):
        block = calculate_reaction_loadings(block)
        compound.reaction = block
    text = str(block.get("formatted_text", "")).strip()
    if not text:
        return ""
    return f"Reaction loadings: {text.rstrip('.')}."


def _nmr_peaks_text(value: str) -> str:
    text = " ".join(str(value or "").strip().split())
    if not text:
        return ""
    if "\u03b4" in text:
        text = text.split("\u03b4", 1)[1].strip()
    previous = None
    while text and text != previous:
        previous = text
        text = re.sub(r"^\s*\u03b4\s*", "", text)
        text = re.sub(r"^\s*=\s*", "", text)
    return text.rstrip(" .")


def _nmr_label_from_text(value: str, fallback: str) -> str:
    text = " ".join(str(value or "").strip().split())
    if not text:
        return fallback
    if not re.match(r"^(?:1H|13C(?:\{1H\})?)\s*NMR\b", text, flags=re.IGNORECASE):
        return fallback
    if "(" in text:
        label = text.split("(", 1)[0].strip()
    elif "\u03b4" in text:
        label = text.split("\u03b4", 1)[0].strip()
    else:
        label = ""
    label = label.rstrip(":")
    return label or fallback


def _hrms_values(compound: Compound) -> dict[str, str]:
    found = hrms_found_text(compound.hrms, compound.hrms_found)
    if not (compound.formula and found):
        return {}
    try:
        block = build_hrms_block(
            formula=compound.formula,
            label=hrms_label_text(compound.hrms, compound.hrms_label),
            adduct=hrms_adduct_text(compound.hrms, compound.hrms_adduct),
            found_text=found,
            isotope_labels=compound.hrms.get("isotope_labels") if compound.hrms else None,
        )
    except Exception:
        return {}
    return {
        "hrms.label": str(block.get("label") or ""),
        "hrms.adduct": str(block.get("adduct") or ""),
        "hrms.formula": _formula_with_isotope_labels(str(block.get("ion_formula") or ""), block.get("isotope_labels", {})),
        "hrms.calculated": f"{float(block.get('calculated_mz', 0.0)):.4f}",
        "hrms.found": str(block.get("found_text") or found).replace(",", "."),
    }


def _elemental_values(compound: Compound) -> dict[str, str]:
    if not (compound.formula and compound.elemental_analysis):
        return {}
    if compound.elemental_analysis.get("skip"):
        return {}
    try:
        block = calculate_elemental_analysis_block(compound.formula, found_from_block(compound.elemental_analysis))
    except Exception:
        return {}
    calculated = block.get("calculated", {})
    found = block.get("found", {})
    elements = list(calculated)
    for element in found:
        if element not in elements:
            elements.append(element)
    return {
        "anal.label": "Anal.",
        "anal.formula": str(block.get("formula") or compound.formula),
        "anal.calculated": "; ".join(f"{element}, {calculated[element]:.2f}" for element in elements if element in calculated),
        "anal.found": "; ".join(f"{element}, {found[element]:.2f}" for element in elements if element in found),
        "anal.text": str(block.get("formatted_text") or ""),
        "anal.formatted.text": str(block.get("formatted_text") or ""),
    }


def _ir_values(compound: Compound) -> dict[str, str]:
    block = parse_ir_block(compound.ir)
    peaks = block.get("peaks_cm1", [])
    if not peaks:
        return {}
    return {
        "ir.label": "IR",
        "ir.method": str(block.get("method") or "KBr"),
        "ir.peaks": ", ".join(str(peak) for peak in peaks),
    }


def _formula_with_isotope_labels(formula: str, isotope_labels: Any) -> str:
    if not formula or not isinstance(isotope_labels, dict) or not isotope_labels:
        return formula

    def replace(match: re.Match[str]) -> str:
        element = match.group(0)
        label = isotope_labels.get(element)
        return f"{label}{element}" if label else element

    return re.sub(r"[A-Z][a-z]?", replace, formula)


def _style_warning_paragraph(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(192, 0, 0)


def _has_loadings_placeholders(keys: set[str]) -> bool:
    prefixes = (
        "product.",
        "reagent.",
        "solvent.",
    )
    if any(key.startswith(prefix) for key in keys for prefix in prefixes):
        return True
    return any(
        re.match(r"^[a-z0-9]+(?:_[a-z0-9]+)?\.(?:name|mg|g|kg|mmol|mol|mcl|ml|l|eq)$", key)
        for key in keys
    )


def _add_bookmark_range(start_paragraph: Paragraph, end_paragraph: Paragraph, name: str) -> None:
    if not name:
        return
    bookmark_id = str(int(hashlib.sha1(name.encode("utf-8")).hexdigest()[:8], 16))
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    insert_at = 1 if len(start_paragraph._p) and start_paragraph._p[0].tag == qn("w:pPr") else 0
    start_paragraph._p.insert(insert_at, start)
    end_paragraph._p.append(end)


def _key(value: Any) -> str:
    text = str(value).strip().replace("\u03bc", "u")
    return re.sub(r"[^a-z0-9]+", ".", text.lower()).strip(".")
