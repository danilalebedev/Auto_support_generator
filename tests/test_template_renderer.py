from __future__ import annotations

import base64
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document

from si_generator.docx_builder import build_document_from_model
from si_generator.domain.compound import Compound
from si_generator.render.document_model import build_si_document_model


class TemplateRendererTests(unittest.TestCase):
    def test_replaces_split_placeholder_and_preserves_template_run_style(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            output = root / "support.docx"
            document = Document()
            paragraph = document.add_paragraph("Name ")
            paragraph.add_run("{").italic = True
            paragraph.add_run("Product.name").italic = True
            paragraph.add_run("}").italic = True
            document.save(template)

            compound = Compound(id="cmp_001", number="2a", name="Example compound")
            model = build_si_document_model([compound], spectra_embed_mode="none")
            build_document_from_model(model, output, template_path=template)
            rendered = Document(output)

        self.assertEqual(rendered.paragraphs[0].text, "Name Example compound")
        replacement_runs = [run for run in rendered.paragraphs[0].runs if "Example compound" in run.text]
        self.assertEqual(len(replacement_runs), 1)
        self.assertTrue(replacement_runs[0].italic)

    def test_new_spectrum_picture_alias_renders_artifact(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            image = root / "2a_1H.png"
            image.write_bytes(_tiny_png())
            template = root / "template.docx"
            output = root / "support.docx"
            document = Document()
            document.add_paragraph("{Product.name} ({Product.number})")
            document.add_page_break()
            document.add_paragraph("{Product.name} ({Product.number})")
            document.add_paragraph("{Product.nmr.1h.picture}")
            document.save(template)

            compound = Compound(
                id="cmp_001",
                number="2a",
                name="Example compound",
                h1_image_path=str(image),
                h1_spectrum_path="fid",
            )
            model = build_si_document_model([compound], spectra_embed_mode="png")
            build_document_from_model(model, output, template_path=template)
            text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
            with ZipFile(output) as archive:
                media_files = [name for name in archive.namelist() if name.startswith("word/media/")]

        self.assertNotIn("{Product.nmr.1h.picture}", text)
        self.assertEqual(len(media_files), 1)

    def test_13c_label_uses_mnova_13c_nmr_text(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            output = root / "support.docx"
            document = Document()
            document.add_paragraph("{nmr.13c.label} {nmr.13c.conditions} δ = {nmr.13c.peaks}.")
            document.save(template)

            compound = Compound(
                id="cmp_001",
                number="2a",
                name="Example compound",
                c13_nmr="13C NMR (151 MHz, CDCl3) δ = 167.1, 140.9.",
                c13_conditions="(CDCl3, 150 MHz)",
            )
            model = build_si_document_model([compound], spectra_embed_mode="none")
            build_document_from_model(model, output, template_path=template)
            text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)

        self.assertIn("13C NMR", text)
        self.assertNotIn("13C{1H} NMR", text)

    def test_renders_elemental_analysis_without_empty_found_fragment(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            output = root / "support.docx"
            document = Document()
            document.add_paragraph("{Product.name} ({Product.number})")
            document.add_paragraph("Anal. Calcd for {anal.formula}: {anal.calculated}. Found: {anal.found}.")
            document.save(template)

            compound = Compound(
                id="cmp_001",
                number="2a",
                name="Example compound",
                formula="C11H11BrO2",
                elemental_analysis={"found": {}},
            )
            model = build_si_document_model([compound], spectra_embed_mode="none")
            build_document_from_model(model, output, template_path=template)
            text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)

        self.assertIn("Example compound (2a)", text)
        self.assertIn("Anal. Calcd for C11H11BrO2: C, 51.79; H, 4.35.", text)
        self.assertNotIn("Found: .", text)

    def test_skips_elemental_analysis_when_input_marks_it_disabled(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            output = root / "support.docx"
            document = Document()
            document.add_paragraph("{Product.name} ({Product.number})")
            document.add_paragraph("Anal. Calcd for {anal.formula}: {anal.calculated}. Found: {anal.found}.")
            document.save(template)

            compound = Compound(
                id="cmp_001",
                number="2a",
                name="Example compound",
                formula="C11H11BrO2",
                elemental_analysis={"skip": True},
            )
            model = build_si_document_model([compound], spectra_embed_mode="none")
            build_document_from_model(model, output, template_path=template)
            text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)

        self.assertIn("Example compound (2a)", text)
        self.assertNotIn("Anal. Calcd", text)

    def test_removes_empty_optional_physical_fragments_from_template_line(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            template = root / "template.docx"
            output = root / "support.docx"
            document = Document()
            document.add_paragraph(
                "Yield {Product.mg} mg ({Product.yield.percent}); {Product.appearance}; "
                "mp {Product.mp} °C. Rf = {Product.rf.value} ({Product.rf.system})."
            )
            document.save(template)

            compound = Compound(id="cmp_001", number="2a", name="Example compound")
            model = build_si_document_model([compound], spectra_embed_mode="none")
            build_document_from_model(model, output, template_path=template)
            text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)

        self.assertNotIn("Yield  mg", text)
        self.assertNotIn("()", text)
        self.assertNotIn("; ;", text)
        self.assertNotIn("mp  °C", text)
        self.assertNotIn("Rf =", text)

    def test_visual_template_uses_current_aliases(self) -> None:
        template = Path(__file__).resolve().parents[1] / "examples" / "templates" / "SI_template.docx"
        document = Document(template)
        text_parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.extend(paragraph.text for paragraph in cell.paragraphs)
        text = "\n".join(text_parts)

        legacy_aliases = (
            "number.Product",
            "number.Reagent",
            "name.Reagent",
            "mg.Reagent",
            "mmol.Reagent",
            "mmol.AcOH",
            "yield.Product",
            "spectrum.structure.marker",
            "compound.name",
            "compound.number",
            "compound.number.structure",
            "compound.number.nmr",
            "[[SPECTRUM",
            "[[STRUCTURE",
        )
        current_aliases = (
            "Product.name",
            "Product.number",
            "Reagent_1.number",
            "Reagent_2.name",
            "AcOH.mmol",
            "AcOH.mcl",
            "Product.yield.percent",
            "Product.structure",
            "Product.nmr.1h.picture",
            "Product.nmr.13c.picture",
        )

        for alias in legacy_aliases:
            self.assertNotIn(alias, text)
        for alias in current_aliases:
            self.assertIn(alias, text)


def _tiny_png() -> bytes:
    return base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNk+M9QDwADhgGA"
        "WjR9awAAAABJRU5ErkJggg=="
    )


if __name__ == "__main__":
    unittest.main()
