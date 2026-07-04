from __future__ import annotations

import tempfile
import unittest
import zipfile
from xml.etree import ElementTree
from pathlib import Path

from docx import Document

from si_generator.chemistry import calc_hrms_mz
from si_generator.docx_builder import build_document_from_model
from si_generator.domain.bookmarks import bookmark_name_for_block_id
from si_generator.domain.compound import Compound
from si_generator.render.document_model import build_si_document_model


class DocumentModelTests(unittest.TestCase):
    def test_builds_compound_and_spectrum_sections(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            image_path = Path(tmp) / "2a_1H.png"
            image_path.write_bytes(b"not-a-real-image-for-model-only")
            compound = Compound(id="cmp_001", number="2a", name="Example", h1_image_path=str(image_path))

            model = build_si_document_model([compound])

        self.assertEqual(model["title"], "Supporting Information")
        self.assertEqual([section["id"] for section in model["sections"]], ["compound_descriptions", "spectra_appendix"])
        compound_block = model["sections"][0]["blocks"][0]
        spectrum_block = model["sections"][1]["blocks"][0]
        self.assertEqual(compound_block["kind"], "compound_description")
        self.assertEqual(compound_block["block_id"], "compound:cmp_001")
        self.assertEqual(compound_block["bookmark"], bookmark_name_for_block_id("compound:cmp_001"))
        self.assertEqual(compound_block["display_number"], "2a")
        self.assertEqual(compound_block["title_text"], "Example (2a)")
        self.assertEqual(compound_block["structure_placeholder"], "STRUCTURE:2a")
        self.assertEqual(spectrum_block["kind"], "spectrum_page")
        self.assertEqual(spectrum_block["block_id"], "spectrum:cmp_001:1H")
        self.assertEqual(spectrum_block["bookmark"], bookmark_name_for_block_id("spectrum:cmp_001:1H"))
        self.assertEqual(spectrum_block["compound_id"], "cmp_001")
        self.assertEqual(spectrum_block["nucleus"], "1H")
        self.assertEqual(spectrum_block["structure_placeholder"], "SPECTRUM_STRUCTURE:2a:1H")
        self.assertEqual(spectrum_block["expected_artifact_path"], str(image_path))
        self.assertEqual(model["metadata"]["spectrum_count"], "1")
        self.assertEqual(model["metadata"]["references_count"], "0")

    def test_renders_docx_from_document_model(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "support_information.docx"
            compound = Compound(id="cmp_001", number="2a", name="Example compound", color="white", state="solid")
            model = build_si_document_model([compound])

            build_document_from_model(model, output_path)

            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

        self.assertIn("Example compound", text)
        self.assertIn("(2a)", text)
        self.assertIn("white solid.", text)

    def test_renders_docx_bookmarks_for_patch_workflow(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "support_information.docx"
            compound = Compound(id="cmp_001", number="2a", name="Example compound", color="white", state="solid")
            model = build_si_document_model([compound])

            build_document_from_model(model, output_path)

            with zipfile.ZipFile(output_path) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
            visible_text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

        self.assertIn('w:bookmarkStart', document_xml)
        self.assertIn(f'w:name="{bookmark_name_for_block_id("compound:cmp_001")}"', document_xml)
        self.assertLess(document_xml.index('w:bookmarkStart'), document_xml.index("Example compound"))
        self.assertLess(document_xml.index("white solid."), document_xml.index('w:bookmarkEnd'))
        self.assertNotIn("asig_compound", visible_text)

    def test_renders_hrms_isotope_labels(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "support_information.docx"
            compound = Compound(
                id="cmp_001",
                number="2a",
                name="Bromide example",
                formula="C11H10BrFO2",
                hrms_found="272.9921",
                hrms_adduct="[M+H]+",
            )
            model = build_si_document_model([compound])

            build_document_from_model(model, output_path)

            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)
            runs = _document_runs(output_path)

        self.assertIn("79Br", text)
        self.assertIn("272.9921", text)
        self.assertTrue(_has_run(runs, "79", vert_align="superscript"))

    def test_renders_hrms_from_structured_block(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "support_information.docx"
            found = f"{calc_hrms_mz('C2H4O2', '[M+Na]+'):.4f}"
            compound = Compound(
                id="cmp_001",
                number="2a",
                name="Sodium adduct example",
                formula="C2H4O2",
                hrms={"adduct": "[M+Na]+", "found_text": found},
            )
            model = build_si_document_model([compound])

            build_document_from_model(model, output_path)

            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

        self.assertIn("[M+Na]+", text)
        self.assertIn("C2H4O2Na+", text)
        self.assertIn(found, text)

    def test_renders_hrms_decimal_comma_found_value(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "support_information.docx"
            compound = Compound(
                id="cmp_001",
                number="2a",
                name="Comma HRMS example",
                formula="C2H6O",
                hrms_found="47,0491",
            )
            model = build_si_document_model([compound])

            build_document_from_model(model, output_path)

            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

        self.assertIn("Found 47.0491.", text)

    def test_renders_elemental_analysis_in_journal_format(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "support_information.docx"
            compound = Compound(
                id="cmp_001",
                number="2a",
                name="Analysis example",
                formula="C17H11FN2O3",
                elemental_analysis={"found": "C, 66.03; H, 3.55; N, 8.92"},
            )
            model = build_si_document_model([compound])

            build_document_from_model(model, output_path)

            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

        self.assertIn("Anal. Calcd for C17H11FN2O3: C, 65.81; H, 3.57; N, 9.03. Found: C, 66.03; H, 3.55; N, 8.92.", text)

    def test_skips_formula_based_lines_when_formula_cannot_be_parsed(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "support_information.docx"
            compound = Compound(
                id="cmp_001",
                number="2a",
                name="Bad formula example",
                formula="C2H6Xx",
                hrms_found="47.0491",
                elemental_analysis={"found": "C, 50.00; H, 10.00"},
            )
            model = build_si_document_model([compound])

            build_document_from_model(model, output_path)

            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

        self.assertIn("Bad formula example", text)
        self.assertNotIn("HRMS", text)
        self.assertNotIn("Anal. Calcd", text)

    def test_renders_ir_method_from_input_line(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "support_information.docx"
            compound = Compound(
                id="cmp_001",
                number="2a",
                name="IR example",
                ir="IR (ATR, cm-1): 3038, 2957, 1711.",
            )
            model = build_si_document_model([compound])

            build_document_from_model(model, output_path)

            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

        self.assertIn("IR (ATR, cm-1): 3038, 2957, 1711.", text)
        self.assertNotIn("IR (KBr, cm-1): IR", text)

    def test_renders_chemical_inline_word_formatting(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "support_information.docx"
            compound = Compound(
                id="cmp_001",
                number="3a",
                name="Methyl (E)-3-(2-(bromomethyl)phenyl)acrylate",
                preparation="Bromide 3a was obtained according to GP2. Rf = 0.4 in CDCl3",
                h1_conditions="CDCl3, 600 MHz",
                h1_nmr="δ = δ = 7.26 (s, 1H).",
                c13_conditions="CDCl3, 150 MHz",
                c13_nmr="δ = 167.2 (CO2Me).",
                formula="C17H18FN2O2",
                hrms={"adduct": "[M+H]+", "found_text": "302.1290"},
            )
            model = build_si_document_model([compound])

            build_document_from_model(model, output_path)

            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)
            runs = _document_runs(output_path)

        self.assertIn("1H NMR (CDCl3, 600 MHz) δ = 7.26 (s, 1H).", text)
        self.assertNotIn("δ = δ =", text)
        self.assertNotIn("..", text)
        self.assertTrue(_has_run(runs, "3a", bold=True))
        self.assertTrue(_has_run(runs, "GP2", bold=True))
        self.assertTrue(_has_run(runs, "f", vert_align="subscript"))
        self.assertTrue(_has_run(runs, "13", vert_align="superscript"))
        self.assertTrue(_has_run(runs, "1", vert_align="superscript"))
        self.assertTrue(_has_run(runs, "3", vert_align="subscript"))
        self.assertTrue(_has_run(runs, "17", vert_align="subscript"))
        self.assertTrue(_has_run(runs, "+", vert_align="superscript"))
        self.assertTrue(_has_run(runs, "E", italic=True))

W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
W_VAL = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"


def _document_runs(path: Path) -> list[dict[str, str | bool]]:
    with zipfile.ZipFile(path) as archive:
        xml = archive.read("word/document.xml")
    root = ElementTree.fromstring(xml)
    runs = []
    for run in root.findall(".//w:r", W_NS):
        text = "".join(node.text or "" for node in run.findall("w:t", W_NS))
        if not text:
            continue
        properties = run.find("w:rPr", W_NS)
        vert_align = ""
        bold = False
        italic = False
        if properties is not None:
            bold = properties.find("w:b", W_NS) is not None
            italic = properties.find("w:i", W_NS) is not None
            vert = properties.find("w:vertAlign", W_NS)
            if vert is not None:
                vert_align = str(vert.attrib.get(W_VAL, ""))
        runs.append({"text": text, "bold": bold, "italic": italic, "vert_align": vert_align})
    return runs


def _has_run(
    runs: list[dict[str, str | bool]],
    text: str,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
    vert_align: str | None = None,
) -> bool:
    for run in runs:
        if run["text"] != text:
            continue
        if bold is not None and run["bold"] != bold:
            continue
        if italic is not None and run["italic"] != italic:
            continue
        if vert_align is not None and run["vert_align"] != vert_align:
            continue
        return True
    return False


if __name__ == "__main__":
    unittest.main()
