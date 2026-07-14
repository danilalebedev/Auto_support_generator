from __future__ import annotations

import tempfile
import unittest
import zipfile
from contextlib import redirect_stderr, redirect_stdout
from io import StringIO
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

from si_generator.cli import main as cli_main
from si_generator.word_input import _first_number, read_word_compounds


REPO_ROOT = Path(__file__).resolve().parents[1]


class WordAndDocxSmokeTests(unittest.TestCase):
    def test_reads_example_word_input(self) -> None:
        compounds = read_word_compounds(REPO_ROOT / "examples" / "example_1" / "Compound_table.docx")
        self.assertEqual([compound.number for compound in compounds[:2]], ["2a", "2b"])
        self.assertTrue(compounds[0].name.startswith("Methyl"))
        self.assertTrue(compounds[0].has_word_structure)

    def test_hrms_number_extraction_accepts_decimal_comma(self) -> None:
        self.assertEqual(_first_number("Found 272,9921"), "272.9921")
        self.assertEqual(_first_number("calcd 100.0000. Found 272,9920"), "272.9920")

    def test_word_input_treats_dash_as_missing_optional_field(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "input.docx"
            document = Document()
            table = document.add_table(rows=2, cols=8)
            headers = ["number", "name", "color", "mp", "Rf", "HRMS", "Elemental_analysis", "IR"]
            values = ["2a", "Example compound", "-", "-", "-", "-", "-", "-"]
            for index, header in enumerate(headers):
                table.rows[0].cells[index].text = header
                table.rows[1].cells[index].text = values[index]
            document.save(input_path)

            compounds = read_word_compounds(input_path)

        self.assertEqual(len(compounds), 1)
        compound = compounds[0]
        self.assertEqual(compound.color, "")
        self.assertEqual(compound.state, "")
        self.assertEqual(compound.melting_point, "")
        self.assertEqual(compound.rf, "")
        self.assertEqual(compound.hrms_found, "")
        self.assertEqual(compound.elemental_analysis, {"skip": True})
        self.assertEqual(compound.ir, "")

    def test_cli_generates_docx_without_mnova(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            requested_output_path = Path(tmp) / "support_information.docx"
            output_path = Path(tmp) / "docx" / "support_information.docx"
            stdout = StringIO()
            with redirect_stdout(stdout), redirect_stderr(StringIO()):
                exit_code = cli_main(
                    [
                        "--word-input",
                        str(REPO_ROOT / "examples" / "example_1" / "Compound_table.docx"),
                        "--spectra-zip",
                        str(REPO_ROOT / "examples" / "example_1" / "Spectra_source"),
                        "--no-extract-nmr",
                        "--no-check-support",
                        "--output",
                        str(requested_output_path),
                    ]
                )

            self.assertEqual(exit_code, 0)
            output_path = _generated_path_from_stdout(stdout.getvalue())
            self.assertTrue(output_path.exists())
            self.assertTrue(output_path.with_suffix(".run_summary.json").exists())
            self.assertIn("Run summary:", stdout.getvalue())
            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)
            self.assertIn("Methyl (E)-3-(2-(bromomethyl)phenyl)acrylate", text)
            self.assertNotIn("Compound 2a", text)
            self.assertTrue(_all_structure_ole_objects_use_chemdraw(output_path))


def _all_structure_ole_objects_use_chemdraw(path: Path) -> bool:
    ns = {
        "o": "urn:schemas-microsoft-com:office:office",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    with zipfile.ZipFile(path) as archive:
        document = ET.fromstring(archive.read("word/document.xml"))
        prog_ids = [ole.attrib.get("ProgID", "") for ole in document.findall(".//o:OLEObject", ns)]
        if not prog_ids or any(prog_id != "ChemDraw.Document" for prog_id in prog_ids):
            return False
        for name in archive.namelist():
            if name.startswith("word/embeddings/si_structure_"):
                data = archive.read(name)
                if b"ACD.ChemSketchCDX" in data or b"ChemDraw.Document" not in data:
                    return False
    return True


def _generated_path_from_stdout(output: str) -> Path:
    for line in output.splitlines():
        if line.startswith("Generated "):
            return Path(line.removeprefix("Generated ").strip())
    raise AssertionError(f"Generated path not found in CLI output:\n{output}")


if __name__ == "__main__":
    unittest.main()
