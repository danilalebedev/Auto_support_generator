from __future__ import annotations

import tempfile
import unittest
import json
from pathlib import Path

from docx import Document

from si_generator.graph.state import GenerateSIRequest
from si_generator.workflows.generate_si import output_path_from_state, run_generate_si


REPO_ROOT = Path(__file__).resolve().parents[1]


class GenerateWorkflowTests(unittest.TestCase):
    def test_graph_generates_docx_without_mnova_or_support_check(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            requested_output_path = Path(tmp) / "support_information.docx"
            request = GenerateSIRequest(
                input_path=REPO_ROOT / "examples" / "example_1" / "Compound_table.docx",
                input_kind="word",
                spectra_zip=REPO_ROOT / "examples" / "example_1" / "Spectra_source",
                output_path=requested_output_path,
                no_extract_nmr=True,
                no_check_support=True,
            )

            state = run_generate_si(request)

            output_path = output_path_from_state(state)
            self.assertEqual(output_path.parent.name, "docx")
            self.assertEqual(output_path.parent.parent.parent, Path(tmp) / "runs")
            self.assertTrue(output_path.parent.parent.name.endswith("_Compound_table"))
            self.assertEqual(output_path_from_state(state), output_path)
            self.assertTrue(output_path.exists())
            manifest_path = output_path.with_suffix(".manifest.json")
            self.assertTrue(manifest_path.exists())
            run_summary_path = output_path.with_suffix(".run_summary.json")
            self.assertTrue(run_summary_path.exists())
            self.assertEqual(Path(state["artifacts"]["support_docx"]), output_path)
            self.assertEqual(Path(state["artifacts"]["manifest"]), manifest_path)
            self.assertEqual(Path(state["artifacts"]["run_summary"]), run_summary_path)
            self.assertIn("spectra_root", state["artifacts"])
            self.assertIn("logs_dir", state["artifacts"])
            self.assertIn("1H", state["spectra_plan"]["cmp_001"])
            self.assertIn("13C", state["spectra_plan"]["cmp_001"])
            self.assertIsInstance(state["compounds"]["cmp_001"].nmr_spectra, dict)
            self.assertEqual(state["document_model"]["sections"][0]["id"], "compound_descriptions")
            self.assertEqual(state["document_model"]["sections"][0]["blocks"][0]["compound_id"], "cmp_001")
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            self.assertEqual(manifest["order"][:2], ["cmp_001", "cmp_002"])
            self.assertEqual(manifest["compounds"]["cmp_001"]["number"], "2a")
            self.assertIn("compound_table", manifest["input_hashes"])
            self.assertIn("spectra_source", manifest["input_hashes"])
            self.assertFalse(manifest["configs"]["spectra"]["extract_nmr"])
            self.assertFalse(manifest["configs"]["generation"]["check_support"])
            self.assertEqual(Path(manifest["artifacts"]["support_docx"]), output_path)
            self.assertEqual(Path(manifest["artifacts"]["manifest"]), manifest_path)
            self.assertEqual(Path(manifest["artifacts"]["run_summary"]), run_summary_path)
            self.assertIn("logs_dir", manifest["artifacts"])
            self.assertIn("logs_dir", manifest["output_paths"])
            self.assertEqual(Path(manifest["output_paths"]["run_summary"]), run_summary_path)
            run_summary = json.loads(run_summary_path.read_text(encoding="utf-8"))
            self.assertEqual(run_summary["status"], "completed")
            self.assertEqual(run_summary["compound_count"], len(state["order"]))
            self.assertEqual(run_summary["output_paths"]["support_docx"], str(output_path))
            self.assertEqual(run_summary["artifacts"]["manifest"], str(manifest_path))
            self.assertEqual(run_summary["issue_counts"]["warning"], len(state["issues"]))
            self.assertEqual(run_summary["compounds"][0]["issue_count"], 0)
            self.assertNotIn("cmp_001", run_summary["compound_issue_counts"])
            self.assertEqual(run_summary["compounds"][0]["issues"], [])
            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)
            self.assertIn("Methyl (E)-3-(2-(bromomethyl)phenyl)acrylate", text)
            self.assertNotIn("Compound 2a", text)

    def test_graph_can_calculate_elemental_analysis_for_all_formula_rows(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            input_csv = root / "input.csv"
            input_csv.write_text("number,name,formula\n1a,Ethanol,C2H6O\n", encoding="utf-8")
            request = GenerateSIRequest(
                input_path=input_csv,
                input_kind="csv",
                output_path=root / "support_information.docx",
                insert_spectra_as="none",
                no_extract_nmr=True,
                no_check_support=True,
                calculate_elemental_analysis=True,
            )

            state = run_generate_si(request)
            output_path = output_path_from_state(state)
            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

        self.assertTrue(state["generation_config"]["calculate_elemental_analysis"])
        self.assertIn("Anal. Calcd for C2H6O: C, 52.14; H, 13.13.", text)


if __name__ == "__main__":
    unittest.main()
