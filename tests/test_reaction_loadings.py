from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from docx import Document

from si_generator.cli import _build_parser
from si_generator.docx_builder import build_document_from_model
from si_generator.domain.loadings_workflow import LoadingsWorkflowPaths, apply_loadings_workflow, read_reaction_schema, read_scope
from si_generator.domain.requests import GenerateSIRequest
from si_generator.domain.reactions import calculate_reaction_loadings, format_reagent_amount, reaction_from_fields
from si_generator.graph.compound_store import make_compound_store
from si_generator.graph.nodes.loadings import calculate_loadings_node
from si_generator.input_table import read_compounds
from si_generator.domain.compound import Compound
from si_generator.render.document_model import build_si_document_model
from si_generator.structure_metadata import extract_structure_metadata_by_cell
from si_generator.workflows.generate_si import request_from_args


REPO_ROOT = Path(__file__).resolve().parents[1]
EXAMPLES_DIR = REPO_ROOT / "examples"
LOADINGS_DIR = EXAMPLES_DIR / "example_3"


class ReactionLoadingsTests(unittest.TestCase):
    def test_calculates_mmol_mass_and_density_volume(self) -> None:
        reaction = {
            "target_mmol": 1.5,
            "reagents": [
                {
                    "name": "Piperidine",
                    "equivalents": 0.1,
                    "mw": 85.15,
                    "density_g_mL": 0.862,
                }
            ],
        }

        result = calculate_reaction_loadings(reaction)
        reagent = result["reagents"][0]

        self.assertEqual(reagent["mmol"], 0.15)
        self.assertEqual(reagent["mass_mg"], 12.77)
        self.assertEqual(reagent["volume_uL"], 14.81)
        self.assertIn("Piperidine", result["formatted_text"])

    def test_calculates_solution_volume_from_concentration(self) -> None:
        reaction = {
            "target_mmol": 2.0,
            "reagents": [
                {
                    "name": "NBS solution",
                    "equivalents": 1.2,
                    "concentration_M": 0.5,
                }
            ],
        }

        result = calculate_reaction_loadings(reaction)

        self.assertEqual(result["reagents"][0]["mmol"], 2.4)
        self.assertEqual(result["reagents"][0]["volume_uL"], 4800.0)

    def test_formats_reagent_amount(self) -> None:
        self.assertEqual(
            format_reagent_amount({"name": "NBS", "mass_mg": 320, "mmol": 1.8, "equivalents": 1.2}),
            "NBS (320 mg, 1.8 mmol, 1.2 equiv)",
        )

    def test_reaction_from_tabular_fields(self) -> None:
        reaction = reaction_from_fields(
            {
                "target_mmol": "1.5",
                "reagent_1_name": "Piperidine",
                "reagent_1_equiv": "0.1",
                "reagent_1_mw": "85.15",
                "reagent_1_density_g_ml": "0.862",
            }
        )

        self.assertEqual(reaction["target_mmol"], 1.5)
        self.assertEqual(reaction["reagents"][0]["name"], "Piperidine")
        self.assertEqual(reaction["reagents"][0]["equivalents"], 0.1)
        self.assertEqual(reaction["reagents"][0]["density_g_mL"], 0.862)

    def test_csv_input_reads_reaction_columns(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            input_path = Path(tmp) / "compounds.csv"
            input_path.write_text(
                "number,name,target_mmol,reagent_1_name,reagent_1_equiv,reagent_1_mw\n"
                "2a,Example,1.0,NBS,1.1,177.98\n",
                encoding="utf-8",
            )

            compounds = read_compounds(input_path)

        self.assertEqual(compounds[0].reaction["target_mmol"], 1.0)
        self.assertEqual(compounds[0].reaction["reagents"][0]["name"], "NBS")

    def test_graph_node_runs_when_reaction_data_is_present(self) -> None:
        compound = Compound(
            number="2a",
            name="Example",
            reaction={
                "target_mmol": 1.0,
                "reagents": [{"name": "NBS", "equivalents": 1.1, "mw": 177.98}],
            },
        )
        compounds, order = make_compound_store([compound])

        applied = calculate_loadings_node({"compounds": compounds, "order": order, "generation_config": {"generate_loadings": False}})

        reagent = applied["compounds"]["cmp_001"].reaction["reagents"][0]
        self.assertEqual(reagent["mmol"], 1.1)
        self.assertEqual(reagent["mass_mg"], 195.78)

    def test_graph_node_skips_when_no_reaction_data_and_flag_is_disabled(self) -> None:
        compounds, order = make_compound_store([Compound(number="2a", name="Example")])

        result = calculate_loadings_node({"compounds": compounds, "order": order, "generation_config": {"generate_loadings": False}})

        self.assertEqual(result, {})

    def test_docx_renders_calculated_reaction_loadings(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            output_path = Path(tmp) / "support_information.docx"
            compound = Compound(
                id="cmp_001",
                number="2a",
                name="Example",
                reaction={
                    "target_mmol": 1.0,
                    "reagents": [{"name": "NBS", "equivalents": 1.1, "mw": 177.98}],
                },
            )

            build_document_from_model(build_si_document_model([compound]), output_path)
            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

        self.assertIn("Reaction loadings: NBS (195.78 mg, 1.1 mmol, 1.1 equiv).", text)

    def test_extracts_scope_structure_metadata_by_cell(self) -> None:
        metadata = extract_structure_metadata_by_cell(LOADINGS_DIR / "Scope.docx")

        self.assertEqual(metadata[(1, 2, 1)].formula, "C11H11BrO2")
        self.assertEqual(metadata[(1, 2, 3)].formula, "C6H8N2")
        self.assertEqual(metadata[(1, 2, 4)].formula, "C17H18N2O2")
        self.assertAlmostEqual(metadata[(1, 2, 1)].molecular_weight, 255.111, places=3)
        self.assertEqual(metadata[(1, 6, 3)].formula, "C6H6Cl2N2")

    def test_reaction_schema_reads_example_values(self) -> None:
        schema = read_reaction_schema(LOADINGS_DIR / "Reaction_schema.docx")

        self.assertEqual(schema["Reagent_1"].equivalents, 1.0)
        self.assertEqual(schema["Reagent_2"].equivalents, 3.0)
        self.assertEqual(schema["K2CO3"].mw, 138.21)
        self.assertEqual(schema["AcOH"].density_g_mL, 1.049)
        self.assertEqual(schema["Solvent_MeCN"].concentration_M, 0.75)

    def test_scope_can_attach_generated_names_by_cell(self) -> None:
        rows = read_scope(
            LOADINGS_DIR / "Scope.docx",
            structure_names_by_cell={(1, 2, 3): "benzene-1,2-diamine"},
        )

        self.assertEqual(rows[0].reagent_2.name, "benzene-1,2-diamine")
        self.assertEqual(rows[0].reagents["Reagent_2"].name, "benzene-1,2-diamine")
        self.assertEqual(rows[0].reagent_2.formula, "C6H8N2")

    def test_cli_args_accept_explicit_loadings_files(self) -> None:
        args = _build_parser().parse_args(
            [
                "--word-input",
                "input.docx",
                "--output",
                "support_information.docx",
                "--generate-loadings",
                "--loadings-schema-docx",
                "Reaction_schema.docx",
                "--loadings-scope-docx",
                "Scope.docx",
            ]
        )

        request = request_from_args(args)

        self.assertTrue(request.generate_loadings)
        self.assertEqual(request.loadings_schema_docx, Path("Reaction_schema.docx"))
        self.assertEqual(request.loadings_scope_docx, Path("Scope.docx"))

    def test_loadings_workflow_generates_preparation_from_examples(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            scope_path = _filtered_example_scope(Path(tmp), "3a")
            compound = Compound(
                number="3a",
                name="Example",
                color="white",
                state="solid",
                melting_point="82",
                rf="0.38 (petroleum ether : ethyl acetate = 7 : 1)",
            )

            issues = apply_loadings_workflow(
                [compound],
                EXAMPLES_DIR,
                paths=LoadingsWorkflowPaths(LOADINGS_DIR / "Reaction_schema.docx", scope_path),
            )

        self.assertTrue(compound.preparation)
        self.assertNotIn("{", compound.preparation)
        self.assertIn("bromide 2a (400 mg, 1.57 mmol)", compound.preparation)
        self.assertIn("K2CO3 (217 mg, 1.57 mmol)", compound.preparation)
        self.assertIn("AcOH (449 µL, 7.84 mmol)", compound.preparation)
        self.assertIn("Rf = 0.38 (petroleum ether : ethyl acetate = 7 : 1)", compound.preparation)
        self.assertEqual(compound.yield_text, "304 mg (69%)")
        self.assertEqual(compound.formula, "C17H18N2O2")
        self.assertTrue(compound.reaction["preparation_includes_summary"])
        self.assertTrue(compound.reaction["hide_loadings_line"])
        self.assertEqual(issues, [])

    def test_loadings_workflow_reports_full_scope_input_mismatch(self) -> None:
        compound = Compound(
            number="2a",
            name="Example",
            color="white",
            state="solid",
            melting_point="82",
            rf="0.38 (petroleum ether : ethyl acetate = 7 : 1)",
        )

        issues = apply_loadings_workflow(
            [compound],
            EXAMPLES_DIR,
            paths=LoadingsWorkflowPaths(
                LOADINGS_DIR / "Reaction_schema.docx",
                LOADINGS_DIR / "Scope.docx",
            ),
        )

        issue_codes = [issue["code"] for issue in issues]
        self.assertIn("LOADINGS_SCOPE_INPUT_MISMATCH", issue_codes)
        self.assertFalse(compound.preparation)
        mismatch = next(issue for issue in issues if issue["code"] == "LOADINGS_SCOPE_INPUT_MISMATCH")
        self.assertEqual(mismatch["severity"], "error")
        self.assertIn("Input compounds: 2a", mismatch["message"])
        self.assertIn("Scope products: 3a", mismatch["message"])

    def test_loadings_workflow_supports_canonical_reagent_aliases(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "SI_template.docx"
            document = Document()
            document.add_paragraph("Alkene {Product.number} used {Reagent_2.name} ({Reagent_2.mg} mg, {Reagent_2.mmol} mmol).")
            document.save(template_path)
            scope_path = _filtered_example_scope(Path(tmp), "3a")
            compound = Compound(number="3a", name="Example", color="white", state="solid")
            issues = apply_loadings_workflow(
                [compound],
                EXAMPLES_DIR,
                paths=LoadingsWorkflowPaths(
                    LOADINGS_DIR / "Reaction_schema.docx",
                    scope_path,
                    template_path,
                ),
                structure_names_by_cell={(1, 2, 3): "benzene-1,2-diamine"},
            )

        self.assertIn("benzene-1,2-diamine (509 mg, 4.7 mmol)", compound.preparation)
        self.assertEqual(issues, [])

    def test_loadings_workflow_supports_entity_first_aliases(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "SI_template.docx"
            document = Document()
            document.add_paragraph(
                "Alkene {Product.precursor_number} from bromide {Reagent_1.number} gave product {Product.number} with "
                "{Reagent_2.name} ({Reagent_2.mg} mg, {Reagent_2.mmol} mmol). "
                "Yield {Product.mg} mg ({Product.yield.percent})."
            )
            document.save(template_path)
            scope_path = _filtered_example_scope(Path(tmp), "3a")
            compound = Compound(number="3a", name="Example", color="white", state="solid")
            apply_loadings_workflow(
                [compound],
                EXAMPLES_DIR,
                paths=LoadingsWorkflowPaths(
                    LOADINGS_DIR / "Reaction_schema.docx",
                    scope_path,
                    template_path,
                ),
                structure_names_by_cell={(1, 2, 3): "benzene-1,2-diamine"},
            )

        self.assertIn("Alkene 2a from bromide 2a gave product 3a", compound.preparation)
        self.assertIn("benzene-1,2-diamine (509 mg, 4.7 mmol)", compound.preparation)
        self.assertIn("Yield 304 mg (69%)", compound.preparation)
        self.assertNotIn("{", compound.preparation)

    def test_loadings_workflow_supports_named_reagents_without_reagent_2(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            schema_path = root / "Reaction_schema.docx"
            schema_doc = Document()
            table = schema_doc.add_table(rows=1, cols=5)
            for index, header in enumerate(["Reagents", "equiv.", "MW, g/mol", "Density, g/ml", "Concentration, M"]):
                table.rows[0].cells[index].text = header
            for values in [
                ["Reagent_1", "1", "", "", ""],
                ["NBS", "1", "177.98", "", ""],
                ["DBP", "0.05", "242.23", "", ""],
                ["Solvent_MeCN", "", "", "", "0.75"],
            ]:
                cells = table.add_row().cells
                for index, value in enumerate(values):
                    cells[index].text = value
            schema_doc.save(schema_path)

            template_path = root / "SI_template.docx"
            template_doc = Document()
            template_doc.add_paragraph(
                "Bromide {Product.number} from {Reagent_1.name} ({Reagent_1.mg} mg, {Reagent_1.mmol} mmol), "
                "NBS ({NBS.mg} mg, {NBS.mmol} mmol), DBP ({DBP.mg} mg, {DBP.mmol} mmol), "
                "MeCN ({Solvent_MeCN.ml} mL). Yield {Product.mg} mg ({Product.yield.percent})."
            )
            template_doc.save(template_path)

            scope_path = _filtered_example_scope(root, "3a")
            compound = Compound(number="3a", name="Example", color="white", state="solid")
            issues = apply_loadings_workflow(
                [compound],
                EXAMPLES_DIR,
                paths=LoadingsWorkflowPaths(schema_path, scope_path, template_path),
                structure_names_by_cell={(1, 2, 1): "bromide 2a"},
            )

        self.assertNotIn("LOADINGS_SCHEMA_INCOMPLETE", {issue["code"] for issue in issues})
        self.assertIn("bromide 2a (400 mg, 1.57 mmol)", compound.preparation)
        self.assertIn("NBS (279 mg, 1.57 mmol)", compound.preparation)
        self.assertIn("DBP (19 mg, 0.08 mmol)", compound.preparation)
        self.assertIn("MeCN (2.09 mL)", compound.preparation)
        self.assertIn("Yield 304 mg (69%)", compound.preparation)
        self.assertNotIn("{", compound.preparation)

    def test_loadings_workflow_supports_arbitrary_numbered_reagents(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            schema_path = root / "Reaction_schema.docx"
            schema_doc = Document()
            schema_table = schema_doc.add_table(rows=1, cols=5)
            for index, header in enumerate(["Reagents", "equiv.", "MW, g/mol", "Density, g/ml", "Concentration, M"]):
                schema_table.rows[0].cells[index].text = header
            for values in [
                ["Reagent_1", "1", "100", "", ""],
                ["Reagent_2", "2", "50", "", ""],
                ["Reagent_3", "0.5", "200", "", ""],
            ]:
                cells = schema_table.add_row().cells
                for index, value in enumerate(values):
                    cells[index].text = value
            schema_doc.save(schema_path)

            scope_path = root / "Scope.docx"
            scope_doc = Document()
            scope_table = scope_doc.add_table(rows=1, cols=7)
            for index, header in enumerate(
                ["Reagent_1", "Mass of Reagent_1, mg", "Reagent_2", "Reagent_3", "Product", "Product_number", "Mass of product, mg"]
            ):
                scope_table.rows[0].cells[index].text = header
            cells = scope_table.add_row().cells
            for index, value in enumerate(["", "100", "", "", "", "7a", "42"]):
                cells[index].text = value
            scope_doc.save(scope_path)

            template_path = root / "SI_template.docx"
            template_doc = Document()
            template_doc.add_paragraph(
                "{Reagent_1.name} ({Reagent_1.mg} mg, {Reagent_1.mmol} mmol); "
                "{Reagent_2.name} ({Reagent_2.mg} mg, {Reagent_2.mmol} mmol); "
                "{Reagent_3.name} ({Reagent_3.mg} mg, {Reagent_3.mmol} mmol)."
            )
            template_doc.save(template_path)

            rows = read_scope(scope_path, structure_names_by_cell={(1, 2, 1): "A", (1, 2, 3): "B", (1, 2, 4): "C"})
            compound = Compound(number="7a", name="Example", color="white", state="solid")
            issues = apply_loadings_workflow(
                [compound],
                root,
                paths=LoadingsWorkflowPaths(schema_path, scope_path, template_path),
                structure_names_by_cell={(1, 2, 1): "A", (1, 2, 3): "B", (1, 2, 4): "C"},
            )

        self.assertEqual(rows[0].reagents["Reagent_3"].name, "C")
        self.assertEqual(rows[0].reagent_masses_mg["Reagent_1"], 100)
        self.assertNotIn("LOADINGS_SCHEMA_INCOMPLETE", {issue["code"] for issue in issues})
        self.assertIn("A (100 mg, 1 mmol)", compound.preparation)
        self.assertIn("B (100 mg, 2 mmol)", compound.preparation)
        self.assertIn("C (100 mg, 0.5 mmol)", compound.preparation)
        self.assertNotIn("{", compound.preparation)

    def test_loadings_template_omits_missing_optional_product_fields(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            schema_path = root / "Reaction_schema.docx"
            schema_doc = Document()
            schema_table = schema_doc.add_table(rows=1, cols=5)
            for index, header in enumerate(["Reagents", "equiv.", "MW, g/mol", "Density, g/ml", "Concentration, M"]):
                schema_table.rows[0].cells[index].text = header
            for values in [["Reagent_1", "1", "100", "", ""]]:
                cells = schema_table.add_row().cells
                for index, value in enumerate(values):
                    cells[index].text = value
            schema_doc.save(schema_path)

            scope_path = root / "Scope.docx"
            scope_doc = Document()
            scope_table = scope_doc.add_table(rows=1, cols=5)
            for index, header in enumerate(["Reagent_1", "Mass of Reagent_1, mg", "Product", "Product_number", "Mass of product, mg"]):
                scope_table.rows[0].cells[index].text = header
            cells = scope_table.add_row().cells
            for index, value in enumerate(["", "100", "", "7a", ""]):
                cells[index].text = value
            scope_doc.save(scope_path)

            template_path = root / "SI_template.docx"
            template_doc = Document()
            template_doc.add_paragraph(
                "Product {Product.number} from {Reagent_1.name} ({Reagent_1.mg} mg, {Reagent_1.mmol} mmol). "
                "Yield {Product.mg} mg ({Product.yield.percent}); {Product.appearance}; "
                "mp {Product.mp} °C. Rf = {Product.rf.value} ({Product.rf.system})."
            )
            template_doc.save(template_path)

            compound = Compound(number="7a", name="Example")
            apply_loadings_workflow(
                [compound],
                root,
                paths=LoadingsWorkflowPaths(schema_path, scope_path, template_path),
                structure_names_by_cell={(1, 2, 1): "A"},
            )

        self.assertIn("Product 7a from A (100 mg, 1 mmol)", compound.preparation)
        self.assertNotIn("Yield  mg", compound.preparation)
        self.assertNotIn("()", compound.preparation)
        self.assertNotIn("; ;", compound.preparation)
        self.assertNotIn("mp  °C", compound.preparation)
        self.assertNotIn("Rf =", compound.preparation)

    def test_loadings_workflow_supports_canonical_unit_conversions(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            template_path = Path(tmp) / "SI_template.docx"
            document = Document()
            document.add_paragraph(
                "Product {Product.mg} mg, {Product.g} g, {Product.kg} kg, {Product.mmol} mmol, {Product.mol} mol; "
                "AcOH {AcOH.mcl} uL, {AcOH.ml} mL, {AcOH.l} L; "
                "{Solvent_MeCN.name} {Solvent_MeCN.mcl} uL, {Solvent_MeCN.ml} mL, {Solvent_MeCN.l} L; "
                "Reagent {Reagent_2.mol} mol."
            )
            document.save(template_path)
            scope_path = _filtered_example_scope(Path(tmp), "3a")
            compound = Compound(number="3a", name="Example", color="white", state="solid")
            apply_loadings_workflow(
                [compound],
                EXAMPLES_DIR,
                paths=LoadingsWorkflowPaths(
                    LOADINGS_DIR / "Reaction_schema.docx",
                    scope_path,
                    template_path,
                ),
                structure_names_by_cell={(1, 2, 3): "benzene-1,2-diamine"},
            )

        self.assertIn("Product 304 mg, 0.304 g, 0.000304 kg, 1.08 mmol, 0.001077 mol", compound.preparation)
        self.assertIn("AcOH 449 uL, 0.45 mL, 0.000449 L", compound.preparation)
        self.assertIn("MeCN 2091 uL, 2.09 mL, 0.002091 L", compound.preparation)
        self.assertIn("Reagent 0.004704 mol", compound.preparation)

    def test_loadings_node_uses_examples_dir_when_enabled(self) -> None:
        compounds, order = make_compound_store(
            [
                Compound(number="3a", name="Example", color="white solid", melting_point="82", rf="0.38 (petroleum ether : ethyl acetate = 7 : 1)"),
                Compound(number="3b", name="Example", color="white solid", melting_point="84", rf="0.63 (petroleum ether : ethyl acetate = 7 : 1)"),
                Compound(number="3c", name="Example", color="white solid", melting_point="85", rf="0.56 (petroleum ether : ethyl acetate = 7 : 1)"),
                Compound(number="3d", name="Example", color="white solid", melting_point="97", rf="0.48 (petroleum ether : ethyl acetate = 7 : 1)"),
                Compound(number="3i", name="Example", color="Green oil", melting_point="98", rf="0.55 (petroleum ether : ethyl acetate = 7 : 1)"),
            ]
        )
        request = GenerateSIRequest(
            input_path=EXAMPLES_DIR / "example_3" / "Compound_table.docx",
            input_kind="word",
            output_path=Path("out.docx"),
        )

        result = calculate_loadings_node(
            {
                "request": request,
                "compounds": compounds,
                "order": order,
                "generation_config": {"generate_loadings": True},
                "issues": [],
            }
        )

        self.assertIn("compounds", result)
        self.assertNotIn("issues", result)
        prepared = result["compounds"]["cmp_001"]
        self.assertIn("Alkene 3a was obtained", prepared.preparation)
        self.assertAlmostEqual(prepared.reaction["target_mmol"], 1.5679, places=4)
        self.assertEqual(prepared.reaction["source"], "loadings_workflow")

    def test_loadings_node_accepts_explicit_workflow_files(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            scope_path = _filtered_example_scope(Path(tmp), "3a")
            compounds, order = make_compound_store(
                [
                    Compound(
                        number="3a",
                        name="Example",
                        color="white",
                        state="solid",
                        melting_point="82",
                        rf="0.38 (petroleum ether : ethyl acetate = 7 : 1)",
                    )
                ]
            )
            request = GenerateSIRequest(
                input_path=Path("input.docx"),
                input_kind="word",
                output_path=Path("out.docx"),
                loadings_schema_docx=LOADINGS_DIR / "Reaction_schema.docx",
                loadings_scope_docx=scope_path,
            )

            result = calculate_loadings_node(
                {
                    "request": request,
                    "compounds": compounds,
                    "order": order,
                    "generation_config": {"generate_loadings": True},
                    "issues": [],
                }
            )

        prepared = result["compounds"]["cmp_001"]
        self.assertIn("bromide 2a (400 mg, 1.57 mmol)", prepared.preparation)
        self.assertEqual(prepared.reaction["source"], "loadings_workflow")

    def test_loadings_workflow_does_not_depend_on_render_fallback(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            scope_path = _filtered_example_scope(Path(tmp), "3a")
            compound = Compound(
                id="cmp_001",
                number="3a",
                name="Example",
                color="white solid",
                melting_point="82",
                rf="0.38 (petroleum ether : ethyl acetate = 7 : 1)",
            )
            apply_loadings_workflow(
                [compound],
                EXAMPLES_DIR,
                paths=LoadingsWorkflowPaths(LOADINGS_DIR / "Reaction_schema.docx", scope_path),
            )
            output_path = Path(tmp) / "support_information.docx"
            renderer = __import__("si_generator.template_renderer", fromlist=["calculate_reaction_loadings"])
            original = renderer.calculate_reaction_loadings
            try:
                renderer.calculate_reaction_loadings = _raise_render_fallback
                build_document_from_model(build_si_document_model([compound]), output_path)
            finally:
                renderer.calculate_reaction_loadings = original

            text = "\n".join(paragraph.text for paragraph in Document(output_path).paragraphs)

        self.assertIn("Alkene 3a was obtained", text)
        self.assertNotIn("Reaction loadings:", text)


def _raise_render_fallback(_reaction):
    raise AssertionError("loadings must be calculated before rendering")


def _filtered_example_scope(root: Path, *numbers: str) -> Path:
    keep = set(numbers)
    document = Document(LOADINGS_DIR / "Scope.docx")
    table = document.tables[0]
    for row in list(table.rows[1:]):
        product_number = row.cells[4].text.strip()
        if product_number not in keep:
            table._tbl.remove(row._tr)
    path = root / "Scope.docx"
    document.save(path)
    return path


if __name__ == "__main__":
    unittest.main()
