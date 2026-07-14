from __future__ import annotations

import json
import base64
import tempfile
import unittest
import zipfile
from contextlib import redirect_stderr, redirect_stdout
from io import StringIO
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

from si_generator.cli import main as cli_main
from si_generator.docx_builder import build_document_from_model
from si_generator.domain.bookmarks import bookmark_name_for_block_id
from si_generator.domain.manifest import check_manifest, manifest_has_errors
from si_generator.domain.patching import parse_remove_list, parse_renumber_map, parse_reorder_list, parse_swap_pairs
from si_generator.graph.state import PatchSIRequest
from si_generator.domain.compound import Compound
from si_generator.render.document_model import build_si_document_model
from si_generator.workflows.patch_si import run_patch_si


class PatchWorkflowTests(unittest.TestCase):
    def test_parse_renumber_map_accepts_comma_separated_pairs(self) -> None:
        self.assertEqual(parse_renumber_map("2a=3a, cmp_002=3b"), {"2a": "3a", "cmp_002": "3b"})

    def test_parse_renumber_map_rejects_invalid_items(self) -> None:
        with self.assertRaisesRegex(ValueError, "OLD=NEW"):
            parse_renumber_map("2a")

    def test_parse_reorder_list_accepts_numbers_or_ids(self) -> None:
        self.assertEqual(parse_reorder_list("2b, cmp_001"), ("2b", "cmp_001"))

    def test_parse_remove_list_accepts_numbers_or_ids(self) -> None:
        self.assertEqual(parse_remove_list("2a, cmp_002"), ("2a", "cmp_002"))

    def test_parse_swap_pairs_accepts_any_number_of_non_overlapping_pairs(self) -> None:
        self.assertEqual(
            parse_swap_pairs("2a=3a, 2b=3b, cmp_005=4a"),
            (("2a", "3a"), ("2b", "3b"), ("cmp_005", "4a")),
        )

    def test_parse_swap_pairs_rejects_reused_compound(self) -> None:
        with self.assertRaisesRegex(ValueError, "more than one swap pair"):
            parse_swap_pairs("2a=3a,2a=4a")

    def test_patch_workflow_renumbers_docx_and_manifest(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source_docx, source_manifest = _write_source_support(root)
            state = run_patch_si(
                PatchSIRequest(
                    manifest_path=source_manifest,
                    renumber={"2a": "5a"},
                    output_folder=root / "patches",
                )
            )
            patched_docx, patched_manifest, report_path = _patch_artifact_paths(state)
            patched = json.loads(patched_manifest.read_text(encoding="utf-8"))
            patch_report = json.loads(report_path.read_text(encoding="utf-8"))
            text = "\n".join(paragraph.text for paragraph in Document(patched_docx).paragraphs)
            issues = check_manifest(patched, manifest_path=patched_manifest)
            patched_docx_exists = patched_docx.exists()
            patched_manifest_exists = patched_manifest.exists()

        self.assertEqual(state["status"], "pass")
        self.assertTrue(patched_docx_exists)
        self.assertTrue(patched_manifest_exists)
        self.assertEqual(Path(state["artifacts"]["patch_report"]), report_path)
        self.assertEqual(patch_report["status"], "pass")
        self.assertEqual(patch_report["operations"]["renumber"], {"2a": "5a"})
        self.assertEqual(state["patch_result"]["renumbered"], {"2a": "5a"})
        self.assertEqual(patch_report["patch_result"]["renumbered"], {"2a": "5a"})
        self.assertEqual(patch_report["patch_result"]["removed_ids"], [])
        self.assertEqual(patch_report["patch_result"]["reordered_ids"], [])
        self.assertEqual(patch_report["compound_issue_counts"], {})
        self.assertEqual(len(patched["patch_history"]), 1)
        self.assertEqual(patched["patch_history"][0]["run_id"], state["run_id"])
        self.assertEqual(patched["patch_history"][0]["source_manifest"], str(source_manifest))
        self.assertEqual(patched["patch_history"][0]["output_manifest"], str(patched_manifest))
        self.assertEqual(patched["patch_history"][0]["output_docx"], str(patched_docx))
        self.assertEqual(patched["patch_history"][0]["operations"]["renumber"], {"2a": "5a"})
        self.assertEqual(patched["patch_history"][0]["result"]["renumbered"], {"2a": "5a"})
        self.assertEqual(patched["compounds"]["cmp_001"]["number"], "5a")
        self.assertEqual(patched["compounds"]["cmp_001"]["domain_snapshot"]["number"], "5a")
        self.assertIn("Example (5a)", text)
        self.assertNotIn("Example (2a)", text)
        self.assertFalse(manifest_has_errors(issues))

    def test_cli_patch_manifest_mode(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, source_manifest = _write_source_support(root)
            stdout = StringIO()

            with redirect_stdout(stdout), redirect_stderr(StringIO()):
                exit_code = cli_main(
                    [
                        "--patch-manifest",
                        str(source_manifest),
                        "--renumber",
                        "2a=6a",
                        "--patch-output-folder",
                        str(root / "patches"),
                    ]
                )
            patched_docx = next((root / "patches" / "runs").glob("*/docx/support_information.docx"))
            text = "\n".join(paragraph.text for paragraph in Document(patched_docx).paragraphs)

        self.assertEqual(exit_code, 0)
        self.assertIn("Patch check passed", stdout.getvalue())
        self.assertIn("Patch report:", stdout.getvalue())
        self.assertIn("Example (6a)", text)

    def test_cli_patch_manifest_remove_mode(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, source_manifest = _write_source_support(
                root,
                [
                    Compound(id="cmp_001", number="2a", name="Example A", color="white", state="solid"),
                    Compound(id="cmp_002", number="2b", name="Example B", color="yellow", state="solid"),
                ],
            )
            stdout = StringIO()

            with redirect_stdout(stdout), redirect_stderr(StringIO()):
                exit_code = cli_main(
                    [
                        "--patch-manifest",
                        str(source_manifest),
                        "--remove",
                        "2a",
                        "--patch-output-folder",
                        str(root / "patches"),
                    ]
                )
            patched_docx = next((root / "patches" / "runs").glob("*/docx/support_information.docx"))
            text = "\n".join(paragraph.text for paragraph in Document(patched_docx).paragraphs)

        self.assertEqual(exit_code, 0)
        self.assertIn("Patch check passed", stdout.getvalue())
        self.assertNotIn("Example A (2a)", text)
        self.assertIn("Example B (2b)", text)

    def test_cli_patch_manifest_supports_multiple_swap_pairs(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, source_manifest = _write_source_support(
                root,
                [
                    Compound(id="cmp_001", number="2a", name="Alpha"),
                    Compound(id="cmp_002", number="2b", name="Beta"),
                    Compound(id="cmp_003", number="3a", name="Gamma"),
                    Compound(id="cmp_004", number="3b", name="Delta"),
                ],
            )
            stdout = StringIO()
            with redirect_stdout(stdout), redirect_stderr(StringIO()):
                exit_code = cli_main(
                    [
                        "--patch-manifest",
                        str(source_manifest),
                        "--swap",
                        "2a=3a,2b=3b",
                        "--patch-output-folder",
                        str(root / "patches"),
                    ]
                )
            patched_docx = next((root / "patches" / "runs").glob("*/docx/support_information.docx"))
            text = "\n".join(paragraph.text for paragraph in Document(patched_docx).paragraphs)

        self.assertEqual(exit_code, 0)
        self.assertIn("Patch check passed", stdout.getvalue())
        self.assertLess(text.index("Gamma (2a)"), text.index("Delta (2b)"))
        self.assertLess(text.index("Delta (2b)"), text.index("Alpha (3a)"))
        self.assertLess(text.index("Alpha (3a)"), text.index("Beta (3b)"))

    def test_patch_workflow_writes_report_for_malformed_manifest(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest = root / "support_information.manifest.json"
            manifest.write_text("{", encoding="utf-8")

            state = run_patch_si(
                PatchSIRequest(manifest_path=manifest, renumber={"2a": "3a"}, output_folder=root / "patches")
            )
            report_path = Path(state["artifacts"]["patch_report"])
            report = json.loads(report_path.read_text(encoding="utf-8"))

        self.assertEqual(state["status"], "fail")
        self.assertEqual(Path(state["artifacts"]["patch_report"]), report_path)
        self.assertIn("MANIFEST_LOAD_FAILED", {issue["code"] for issue in state["issues"]})
        self.assertEqual(report["patch_result"]["renumbered"], {})

    def test_cli_patch_manifest_mode_returns_report_for_malformed_manifest(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest = root / "support_information.manifest.json"
            manifest.write_text("{", encoding="utf-8")
            stdout = StringIO()

            with redirect_stdout(stdout), redirect_stderr(StringIO()):
                exit_code = cli_main(["--patch-manifest", str(manifest), "--renumber", "2a=3a"])

        self.assertEqual(exit_code, 1)
        self.assertIn("MANIFEST_LOAD_FAILED", stdout.getvalue())
        self.assertIn("Patch report:", stdout.getvalue())

    def test_patch_workflow_writes_report_when_docx_bookmark_is_missing(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source_docx = root / "support_information.docx"
            source_manifest = root / "support_information.manifest.json"
            Document().save(source_docx)
            source_manifest.write_text(
                json.dumps(
                    {
                        "run_id": "run",
                        "artifacts": {"support_docx": str(source_docx), "manifest": str(source_manifest)},
                        "output_paths": {"support_docx": str(source_docx), "manifest": str(source_manifest)},
                        "order": ["cmp_001"],
                        "compounds": {
                            "cmp_001": {
                                "id": "cmp_001",
                                "number": "2a",
                                "docx_block_id": "compound:cmp_001",
                                "docx_bookmark": bookmark_name_for_block_id("compound:cmp_001"),
                            }
                        },
                    }
                ),
                encoding="utf-8",
            )

            state = run_patch_si(
                PatchSIRequest(
                    manifest_path=source_manifest,
                    remove=("2a",),
                    output_folder=root / "patches",
                )
            )
            report_path = Path(state["artifacts"]["patch_report"])
            report = json.loads(report_path.read_text(encoding="utf-8"))
            patched_docx = Path(state["artifacts"]["support_docx"])
            patched_manifest = Path(state["artifacts"]["manifest"])
            patched_docx_exists = patched_docx.exists()
            patched_manifest_exists = patched_manifest.exists()
            temp_outputs = [path.name for path in patched_docx.parent.iterdir() if path.name.startswith(".support_information")]

        self.assertEqual(state["status"], "fail")
        self.assertEqual(Path(state["artifacts"]["patch_report"]), report_path)
        self.assertIn("PATCH_APPLY_FAILED", {issue["code"] for issue in state["issues"]})
        self.assertIn("PATCH_APPLY_FAILED", {issue["code"] for issue in report["issues"]})
        self.assertFalse(patched_docx_exists)
        self.assertFalse(patched_manifest_exists)
        self.assertEqual(temp_outputs, [])

    def test_patch_workflow_reorders_docx_and_manifest(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, source_manifest = _write_source_support(
                root,
                [
                    Compound(id="cmp_001", number="2a", name="Example A", color="white", state="solid"),
                    Compound(id="cmp_002", number="2b", name="Example B", color="yellow", state="solid"),
                ],
            )
            state = run_patch_si(
                PatchSIRequest(
                    manifest_path=source_manifest,
                    reorder=("2b", "2a"),
                    output_folder=root / "patches",
                )
            )
            patched_docx, patched_manifest, report_path = _patch_artifact_paths(state)
            patched = json.loads(patched_manifest.read_text(encoding="utf-8"))
            patch_report = json.loads(report_path.read_text(encoding="utf-8"))
            text = "\n".join(paragraph.text for paragraph in Document(patched_docx).paragraphs)

        self.assertEqual(state["status"], "pass")
        self.assertEqual(state["patch_result"]["reordered_ids"], ["cmp_002", "cmp_001"])
        self.assertEqual(patch_report["patch_result"]["reordered_ids"], ["cmp_002", "cmp_001"])
        self.assertEqual(patched["order"], ["cmp_002", "cmp_001"])
        self.assertLess(text.index("Example B (2b)"), text.index("Example A (2a)"))

    def test_patch_workflow_removes_docx_block_and_manifest_entry(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, source_manifest = _write_source_support(
                root,
                [
                    Compound(id="cmp_001", number="2a", name="Example A", color="white", state="solid"),
                    Compound(id="cmp_002", number="2b", name="Example B", color="yellow", state="solid"),
                ],
            )
            state = run_patch_si(
                PatchSIRequest(
                    manifest_path=source_manifest,
                    remove=("2a",),
                    output_folder=root / "patches",
                )
            )
            patched_docx, patched_manifest, report_path = _patch_artifact_paths(state)
            patched = json.loads(patched_manifest.read_text(encoding="utf-8"))
            patch_report = json.loads(report_path.read_text(encoding="utf-8"))
            text = "\n".join(paragraph.text for paragraph in Document(patched_docx).paragraphs)

        self.assertEqual(state["status"], "pass")
        self.assertEqual(state["patch_result"]["removed_ids"], ["cmp_001"])
        self.assertEqual(patch_report["patch_result"]["removed_ids"], ["cmp_001"])
        self.assertEqual(
            patch_report["patch_result"]["removed_bookmarks"],
            [bookmark_name_for_block_id("compound:cmp_001")],
        )
        self.assertEqual(patched["order"], ["cmp_002"])
        self.assertNotIn("cmp_001", patched["compounds"])
        self.assertNotIn("Example A (2a)", text)
        self.assertIn("Example B (2b)", text)

    def test_patch_workflow_swaps_multiple_compound_pairs_and_preserves_visible_number_order(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            source_docx, source_manifest = _write_source_support(
                root,
                [
                    Compound(id="cmp_001", number="2a", name="Alpha"),
                    Compound(id="cmp_002", number="2b", name="Beta"),
                    Compound(id="cmp_003", number="3a", name="Gamma"),
                    Compound(id="cmp_004", number="3b", name="Delta"),
                ],
            )
            source_text_before = "\n".join(paragraph.text for paragraph in Document(source_docx).paragraphs)

            state = run_patch_si(
                PatchSIRequest(
                    manifest_path=source_manifest,
                    swap=(("2a", "3a"), ("2b", "3b")),
                    output_folder=root / "patches",
                )
            )
            patched_docx, patched_manifest, _ = _patch_artifact_paths(state)
            manifest = json.loads(patched_manifest.read_text(encoding="utf-8"))
            patched_text = "\n".join(paragraph.text for paragraph in Document(patched_docx).paragraphs)
            source_text_after = "\n".join(paragraph.text for paragraph in Document(source_docx).paragraphs)

        self.assertEqual(state["status"], "pass")
        self.assertEqual(manifest["order"], ["cmp_003", "cmp_004", "cmp_001", "cmp_002"])
        self.assertEqual(
            [manifest["compounds"][compound_id]["number"] for compound_id in manifest["order"]],
            ["2a", "2b", "3a", "3b"],
        )
        self.assertLess(patched_text.index("Gamma (2a)"), patched_text.index("Delta (2b)"))
        self.assertLess(patched_text.index("Delta (2b)"), patched_text.index("Alpha (3a)"))
        self.assertLess(patched_text.index("Alpha (3a)"), patched_text.index("Beta (3b)"))
        self.assertEqual(len(state["patch_result"]["swapped_pairs"]), 2)
        self.assertEqual(source_text_after, source_text_before)

    def test_swap_reorders_the_linked_1h_and_13c_spectrum_pages(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            image = root / "spectrum.png"
            image.write_bytes(
                base64.b64decode(
                    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
                )
            )
            compounds = [
                Compound(
                    id="cmp_001",
                    number="2a",
                    name="Alpha",
                    h1_image_path=str(image),
                    c13_image_path=str(image),
                ),
                Compound(
                    id="cmp_002",
                    number="3a",
                    name="Gamma",
                    h1_image_path=str(image),
                    c13_image_path=str(image),
                ),
            ]
            source_docx, source_manifest = _write_source_support(root, compounds)
            _clear_spectrum_structure_placeholders(source_docx)
            state = run_patch_si(
                PatchSIRequest(
                    manifest_path=source_manifest,
                    swap=(("2a", "3a"),),
                    output_folder=root / "patches",
                )
            )
            patched_docx = Path(state["artifacts"]["support_docx"])
            positions = _bookmark_positions(patched_docx)
            expected = [
                bookmark_name_for_block_id("spectrum:cmp_002:1H"),
                bookmark_name_for_block_id("spectrum:cmp_002:13C"),
                bookmark_name_for_block_id("spectrum:cmp_001:1H"),
                bookmark_name_for_block_id("spectrum:cmp_001:13C"),
            ]
            page_starts = _spectrum_page_starts(patched_docx, expected)
            trailing_blanks = _trailing_blank_paragraph_count(patched_docx)

        self.assertEqual(state["status"], "pass")
        self.assertTrue(all(name in positions for name in expected))
        self.assertEqual([positions[name] for name in expected], sorted(positions[name] for name in expected))
        self.assertEqual(page_starts, {name: True for name in expected})
        self.assertEqual(trailing_blanks, 0)

    def test_remove_deletes_the_linked_spectrum_pages(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            image = root / "spectrum.png"
            image.write_bytes(
                base64.b64decode(
                    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
                )
            )
            compounds = [
                Compound(
                    id="cmp_001",
                    number="2a",
                    name="Alpha",
                    h1_image_path=str(image),
                    c13_image_path=str(image),
                ),
                Compound(
                    id="cmp_002",
                    number="2b",
                    name="Beta",
                    h1_image_path=str(image),
                    c13_image_path=str(image),
                ),
            ]
            source_docx, source_manifest = _write_source_support(root, compounds)
            _clear_spectrum_structure_placeholders(source_docx)
            state = run_patch_si(
                PatchSIRequest(
                    manifest_path=source_manifest,
                    remove=("2a",),
                    output_folder=root / "patches",
                )
            )
            patched_docx = Path(state["artifacts"]["support_docx"])
            positions = _bookmark_positions(patched_docx)
            remaining_spectra = [
                bookmark_name_for_block_id("spectrum:cmp_002:1H"),
                bookmark_name_for_block_id("spectrum:cmp_002:13C"),
            ]
            page_starts = _spectrum_page_starts(patched_docx, remaining_spectra)
            trailing_blanks = _trailing_blank_paragraph_count(patched_docx)

        self.assertEqual(state["status"], "pass")
        self.assertNotIn(bookmark_name_for_block_id("compound:cmp_001"), positions)
        self.assertNotIn(bookmark_name_for_block_id("spectrum:cmp_001:1H"), positions)
        self.assertNotIn(bookmark_name_for_block_id("spectrum:cmp_001:13C"), positions)
        self.assertEqual(page_starts, {name: True for name in remaining_spectra})
        self.assertEqual(trailing_blanks, 0)

    def test_reorder_moves_the_linked_spectrum_pages(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            image = root / "spectrum.png"
            image.write_bytes(
                base64.b64decode(
                    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
                )
            )
            compounds = [
                Compound(
                    id="cmp_001",
                    number="2a",
                    name="Alpha",
                    h1_image_path=str(image),
                    c13_image_path=str(image),
                ),
                Compound(
                    id="cmp_002",
                    number="3a",
                    name="Gamma",
                    h1_image_path=str(image),
                    c13_image_path=str(image),
                ),
            ]
            source_docx, source_manifest = _write_source_support(root, compounds)
            _clear_spectrum_structure_placeholders(source_docx)
            state = run_patch_si(
                PatchSIRequest(
                    manifest_path=source_manifest,
                    reorder=("3a", "2a"),
                    output_folder=root / "patches",
                )
            )
            positions = _bookmark_positions(Path(state["artifacts"]["support_docx"]))
            expected = [
                bookmark_name_for_block_id("spectrum:cmp_002:1H"),
                bookmark_name_for_block_id("spectrum:cmp_002:13C"),
                bookmark_name_for_block_id("spectrum:cmp_001:1H"),
                bookmark_name_for_block_id("spectrum:cmp_001:13C"),
            ]
            page_starts = _spectrum_page_starts(Path(state["artifacts"]["support_docx"]), expected)
            trailing_blanks = _trailing_blank_paragraph_count(Path(state["artifacts"]["support_docx"]))

        self.assertEqual(state["status"], "pass")
        self.assertEqual([positions[name] for name in expected], sorted(positions[name] for name in expected))
        self.assertEqual(page_starts, {name: True for name in expected})
        self.assertEqual(trailing_blanks, 0)

    def test_patch_workflow_rejects_more_than_one_operation(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, source_manifest = _write_source_support(root)
            state = run_patch_si(
                PatchSIRequest(
                    manifest_path=source_manifest,
                    renumber={"2a": "3a"},
                    remove=("2a",),
                    output_folder=root / "patches",
                )
            )

        self.assertEqual(state["status"], "fail")
        self.assertIn("PATCH_OPERATION_COUNT_INVALID", {issue["code"] for issue in state["issues"]})

    def test_each_patch_creates_a_unique_run_folder(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, source_manifest = _write_source_support(root)
            request = PatchSIRequest(
                manifest_path=source_manifest,
                renumber={"2a": "3a"},
                output_folder=root / "patches",
            )
            first = run_patch_si(request)
            second = run_patch_si(request)

        self.assertNotEqual(first["artifacts"]["output_root"], second["artifacts"]["output_root"])

    def test_a_new_patch_can_use_the_previous_patch_manifest(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            _, source_manifest = _write_source_support(root)
            first = run_patch_si(
                PatchSIRequest(
                    manifest_path=source_manifest,
                    renumber={"2a": "3a"},
                    output_folder=root / "patches",
                )
            )
            first_manifest_path = Path(first["artifacts"]["manifest"])
            second = run_patch_si(
                PatchSIRequest(
                    manifest_path=first_manifest_path,
                    renumber={"3a": "4a"},
                    output_folder=root / "patches",
                )
            )
            first_manifest = json.loads(first_manifest_path.read_text(encoding="utf-8"))
            second_manifest = json.loads(Path(second["artifacts"]["manifest"]).read_text(encoding="utf-8"))

        self.assertEqual(first["status"], "pass")
        self.assertEqual(second["status"], "pass")
        self.assertEqual(first_manifest["compounds"]["cmp_001"]["number"], "3a")
        self.assertEqual(second_manifest["compounds"]["cmp_001"]["number"], "4a")
        self.assertNotEqual(first["artifacts"]["output_root"], second["artifacts"]["output_root"])


def _patch_artifact_paths(state: dict) -> tuple[Path, Path, Path]:
    artifacts = state["artifacts"]
    return Path(artifacts["support_docx"]), Path(artifacts["manifest"]), Path(artifacts["patch_report"])


def _bookmark_positions(docx_path: Path) -> dict[str, int]:
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(docx_path, "r") as source:
        root = ET.fromstring(source.read("word/document.xml"))
    body = root.find(f"{{{namespace}}}body")
    assert body is not None
    name_attr = f"{{{namespace}}}name"
    return {
        str(bookmark.attrib[name_attr]): index
        for index, child in enumerate(list(body))
        for bookmark in child.iter(f"{{{namespace}}}bookmarkStart")
        if name_attr in bookmark.attrib
    }


def _spectrum_page_starts(docx_path: Path, bookmark_names: list[str]) -> dict[str, bool]:
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(docx_path, "r") as source:
        root = ET.fromstring(source.read("word/document.xml"))
    body = root.find(f"{{{namespace}}}body")
    assert body is not None
    children = list(body)
    positions = _bookmark_positions(docx_path)
    return {
        name: positions[name] > 0 and _is_page_break_only(children[positions[name] - 1], namespace)
        for name in bookmark_names
    }


def _trailing_blank_paragraph_count(docx_path: Path) -> int:
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(docx_path, "r") as source:
        root = ET.fromstring(source.read("word/document.xml"))
    body = root.find(f"{{{namespace}}}body")
    assert body is not None
    children = list(body)
    if children and children[-1].tag == f"{{{namespace}}}sectPr":
        children.pop()
    count = 0
    for child in reversed(children):
        if not _is_empty_paragraph(child, namespace):
            break
        count += 1
    return count


def _is_page_break_only(element: ET.Element, namespace: str) -> bool:
    return _is_empty_paragraph(element, namespace) and any(
        item.attrib.get(f"{{{namespace}}}type") == "page"
        for item in element.iter(f"{{{namespace}}}br")
    )


def _is_empty_paragraph(element: ET.Element, namespace: str) -> bool:
    if element.tag != f"{{{namespace}}}p":
        return False
    if any((item.text or "").strip() for item in element.iter(f"{{{namespace}}}t")):
        return False
    for tag in ("drawing", "object", "pict", "sym", "fldSimple"):
        if element.find(f".//{{{namespace}}}{tag}") is not None:
            return False
    return True


def _clear_spectrum_structure_placeholders(docx_path: Path) -> None:
    with zipfile.ZipFile(docx_path, "r") as source:
        members = {item.filename: source.read(item.filename) for item in source.infolist()}
    root = ET.fromstring(members["word/document.xml"])
    namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for text_node in root.iter(f"{{{namespace}}}t"):
        if text_node.text and "[[SPECTRUM_STRUCTURE:" in text_node.text:
            text_node.text = ""
    members["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as target:
        for name, data in members.items():
            target.writestr(name, data)


def _write_source_support(root: Path, compounds: list[Compound] | None = None) -> tuple[Path, Path]:
    source_docx = root / "support_information.docx"
    source_manifest = root / "support_information.manifest.json"
    compounds = compounds or [Compound(id="cmp_001", number="2a", name="Example")]
    build_document_from_model(build_si_document_model(compounds), source_docx)
    compound_entries = {}
    for index, compound in enumerate(compounds, start=2):
        compound_entries[compound.id] = {
            "id": compound.id,
            "number": compound.number,
            "domain_snapshot": compound.to_domain_dict(),
            "source_row": index,
            "structure_placeholder": f"STRUCTURE:{compound.number}",
            "docx_block_id": f"compound:{compound.id}",
            "docx_bookmark": bookmark_name_for_block_id(f"compound:{compound.id}"),
            "artifacts": {},
        }
    manifest = {
        "run_id": "run",
        "artifacts": {"support_docx": str(source_docx), "manifest": str(source_manifest)},
        "output_paths": {"support_docx": str(source_docx), "manifest": str(source_manifest)},
        "order": [compound.id for compound in compounds],
        "compounds": compound_entries,
    }
    source_manifest.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return source_docx, source_manifest


if __name__ == "__main__":
    unittest.main()
