from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path
from zipfile import ZipFile

from docx import Document

from si_generator.docx_builder import build_document_from_model
from si_generator.domain.bookmarks import bookmark_name_for_block_id
from si_generator.domain.requests import AddCompoundsRequest
from si_generator.graph.nodes.add_compounds import _append_generated_docx_blocks, _new_compound_id_map, resolve_add_method_config_node
from si_generator.domain.compound import Compound
from si_generator.render.document_model import build_si_document_model
from si_generator.workflows.add_compounds import run_add_compounds


class AddCompoundsWorkflowTests(unittest.TestCase):
    def test_duplicate_compound_number_blocks_before_output_mutation(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = _write_manifest(root, number="2a")
            new_table = root / "new_compounds.csv"
            new_table.write_text("number,name\n2a,Duplicate\n", encoding="utf-8")
            output_docx = root / "patched_support.docx"

            state = run_add_compounds(
                AddCompoundsRequest(
                    manifest_path=manifest_path,
                    input_path=new_table,
                    input_kind="csv",
                    output_docx=output_docx,
                )
            )
            report = json.loads(output_docx.with_suffix(".add_report.json").read_text(encoding="utf-8"))

        self.assertEqual(state["status"], "fail")
        self.assertIn("DUPLICATE_COMPOUND_NUMBER", {issue["code"] for issue in state["issues"]})
        self.assertEqual(state["add_result"]["duplicate_numbers"], ["2a"])
        self.assertFalse(output_docx.exists())
        self.assertNotIn("generated_support_docx", state.get("artifacts", {}))
        self.assertEqual(report["status"], "fail")
        self.assertEqual(report["add_result"]["duplicate_numbers"], ["2a"])

    def test_non_duplicate_adds_docx_blocks_and_manifest_entry(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = _write_manifest(root, number="2a", valid_docx=True)
            new_table = root / "new_compounds.csv"
            new_table.write_text("number,name,color,state\n2b,Added,white,solid\n", encoding="utf-8")
            output_docx = root / "patched_support.docx"

            state = run_add_compounds(
                AddCompoundsRequest(
                    manifest_path=manifest_path,
                    input_path=new_table,
                    input_kind="csv",
                    output_docx=output_docx,
                    method_mode="new_method",
                    no_extract_nmr=True,
                    no_check_support=True,
                )
            )
            merged_manifest = json.loads(output_docx.with_suffix(".manifest.json").read_text(encoding="utf-8"))
            report = json.loads(output_docx.with_suffix(".add_report.json").read_text(encoding="utf-8"))
            text = "\n".join(paragraph.text for paragraph in Document(output_docx).paragraphs)
            output_exists = output_docx.exists()
            document_xml = _docx_text(output_docx, "word/document.xml")

        self.assertEqual(state["status"], "pass")
        self.assertTrue(output_exists)
        self.assertEqual([merged_manifest["compounds"][compound_id]["number"] for compound_id in merged_manifest["order"]], ["2a", "2b"])
        self.assertIn("Existing compound (2a)", text)
        self.assertIn("Added (2b)", text)
        self.assertTrue(merged_manifest["compounds"]["added_cmp_001_1"]["docx_bookmark"])
        self.assertIn(merged_manifest["compounds"]["added_cmp_001_1"]["docx_bookmark"], document_xml)
        self.assertEqual(report["status"], "pass")
        self.assertNotIn("ADD_COMPOUNDS_TEXT_ONLY_MERGE", {issue["code"] for issue in report["issues"]})
        self.assertEqual(report["method_mode"], "new_method")

    def test_output_folder_creates_per_run_add_output(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = _write_manifest(root, number="2a", valid_docx=True)
            new_table = root / "new_compounds.csv"
            new_table.write_text("number,name,color,state\n2b,Added,white,solid\n", encoding="utf-8")
            output_folder = root / "add_output"

            state = run_add_compounds(
                AddCompoundsRequest(
                    manifest_path=manifest_path,
                    input_path=new_table,
                    input_kind="csv",
                    output_folder=output_folder,
                    method_mode="new_method",
                    no_extract_nmr=True,
                    no_check_support=True,
                )
            )
            output_docx = Path(state["artifacts"]["support_docx"])
            report = json.loads(Path(state["artifacts"]["add_report"]).read_text(encoding="utf-8"))

            self.assertEqual(state["status"], "pass")
            self.assertEqual(output_docx.name, "support_information.docx")
            self.assertEqual(output_docx.parent.name, "docx")
            self.assertEqual(output_docx.parent.parent.parent, output_folder / "runs")
            self.assertTrue(output_docx.exists())
            self.assertEqual(report["output_docx"], str(output_docx))
            self.assertTrue((output_docx.parent / "support_information.manifest.json").exists())

    def test_same_series_requires_manifest_run_config(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = _write_manifest(root, number="2a", valid_docx=True)
            new_table = root / "new_compounds.csv"
            new_table.write_text("number,name,color,state\n2b,Added,white,solid\n", encoding="utf-8")
            output_docx = root / "patched_support.docx"

            state = run_add_compounds(
                AddCompoundsRequest(
                    manifest_path=manifest_path,
                    input_path=new_table,
                    input_kind="csv",
                    output_docx=output_docx,
                )
            )
            report = json.loads(output_docx.with_suffix(".add_report.json").read_text(encoding="utf-8"))

        self.assertEqual(state["status"], "fail")
        self.assertIn("ADD_RUN_CONFIG_MISSING", {issue["code"] for issue in state["issues"]})
        self.assertFalse(output_docx.exists())
        self.assertEqual(report["method_mode"], "same_series")

    def test_same_series_reuses_manifest_run_config(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = _write_manifest(
                root,
                number="2a",
                valid_docx=True,
                run_config={
                    "version": 1,
                    "no_extract_nmr": True,
                    "insert_spectra_as": "none",
                    "no_check_support": True,
                },
            )
            new_table = root / "new_compounds.csv"
            new_table.write_text("number,name,color,state\n2b,Added,white,solid\n", encoding="utf-8")
            output_docx = root / "patched_support.docx"

            state = run_add_compounds(
                AddCompoundsRequest(
                    manifest_path=manifest_path,
                    input_path=new_table,
                    input_kind="csv",
                    output_docx=output_docx,
                )
            )
            merged_manifest = json.loads(output_docx.with_suffix(".manifest.json").read_text(encoding="utf-8"))
            report = json.loads(output_docx.with_suffix(".add_report.json").read_text(encoding="utf-8"))

        self.assertEqual(state["status"], "pass")
        self.assertEqual(report["method_mode"], "same_series")
        self.assertEqual(report["method_config"]["source"], "manifest")
        self.assertTrue(report["method_config"]["no_extract_nmr"])
        self.assertEqual(merged_manifest["add_compounds_history"][-1]["method_mode"], "same_series")
        self.assertEqual(merged_manifest["compounds"]["added_cmp_001_1"]["add_method_mode"], "same_series")

    def test_same_series_does_not_reuse_old_scope_file(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            old_schema = root / "Reaction_schema.docx"
            old_scope = root / "Scope.docx"
            old_schema.write_text("schema", encoding="utf-8")
            manifest_path = _write_manifest(
                root,
                number="2a",
                valid_docx=True,
                run_config={
                    "version": 1,
                    "no_extract_nmr": True,
                    "insert_spectra_as": "none",
                    "generate_loadings": True,
                    "loadings_schema_docx": str(old_schema),
                    "loadings_scope_docx": str(old_scope),
                    "no_check_support": True,
                },
            )
            new_table = root / "new_compounds.csv"
            new_table.write_text("number,name,color,state\n2b,Added,white,solid\n", encoding="utf-8")
            output_docx = root / "patched_support.docx"

            state = run_add_compounds(
                AddCompoundsRequest(
                    manifest_path=manifest_path,
                    input_path=new_table,
                    input_kind="csv",
                    output_docx=output_docx,
                )
            )
            report = json.loads(output_docx.with_suffix(".add_report.json").read_text(encoding="utf-8"))

        self.assertEqual(state["status"], "fail")
        issue_codes = {issue["code"] for issue in state["issues"]}
        self.assertIn("ADD_LOADINGS_CONFIG_MISSING", issue_codes)
        self.assertNotIn("ADD_RUN_CONFIG_FILE_MISSING", issue_codes)
        self.assertEqual(report["method_config"]["source"], "manifest")
        self.assertEqual(report["method_config"]["loadings_schema_docx"], str(old_schema))
        self.assertEqual(report["method_config"]["loadings_scope_docx"], None)

    def test_same_series_errors_when_scope_does_not_match_new_compounds(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            schema = root / "Reaction_schema.docx"
            schema.write_text("schema", encoding="utf-8")
            old_scope = root / "Scope_1.docx"
            _write_scope_docx(old_scope, ["2a"])
            new_dir = root / "new"
            new_dir.mkdir()
            new_table = new_dir / "test_input_2.docx"
            new_table.write_text("table", encoding="utf-8")
            new_scope = new_dir / "Scope_2.docx"
            _write_scope_docx(new_scope, ["2b"])
            manifest_path = root / "support_information.manifest.json"
            manifest_path.write_text("{}", encoding="utf-8")

            result = resolve_add_method_config_node(
                {
                    "request": AddCompoundsRequest(
                        manifest_path=manifest_path,
                        input_path=new_table,
                        input_kind="word",
                        output_docx=root / "out.docx",
                        loadings_scope_docx=old_scope,
                    ),
                    "manifest": {
                        "run_config": {
                            "version": 1,
                            "generate_loadings": True,
                            "loadings_schema_docx": str(schema),
                            "loadings_scope_docx": str(old_scope),
                            "no_extract_nmr": True,
                            "insert_spectra_as": "none",
                            "no_check_support": True,
                        }
                    },
                    "new_compounds": [Compound(number="2b", name="Added")],
                    "issues": [],
                }
            )

        self.assertEqual(result["status"], "fail")
        self.assertEqual(result["add_method_config"]["loadings_scope_docx"], old_scope)
        issue_codes = {issue["code"] for issue in result["issues"]}
        self.assertIn("ADD_LOADINGS_SCOPE_INPUT_MISMATCH", issue_codes)
        mismatch = next(issue for issue in result["issues"] if issue["code"] == "ADD_LOADINGS_SCOPE_INPUT_MISMATCH")
        self.assertEqual(mismatch["severity"], "error")
        self.assertIn("New compounds: 2b", mismatch["message"])
        self.assertIn("Scope products: 2a", mismatch["message"])

    def test_new_method_preserves_spectrum_settings_from_manifest(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            manifest_path = _write_manifest(
                root,
                number="2a",
                valid_docx=True,
                run_config={
                    "version": 1,
                    "no_extract_nmr": True,
                    "insert_spectra_as": "none",
                    "target_signal_height_fraction": 0.65,
                    "x_range_ppm_1h": [0, 10],
                    "x_range_ppm_13c": [0, 180],
                    "no_check_support": True,
                },
            )
            new_table = root / "new_compounds.csv"
            new_table.write_text("number,name,color,state\n2b,Added,white,solid\n", encoding="utf-8")
            output_docx = root / "patched_support.docx"

            state = run_add_compounds(
                AddCompoundsRequest(
                    manifest_path=manifest_path,
                    input_path=new_table,
                    input_kind="csv",
                    output_docx=output_docx,
                    method_mode="new_method",
                    insert_spectra_as="png",
                    target_signal_height_fraction=0.9,
                    x_range_ppm_1h=(-1, 12),
                    x_range_ppm_13c=(-10, 210),
                )
            )
            report = json.loads(output_docx.with_suffix(".add_report.json").read_text(encoding="utf-8"))

        self.assertEqual(state["status"], "pass")
        self.assertEqual(report["method_config"]["source"], "manifest+request")
        self.assertEqual(report["method_config"]["insert_spectra_as"], "none")
        self.assertEqual(report["method_config"]["target_signal_height_fraction"], 0.65)
        self.assertEqual(report["method_config"]["x_range_ppm_1h"], [0.0, 10.0])
        self.assertEqual(report["method_config"]["x_range_ppm_13c"], [0.0, 180.0])

    def test_docx_block_merge_preserves_media_and_remaps_bookmarks(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            root = Path(tmp)
            image_path = root / "spectrum.png"
            image_path.write_bytes(_tiny_png())
            old_docx = root / "old.docx"
            new_docx = root / "new.docx"
            output_docx = root / "merged.docx"

            old_manifest = _build_support_docx_and_manifest(old_docx, Compound(id="cmp_001", number="2a", name="Existing"))
            new_manifest = _build_support_docx_and_manifest(
                new_docx,
                Compound(
                    id="cmp_001",
                    number="2b",
                    name="Added",
                    h1_spectrum_path="fid",
                    h1_image_path=str(image_path),
                ),
                spectra_embed_mode="png",
            )
            id_map = _new_compound_id_map(old_manifest, new_manifest)

            _append_generated_docx_blocks(old_docx, new_docx, output_docx, old_manifest=old_manifest, new_manifest=new_manifest, id_map=id_map)

            document_xml = _docx_text(output_docx, "word/document.xml")
            rels_xml = _docx_text(output_docx, "word/_rels/document.xml.rels")
            with ZipFile(output_docx) as archive:
                media_files = [name for name in archive.namelist() if name.startswith("word/media/")]

        added_id = "added_cmp_001_1"
        self.assertEqual(id_map, {"cmp_001": added_id})
        self.assertIn(bookmark_name_for_block_id("compound:cmp_001"), document_xml)
        self.assertIn(bookmark_name_for_block_id(f"compound:{added_id}"), document_xml)
        self.assertIn(bookmark_name_for_block_id(f"spectrum:{added_id}:1H"), document_xml)
        self.assertIn("Added", document_xml)
        self.assertIn("add_compounds_1.png", rels_xml)
        self.assertIn("word/media/add_compounds_1.png", media_files)


def _write_scope_docx(path: Path, numbers: list[str]) -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.rows[0].cells[0].text = "Product Number"
    for number in numbers:
        table.add_row().cells[0].text = number
    document.save(path)


def _write_manifest(root: Path, *, number: str, valid_docx: bool = False, run_config: dict | None = None) -> Path:
    support_docx = root / "support_information.docx"
    if valid_docx:
        _build_support_docx_and_manifest(support_docx, Compound(id="cmp_001", number=number, name="Existing compound"))
    else:
        support_docx.write_bytes(b"placeholder")
    manifest_path = root / "support_information.manifest.json"
    manifest = {
        "run_id": "run",
        "artifacts": {"support_docx": str(support_docx), "manifest": str(manifest_path)},
        "output_paths": {"support_docx": str(support_docx), "manifest": str(manifest_path)},
        "order": ["cmp_001"],
        "compounds": {
            "cmp_001": {
                "id": "cmp_001",
                "number": number,
                "docx_block_id": "compound:cmp_001",
                "docx_bookmark": bookmark_name_for_block_id("compound:cmp_001"),
            }
        },
    }
    if run_config is not None:
        manifest["run_config"] = run_config
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest_path


def _build_support_docx_and_manifest(docx_path: Path, compound: Compound, *, spectra_embed_mode: str = "none") -> dict:
    model = build_si_document_model([compound], spectra_embed_mode=spectra_embed_mode)
    build_document_from_model(model, docx_path)
    compound_id = compound.id or "cmp_001"
    return {
        "run_id": "run",
        "artifacts": {"support_docx": str(docx_path)},
        "output_paths": {"support_docx": str(docx_path)},
        "order": [compound_id],
        "compounds": {
            compound_id: {
                "id": compound_id,
                "number": compound.number,
                "docx_block_id": f"compound:{compound_id}",
                "docx_bookmark": bookmark_name_for_block_id(f"compound:{compound_id}"),
                "artifacts": {"h1_png": compound.h1_image_path} if compound.h1_image_path else {},
            }
        },
    }


def _docx_text(docx_path: Path, member: str) -> str:
    with ZipFile(docx_path) as archive:
        return archive.read(member).decode("utf-8")


def _tiny_png() -> bytes:
    return (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
        b"\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4"
        b"\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05"
        b"\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
    )


if __name__ == "__main__":
    unittest.main()
